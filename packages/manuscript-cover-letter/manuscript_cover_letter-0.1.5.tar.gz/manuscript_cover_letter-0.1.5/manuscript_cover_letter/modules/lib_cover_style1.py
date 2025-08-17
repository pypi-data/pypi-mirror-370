#!/usr/bin/env python3
# pip install python-docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import re
import locale




def generate_cover_style1( filename, dict_info ):

    generate_cover_style1_raw(  dict_info["title"],
                                dict_info["journal"],
                                dict_info["complete_name"],
                                dict_info["university"],
                                dict_info["country"],
                                dict_info["address"],
                                dict_info["emails"],
                                dict_info["telephone"],
                                dict_info["summary"],
                                filename )    

def generate_cover_style1_raw(  title,
                                journal,
                                complete_name,
                                university,
                                country,
                                address,
                                emails,
                                telephone,
                                summary,
                                filename ):

    # Criação do documento
    doc = Document()
    
    # Defina a fonte padrão para o documento
    estilo_padrao = doc.styles['Normal']
    fonte_padrao = estilo_padrao.font
    fonte_padrao.name = 'Times New Roman'
    fonte_padrao.size = Pt(12)

    # Função para adicionar parágrafo com estilo
    def add_paragraph(doc, text, bold=False, align=None, font_size=12):
        p = doc.add_paragraph()
        
        # Se o texto tiver partes em **negrito**, processamos por partes
        if "**" in text:
            partes = re.split(r'(\*\*.*?\*\*)', text)
            for parte in partes:
                if parte.startswith("**") and parte.endswith("**"):
                    run = p.add_run(parte[2:-2])
                    run.bold = True
                else:
                    run = p.add_run(parte)
                    run.bold = bold
                run.font.size = Pt(font_size)
        else:
            # Caso não haja marcações com **, aplica o bold global
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(font_size)
        
        if align:
            p.alignment = align
        return p

    locale.setlocale(locale.LC_TIME, 'en_US.UTF-8')
    data = datetime.now()
    data_formatada = data.strftime(f"%-d %B of %Y") 

    # Cabeçalho com alinhamento à direita
    add_paragraph(doc, f"{country}, {data_formatada}", align=WD_PARAGRAPH_ALIGNMENT.RIGHT)

    # Destinatário
    add_paragraph(doc, "To")
    add_paragraph(doc, f"{journal} Editor")
    add_paragraph(doc, "**Ref. Cover letter:** New Manuscript")
    add_paragraph(doc, f"**Title:** {title}",align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY)

    # Saudação
    add_paragraph(doc, "Dear Editor")

    # Corpo da carta
    add_paragraph(
        doc,
        f"We are pleased to submit our manuscript entitled **\"{title}\"** for consideration for publication in {journal}.",
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    )

    add_paragraph(
        doc,
        summary,
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    )


    # Autor correspondente
    add_paragraph(doc, "**Corresponding author:**")
    data=""
    informations = [complete_name,university, address, emails, telephone]
    for l,info in enumerate(informations):
        if info !="":
            data+=info
            if (l+1)!=len(informations):
                data+="\n"
    add_paragraph(doc, data,align=WD_PARAGRAPH_ALIGNMENT.RIGHT)


    # Declaração de originalidade
    add_paragraph(
        doc,
        f"The manuscript is original, has not been published elsewhere, and is not under consideration by any other journal. All authors have read and approved the manuscript and have no conflicts of interest to declare.",
        align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    )

    add_paragraph(doc, "We appreciate your time and consideration, and we look forward to hearing from you.")
    # Despedida
    add_paragraph(doc, "Sincerely,\n" + complete_name + "\n" + university)

    # Salvar o documento
    doc.save(filename)

if __name__ == '__main__':
    title = "Selection of Statistical Indices in the Biospeckle Laser Analysis Regarding Filtering Actions"
    journal = "Optics Communications"
    complete_name = "Fernando Pujaico Rivera"
    university = "Universidade Federal de Lavras"
    country = "Brazil"
    address = "Lavras MG, CEP 37.200-000, Caixa Postal 3037"
    emails = "ferrnando.pujaico.rivera@gmail.com\n201518201@posgrad.ufla.br"
    telephone = "Tel. +55 35 3829 1672 Office"
    summary = "We tested some traditional methods to analyze dynamic laser speckle regarding their filtering actions. Additionally, we propose two new biospeckle indices based in the binary entropy, including one that avoids filtering the original signals. The work was based on theoretical developments and was validated using a drying paint monitoring test. We proved that the dynamic laser speckle, or biospeckle, are compromised by filtering actions and that it is possible to elect an index that does not provide filtering."

    generate_cover_style1_raw(  title,
                                journal,
                                complete_name,
                                university,
                                country,
                                address,
                                emails,
                                telephone,
                                summary,
                                "salida.docx" )
