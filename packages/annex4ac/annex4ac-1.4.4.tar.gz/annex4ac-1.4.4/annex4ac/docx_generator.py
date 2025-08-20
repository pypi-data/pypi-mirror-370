"""
docx_generator.py

Module for generating Annex IV DOCX documents.
Contains all the logic for working with python-docx to create technical documentation.
"""

import re
import os
import yaml
from pathlib import Path
from datetime import datetime
from hashlib import sha256
from typing import Dict

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from ftfy import fix_text

# Regular expressions for parsing lists
BULLET_RE = re.compile(r'^\s*(?:[\u2022\u25CF\u25AA\-\*])\s+')
SUBPOINT_RE = re.compile(r'^\(([a-z])\)\s+', re.I)  # (a), (b)...
PARA_SPLIT = re.compile(r'\n{2,}')  # paragraphs

from .constants import DOC_CTRL_FIELDS, SECTION_MAPPING, SCHEMA_VERSION


def _enable_auto_update_fields(doc):
    """Enables automatic field updates in the document."""
    settings = doc.settings._element
    upd = settings.find(qn('w:updateFields'))
    if upd is None:
        upd = OxmlElement('w:updateFields')
        upd.set(qn('w:val'), 'true')
        settings.append(upd)


def _ensure_toc_styles(doc):
    """Creates styles for table of contents."""
    if 'TOC Heading' not in doc.styles:
        st = doc.styles.add_style('TOC Heading', WD_STYLE_TYPE.PARAGRAPH)
        st.base_style = doc.styles['Heading 1']
        st.font.name = 'Times New Roman'
        st.font.size = Pt(14)
    # Fix fonts for built-in TOC 1..9
    for i in range(1, 10):
        name = f'TOC {i}'
        if name in doc.styles:
            st = doc.styles[name]
            st.font.name = 'Times New Roman'
            st.font.size = Pt(12)





def _add_page_number(paragraph):
    """Adds page number to paragraph."""
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
    paragraph._p.append(fld)


def _new_alpha_list(doc):
    """Creates new numbering for lists with letters (a), (b), (c)..."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering
    nid = str(max([int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))] or [0]) + 1)

    absNum = OxmlElement('w:abstractNum')
    absNum.set(qn('w:abstractNumId'), nid)
    lvl = OxmlElement('w:lvl')
    lvl.set(qn('w:ilvl'), '0')
    fmt = OxmlElement('w:numFmt')
    fmt.set(qn('w:val'), 'lowerLetter')
    lvl.append(fmt)
    text = OxmlElement('w:lvlText')
    text.set(qn('w:val'), '(%1)')
    lvl.append(text)
    suff = OxmlElement('w:suff')
    suff.set(qn('w:val'), 'space')
    lvl.append(suff)
    start = OxmlElement('w:start')
    start.set(qn('w:val'), '1')
    lvl.append(start)
    # optional restart each section
    restart = OxmlElement('w:lvlRestart')
    restart.set(qn('w:val'), '1')
    lvl.append(restart)

    absNum.append(lvl)
    numbering.append(absNum)

    num = OxmlElement('w:num')
    num.set(qn('w:numId'), nid)
    abs_ref = OxmlElement('w:abstractNumId')
    abs_ref.set(qn('w:val'), nid)
    num.append(abs_ref)
    numbering.append(num)
    return int(nid)


def _is_last_in_block(line: str, para: str) -> bool:
    """Determines if the line is the last in the paragraph block."""
    lines = para.split('\n')
    return line.strip() == lines[-1].strip()


def _fix_escapes(text: str) -> str:
    """Unescapes \n and normalizes line breaks"""
    if not isinstance(text, str):
        return text
    # first normalize CRLF
    text = text.replace('\\r\\n', '\n').replace('\\r', '\n')
    # then literal \n -> real line breaks
    return text.replace('\\n', '\n')


def _apply_indent(p, left=720, hanging=360):
    """Applies indentation to paragraph"""
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))     # 720 twips ≈ 0.5"
    ind.set(qn('w:hanging'), str(hanging))  # 360 twips ≈ 0.25"
    pPr = p._p.get_or_add_pPr()
    pPr.append(ind)


def _apply_numbering(p, num_id, left=720, hanging=360):
    """Applies numbering to paragraph"""
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.extend([ilvl, numId])
    pPr = p._p.get_or_add_pPr()
    pPr.append(numPr)

    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left))
    ind.set(qn('w:hanging'), str(hanging))
    pPr.append(ind)


def render_docx(payload: dict, output_path: Path, meta: dict):
    """
    Main function for generating Annex IV DOCX document.
    
    Args:
        payload: Dictionary with data to fill sections
        output_path: Path to save DOCX file
    """
    doc = Document()

    # --- Auto-update fields and TOC styles ---
    _enable_auto_update_fields(doc)
    _ensure_toc_styles(doc)

    # --- Page margins ---
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(2.5)
        s.top_margin = s.bottom_margin = Cm(2.5)

    # --- Base style (TNR 12 pt, 1.5 line) ---
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.paragraph_format.line_spacing = 1.5
    if not normal.font.name:
        normal.font.name = 'Liberation Serif'   # Linux fallback

    # Headings
    for name, size in (('Title', 18), ('Heading 1', 16), ('Heading 2', 14)):
        st = doc.styles[name]
        st.font.name = 'Times New Roman'
        st.font.size = Pt(size)
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Create list styles if they don't exist
    if 'List Bullet' not in doc.styles:
        bullet_style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Times New Roman'
        bullet_style.font.size = Pt(12)
        bullet_style.paragraph_format.left_indent = Pt(18)
        bullet_style.paragraph_format.space_after = Pt(6)

    # Metadata
    doc.core_properties.author = "Annex4AC"
    doc.core_properties.created = datetime.now()
    doc.core_properties.version = str(payload.get('_schema_version', 'unknown'))
    
    # Add retention period information
    retention_until = payload.get('retention_until')
    if retention_until:
        doc.core_properties.comments = f"Generated by Annex4AC — EU AI Act Annex IV tool. Retention until: {retention_until}"
    else:
        doc.core_properties.comments = "Generated by Annex4AC — EU AI Act Annex IV tool"
    
    doc.core_properties.title = "Annex IV Technical Documentation"
    doc.core_properties.subject = "EU AI Act Compliance"
    doc.core_properties.identifier = f"annex4-{payload.get('_schema_version', 'unknown')}"
    try:
        doc.part.core_properties.category = "Annex IV Tech Doc"
        doc.part.core_properties.keywords = sha256(
            yaml.safe_dump(payload, allow_unicode=True).encode('utf-8')
        ).hexdigest()
    except Exception:
        pass

    # --- Title page ---
    title = doc.add_heading('Annex IV Technical Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Document control block ---
    # Use passed metadata (ensure generation_date is current)
    meta = meta.copy()
    meta["generation_date"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    doc.add_heading('Document control', level=1)
    table = doc.add_table(rows=0, cols=2)
    table.autofit = True
    for label, key in DOC_CTRL_FIELDS:
        row_cells = table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = str(meta.get(key, "—"))
    doc.add_paragraph()  # Empty line after table


    # --- Main Annex IV sections ---
    # bookmark_counter = 0
    for heading, key in SECTION_MAPPING:
        raw = payload.get(key, "")
        if not raw:
            continue
        raw = _fix_escapes(raw)
        raw = fix_text(raw)
        raw = re.sub(r'\r?\n', '\n', raw)  # unification
        # If YAML "folded" line breaks before lists — restore them
        raw = re.sub(r'\s+(?=(?:[-•*]\s))', '\n', raw)
        raw = re.sub(r'\s+(?=\([a-z]\)\s+)', '\n', raw, flags=re.I)

        # create heading with unique bookmark for navigation
        heading_para = doc.add_heading(heading, level=1)
        # bookmark_name = f"section_{key}"
        # bookmark_counter += 1
        # bookmark = OxmlElement('w:bookmarkStart')
        # bookmark.set(qn('w:id'), str(bookmark_counter))
        # bookmark.set(qn('w:name'), bookmark_name)
        # heading_para._p.append(bookmark)
        # add bookmark end
        # bookmark_end = OxmlElement('w:bookmarkEnd')
        # bookmark_end.set(qn('w:id'), str(bookmark_counter))
        # heading_para._p.append(bookmark_end)

        for para in re.split(r'\n{2,}', raw):  # paragraph = 2+ line breaks
            if not para.strip():
                continue
            
            mode = None               # None | 'ul' | 'ol'
            alpha_id = None
            for line in para.split('\n'):
                if SUBPOINT_RE.match(line):
                    txt = SUBPOINT_RE.sub('', line, 1).strip()
                    if alpha_id is None:
                        alpha_id = _new_alpha_list(doc)
                    # add ;/. automatically, but not after colon
                    txt_clean = txt.rstrip(' ;.')
                    if txt_clean.endswith(':'):
                        suffix = ''
                    else:
                        suffix = '.' if _is_last_in_block(line, para) else ';'
                    p = doc.add_paragraph(txt_clean + suffix)
                    _apply_numbering(p, alpha_id)      # sets w:numPr + indentation
                    mode = 'ol'
                elif BULLET_RE.match(line):
                    txt = BULLET_RE.sub('', line, 1).strip()
                    # add ;/. automatically, but not after colon
                    txt_clean = txt.rstrip(' ;.')
                    if txt_clean.endswith(':'):
                        suffix = ''
                    else:
                        suffix = '.' if _is_last_in_block(line, para) else ';'
                    p = doc.add_paragraph(txt_clean + suffix, style='List Bullet')
                    _apply_indent(p, left=720, hanging=360)   # same indentation as ol
                    mode = 'ul'
                else:
                    # regular text with line breaks
                    p = doc.add_paragraph()
                    for i, chunk in enumerate(line.split('\n')):
                        if i:
                            p.add_run().add_break()
                        p.add_run(chunk)

    # --- Footer with page number ---
    for s in doc.sections:
        footer_p = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(footer_p)

    doc.save(output_path)

if __name__ == "__main__":
    import argparse, json
    ap = argparse.ArgumentParser()
    ap.add_argument("payload_json")
    ap.add_argument("output_docx")
    args = ap.parse_args()
    payload = json.load(open(args.payload_json))
    # Build metadata for standalone usage
    from constants import DOC_CTRL_FIELDS, SCHEMA_VERSION
    meta = {}
    for label, key in DOC_CTRL_FIELDS:
        if key == "generation_date":
            meta[key] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        elif key == "_schema_version":
            sv = payload.get(key, SCHEMA_VERSION)
            if isinstance(sv, str) and len(sv) == 8 and sv.isdigit():
                sv = f"{sv[:4]}-{sv[4:6]}-{sv[6:8]}"
            meta[key] = sv
        else:
            meta[key] = payload.get(key, "—")
    render_docx(payload, Path(args.output_docx), meta) 