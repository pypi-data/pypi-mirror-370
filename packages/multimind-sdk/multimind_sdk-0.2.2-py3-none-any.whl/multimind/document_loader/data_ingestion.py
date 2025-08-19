"""
Advanced data ingestion module supporting multiple document types and real-time ingestion.
"""

from typing import List, Dict, Any, Optional, Union, Tuple, Protocol, runtime_checkable
from dataclasses import dataclass
from enum import Enum
import asyncio
import json
import csv
import io
from datetime import datetime
import aiohttp
try:
    import aiofiles
except ImportError:
    aiofiles = None
try:
    import pdfplumber
except ImportError:
    pdfplumber = None
try:
    from docx import Document
except ImportError:
    Document = None
try:
    import html2text
except ImportError:
    html2text = None
try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None
try:
    import pandas as pd
except ImportError:
    pd = None
try:
    from kafka import KafkaConsumer
except ImportError:
    KafkaConsumer = None
try:
    from notion_client import Client as NotionClient
except ImportError:
    NotionClient = None
try:
    from google.oauth2.credentials import Credentials
    from googleapiclient.discovery import build
except ImportError:
    Credentials = None
    build = None
from ..models.base import BaseLLM

@dataclass
class DocumentMetadata:
    """Metadata for ingested documents."""
    source: str
    source_type: str
    timestamp: float
    content_type: str
    language: str
    size: int
    author: Optional[str]
    created_at: Optional[float]
    modified_at: Optional[float]
    custom_metadata: Dict[str, Any]

@dataclass
class IngestedDocument:
    """Represents an ingested document."""
    content: str
    metadata: DocumentMetadata
    chunks: List[Dict[str, Any]]
    raw_content: Optional[Any]
    embeddings: Optional[List[float]]

class SourceType(Enum):
    """Types of document sources."""
    FILE = "file"
    WEB = "web"
    API = "api"
    DATABASE = "database"
    STREAM = "stream"
    NOTION = "notion"
    GOOGLE_DOCS = "google_docs"

class DocumentType(Enum):
    """Types of documents."""
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    TXT = "txt"
    JSON = "json"
    CSV = "csv"
    MARKDOWN = "markdown"
    UNKNOWN = "unknown"

class DataIngestion:
    """Advanced data ingestion system."""

    def __init__(
        self,
        model: BaseLLM,
        notion_token: Optional[str] = None,
        google_credentials: Optional[Dict[str, Any]] = None,
        kafka_config: Optional[Dict[str, Any]] = None,
        **kwargs
    ):
        """
        Initialize data ingestion system.
        
        Args:
            model: Language model for content analysis
            notion_token: Optional Notion API token
            google_credentials: Optional Google API credentials
            kafka_config: Optional Kafka configuration
            **kwargs: Additional parameters
        """
        self.model = model
        self.notion_client = NotionClient(auth=notion_token) if notion_token else None
        self.google_credentials = google_credentials
        self.kafka_config = kafka_config
        self.kwargs = kwargs
        
        # Initialize HTML converter
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = False
        self.html_converter.ignore_images = False
        
        # Initialize session for web requests
        self.session = None

    async def ingest_document(
        self,
        source: str,
        source_type: SourceType,
        **kwargs
    ) -> IngestedDocument:
        """
        Ingest document from source.
        
        Args:
            source: Document source (file path, URL, etc.)
            source_type: Type of source
            **kwargs: Additional parameters
            
        Returns:
            Ingested document
        """
        # Initialize session if needed
        if self.session is None:
            self.session = aiohttp.ClientSession()
        
        try:
            # Get content based on source type
            if source_type == SourceType.FILE:
                content, doc_type = await self._read_file(source)
            elif source_type == SourceType.WEB:
                content, doc_type = await self._read_web_page(source)
            elif source_type == SourceType.API:
                content, doc_type = await self._read_api(source)
            elif source_type == SourceType.DATABASE:
                content, doc_type = await self._read_database(source)
            elif source_type == SourceType.STREAM:
                content, doc_type = await self._read_stream(source)
            elif source_type == SourceType.NOTION:
                content, doc_type = await self._read_notion(source)
            elif source_type == SourceType.GOOGLE_DOCS:
                content, doc_type = await self._read_google_docs(source)
            else:
                raise ValueError(f"Unsupported source type: {source_type}")
            
            # Extract metadata
            metadata = await self._extract_metadata(
                content=content,
                source=source,
                source_type=source_type,
                doc_type=doc_type,
                **kwargs
            )
            
            # Process content
            processed_content = await self._process_content(
                content=content,
                doc_type=doc_type,
                **kwargs
            )
            
            # Create chunks
            chunks = await self._create_chunks(
                content=processed_content,
                metadata=metadata,
                **kwargs
            )
            
            return IngestedDocument(
                content=processed_content,
                metadata=metadata,
                chunks=chunks,
                raw_content=content,
                embeddings=None
            )
        
        finally:
            # Close session if it was created
            if self.session is not None:
                await self.session.close()
                self.session = None

    async def _read_file(
        self,
        file_path: str
    ) -> Tuple[Any, DocumentType]:
        """Read content from file."""
        # Determine file type
        if file_path.endswith(".pdf"):
            doc_type = DocumentType.PDF
            async with aiofiles.open(file_path, "rb") as f:
                content = await f.read()
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                text = "\n".join(page.extract_text() for page in pdf.pages)
            return text, doc_type
        
        elif file_path.endswith(".docx"):
            doc_type = DocumentType.DOCX
            doc = Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return text, doc_type
        
        elif file_path.endswith(".html"):
            doc_type = DocumentType.HTML
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            text = self.html_converter.handle(content)
            return text, doc_type
        
        elif file_path.endswith(".txt"):
            doc_type = DocumentType.TXT
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            return content, doc_type
        
        elif file_path.endswith(".json"):
            doc_type = DocumentType.JSON
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            data = json.loads(content)
            text = json.dumps(data, indent=2)
            return text, doc_type
        
        elif file_path.endswith(".csv"):
            doc_type = DocumentType.CSV
            df = pd.read_csv(file_path)
            text = df.to_string()
            return text, doc_type
        
        elif file_path.endswith(".md"):
            doc_type = DocumentType.MARKDOWN
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            return content, doc_type
        
        else:
            doc_type = DocumentType.UNKNOWN
            async with aiofiles.open(file_path, "r", encoding="utf-8") as f:
                content = await f.read()
            return content, doc_type

    async def _read_web_page(
        self,
        url: str
    ) -> Tuple[str, DocumentType]:
        """Read content from web page."""
        async with self.session.get(url) as response:
            content = await response.text()
        
        # Parse HTML
        soup = BeautifulSoup(content, "html.parser")
        
        # Remove unwanted elements
        for element in soup(["script", "style", "nav", "footer"]):
            element.decompose()
        
        # Extract text
        text = soup.get_text(separator="\n", strip=True)
        
        return text, DocumentType.HTML

    async def _read_api(
        self,
        url: str
    ) -> Tuple[str, DocumentType]:
        """Read content from API."""
        async with self.session.get(url) as response:
            content = await response.json()
        
        # Convert to text
        text = json.dumps(content, indent=2)
        
        return text, DocumentType.JSON

    async def _read_database(
        self,
        query: str
    ) -> Tuple[str, DocumentType]:
        """Read content from database."""
        # This is a placeholder implementation
        # Implement database connection and query execution
        return "", DocumentType.UNKNOWN

    async def _read_stream(
        self,
        topic: str
    ) -> Tuple[str, DocumentType]:
        """Read content from stream."""
        if not self.kafka_config:
            raise ValueError("Kafka configuration required for stream reading")
        
        # Initialize consumer
        consumer = KafkaConsumer(
            topic,
            **self.kafka_config
        )
        
        # Read messages
        messages = []
        for message in consumer:
            messages.append(message.value.decode())
        
        # Combine messages
        text = "\n".join(messages)
        
        return text, DocumentType.TXT

    async def _read_notion(
        self,
        page_id: str
    ) -> Tuple[str, DocumentType]:
        """Read content from Notion."""
        if not self.notion_client:
            raise ValueError("Notion client not initialized")
        
        # Get page content
        page = self.notion_client.pages.retrieve(page_id=page_id)
        blocks = self.notion_client.blocks.children.list(block_id=page_id)
        
        # Extract text from blocks
        text_blocks = []
        for block in blocks["results"]:
            if block["type"] == "paragraph":
                text_blocks.append(block["paragraph"]["rich_text"][0]["text"]["content"])
            elif block["type"] == "heading_1":
                text_blocks.append(f"# {block['heading_1']['rich_text'][0]['text']['content']}")
            elif block["type"] == "heading_2":
                text_blocks.append(f"## {block['heading_2']['rich_text'][0]['text']['content']}")
            elif block["type"] == "heading_3":
                text_blocks.append(f"### {block['heading_3']['rich_text'][0]['text']['content']}")
            elif block["type"] == "bulleted_list_item":
                text_blocks.append(f"* {block['bulleted_list_item']['rich_text'][0]['text']['content']}")
            elif block["type"] == "numbered_list_item":
                text_blocks.append(f"1. {block['numbered_list_item']['rich_text'][0]['text']['content']}")
        
        text = "\n".join(text_blocks)
        
        return text, DocumentType.MARKDOWN

    async def _read_google_docs(
        self,
        doc_id: str
    ) -> Tuple[str, DocumentType]:
        """Read content from Google Docs."""
        if not self.google_credentials:
            raise ValueError("Google credentials not provided")
        
        # Build service
        creds = Credentials.from_authorized_user_info(self.google_credentials)
        service = build("docs", "v1", credentials=creds)
        
        # Get document
        doc = service.documents().get(documentId=doc_id).execute()
        
        # Extract text
        text_blocks = []
        for element in doc["body"]["content"]:
            if "paragraph" in element:
                paragraph = element["paragraph"]
                text = ""
                for elem in paragraph["elements"]:
                    if "textRun" in elem:
                        text += elem["textRun"]["content"]
                text_blocks.append(text)
        
        text = "\n".join(text_blocks)
        
        return text, DocumentType.DOCX

    async def _extract_metadata(
        self,
        content: str,
        source: str,
        source_type: SourceType,
        doc_type: DocumentType,
        **kwargs
    ) -> DocumentMetadata:
        """Extract metadata from content."""
        # Get basic metadata
        metadata = DocumentMetadata(
            source=source,
            source_type=source_type.value,
            timestamp=datetime.now().timestamp(),
            content_type=doc_type.value,
            language=await self._detect_language(content),
            size=len(content.encode()),
            author=None,
            created_at=None,
            modified_at=None,
            custom_metadata={}
        )
        
        # Extract additional metadata based on document type
        if doc_type == DocumentType.PDF:
            # Extract PDF metadata
            with pdfplumber.open(io.BytesIO(content.encode())) as pdf:
                info = pdf.metadata
                metadata.author = info.get("Author")
                metadata.created_at = info.get("CreationDate")
                metadata.modified_at = info.get("ModDate")
        
        elif doc_type == DocumentType.DOCX:
            # Extract DOCX metadata
            doc = Document(io.BytesIO(content.encode()))
            core_props = doc.core_properties
            metadata.author = core_props.author
            metadata.created_at = core_props.created.timestamp() if core_props.created else None
            metadata.modified_at = core_props.modified.timestamp() if core_props.modified else None
        
        elif doc_type == DocumentType.HTML:
            # Extract HTML metadata
            soup = BeautifulSoup(content, "html.parser")
            metadata.author = soup.find("meta", {"name": "author"})
            metadata.author = metadata.author["content"] if metadata.author else None
        
        return metadata

    async def _process_content(
        self,
        content: str,
        doc_type: DocumentType,
        **kwargs
    ) -> str:
        """Process content based on document type."""
        if doc_type == DocumentType.HTML:
            # Clean HTML content
            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text(separator="\n", strip=True)
            return text
        
        elif doc_type == DocumentType.JSON:
            # Format JSON content
            try:
                data = json.loads(content)
                return json.dumps(data, indent=2)
            except json.JSONDecodeError:
                return content
        
        elif doc_type == DocumentType.CSV:
            # Format CSV content
            try:
                df = pd.read_csv(io.StringIO(content))
                return df.to_string()
            except pd.errors.EmptyDataError:
                return content
        
        return content

    async def _create_chunks(
        self,
        content: str,
        metadata: DocumentMetadata,
        **kwargs
    ) -> List[Dict[str, Any]]:
        """Create chunks from content."""
        # Use LLM to create semantic chunks
        prompt = f"""
        Split the following content into meaningful chunks.
        Consider:
        1. Semantic boundaries
        2. Context preservation
        3. Chunk size (max 1000 tokens)
        4. Topic continuity
        
        Content:
        {content}
        """
        
        response = await self.model.generate(prompt=prompt, **kwargs)
        
        # Parse chunks from response
        chunks = []
        current_chunk = {"content": "", "metadata": {}}
        
        for line in response.split("\n"):
            if line.strip():
                if len(current_chunk["content"]) + len(line) > 1000:
                    # Save current chunk
                    current_chunk["metadata"] = {
                        **metadata.__dict__,
                        "chunk_index": len(chunks)
                    }
                    chunks.append(current_chunk)
                    
                    # Start new chunk
                    current_chunk = {"content": line, "metadata": {}}
                else:
                    current_chunk["content"] += "\n" + line
        
        # Add last chunk
        if current_chunk["content"]:
            current_chunk["metadata"] = {
                **metadata.__dict__,
                "chunk_index": len(chunks)
            }
            chunks.append(current_chunk)
        
        return chunks

    async def _detect_language(
        self,
        text: str
    ) -> str:
        """Detect language of text."""
        # Use LLM to detect language
        prompt = f"""
        Detect the language of the following text.
        Return only the ISO 639-1 language code.
        
        Text:
        {text[:1000]}  # Use first 1000 chars for detection
        """
        
        response = await self.model.generate(prompt=prompt)
        return response.strip().lower() 