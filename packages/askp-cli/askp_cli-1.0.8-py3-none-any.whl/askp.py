#!/usr/bin/python3
import os
import sys
import shutil
import json
import logging
import time
import threading
import requests
import argparse
import shlex
import base64
import re
import tempfile
import textwrap
from dataclasses import dataclass
from typing import Optional, List, Dict, Any, Tuple, Set, cast
import platform
import subprocess
import getpass
#
# colorama (optional)
try:
    import colorama

    colorama_init = colorama.init
    from colorama import Fore as _ForeColors, Style as _StyleColors

    Fore: Any = cast(Any, _ForeColors)
    Style: Any = cast(Any, _StyleColors)
except Exception:  # colorama optional

    def colorama_init(*args: Any, **kwargs: Any) -> None:
        pass

    class _ForeFallback:
        CYAN = ""
        GREEN = ""
        RED = ""
        YELLOW = ""
        RESET = ""

    class _StyleFallback:
        BRIGHT = ""
        NORMAL = ""
        RESET_ALL = ""

    Fore: Any = _ForeFallback()
    Style: Any = _StyleFallback()

# Optional enhanced input with slash-commands completion
try:
    from prompt_toolkit import PromptSession as PT_PromptSession
    from prompt_toolkit.completion import (
        Completer as PTCompleter,
        Completion as PTCompletion,
    )

    try:
        from prompt_toolkit.formatted_text import HTML as PT_HTML
    except Exception:
        PT_HTML = None  # fallback to plain text
    try:
        from prompt_toolkit.formatted_text import ANSI as PT_ANSI
    except Exception:
        PT_ANSI = None
    try:
        from prompt_toolkit.application.current import get_app as PT_get_app
    except Exception:
        PT_get_app = None
    try:
        from prompt_toolkit.history import FileHistory as PT_FileHistory
    except Exception:
        PT_FileHistory = None
    PROMPT_TOOLKIT = True
except Exception:
    # Provide permissive aliases so type checker doesn't complain
    PT_PromptSession: Any = None
    PTCompleter: Any = object
    PTCompletion: Any = None
    PT_HTML: Any = None
    PT_get_app: Any = None
    PT_FileHistory: Any = None
    PROMPT_TOOLKIT = False


# Available Perplexity models
AVAILABLE_MODELS = [
    "sonar",
    "sonar-pro",
    "sonar-reasoning",
    "sonar-reasoning-pro",
    "sonar-deep-research",
]

DEFAULT_MODEL = "sonar"  # default for sync; async path will force sonar-deep-research
API_URL = "https://api.perplexity.ai/chat/completions"
ASYNC_API_URL = "https://api.perplexity.ai/async/chat/completions"


@dataclass
class ApiConfig:
    api_url: str = API_URL
    api_key: Optional[str] = None


class AttachmentState:
    """Manages pending document attachments for pre-prompt inclusion."""
    
    def __init__(self):
        self.pending_attachments: List[Dict[str, Any]] = []
        self.auto_clear: bool = True
        self.max_total_size: int = 50000  # Character limit for all attachments
    
    def add_attachment(self, filename: str, content: str):
        """Add a document attachment to pending list."""
        self.pending_attachments.append({
            'filename': filename,
            'content': content,
            'size': len(content)
        })
    
    def get_total_size(self) -> int:
        """Get total character count of all pending attachments."""
        return sum(att['size'] for att in self.pending_attachments)
    
    def build_preprompt(self, user_query: str) -> str:
        """Build user message with attached documents as pre-prompt."""
        if not self.pending_attachments:
            return user_query
        
        # Check if total size exceeds limit
        total_size = self.get_total_size()
        if total_size > self.max_total_size:
            return self._build_truncated_preprompt(user_query)
        
        return self._build_full_preprompt(user_query)
    
    def _build_full_preprompt(self, user_query: str) -> str:
        """Build preprompt with full document content."""
        if len(self.pending_attachments) == 1:
            # Single document - simpler format
            att = self.pending_attachments[0]
            return f"""[ATTACHED DOCUMENT: {att['filename']}]
{att['content']}
[END ATTACHED DOCUMENT]

{user_query}"""
        
        # Multiple documents - structured format
        content_parts = []
        for att in self.pending_attachments:
            content_parts.append(f"""
--- DOCUMENT: {att['filename']} ---
{att['content']}
--- END DOCUMENT ---""")
        
        return f"""[ATTACHED DOCUMENTS]
{''.join(content_parts)}
[END ATTACHED DOCUMENTS]

{user_query}"""
    
    def _build_truncated_preprompt(self, user_query: str) -> str:
        """Build preprompt with truncated content when size limit exceeded."""
        available_size = self.max_total_size - len(user_query) - 500  # Reserve space for formatting
        
        content_parts = []
        used_size = 0
        
        for att in self.pending_attachments:
            header = f"\n--- DOCUMENT: {att['filename']} ---\n"
            footer = "\n--- END DOCUMENT ---"
            
            available_for_content = available_size - used_size - len(header) - len(footer)
            
            if available_for_content <= 100:
                content_parts.append(f"{header}[Content truncated - document too large]{footer}")
                break
            
            if len(att['content']) <= available_for_content:
                content_parts.append(f"{header}{att['content']}{footer}")
                used_size += len(header) + len(att['content']) + len(footer)
            else:
                truncated = att['content'][:available_for_content-20] + "\n... [truncated]"
                content_parts.append(f"{header}{truncated}{footer}")
                break
        
        return f"""[ATTACHED DOCUMENTS - Some content may be truncated due to size limits]
{''.join(content_parts)}
[END ATTACHED DOCUMENTS]

{user_query}"""
    
    def clear(self):
        """Clear all pending attachments."""
        self.pending_attachments.clear()
    
    def has_attachments(self) -> bool:
        """Check if there are pending attachments."""
        return len(self.pending_attachments) > 0
    
    def get_summary(self) -> str:
        """Get summary of pending attachments for display."""
        if not self.pending_attachments:
            return "No pending attachments"
        
        total_size = self.get_total_size()
        files = [att['filename'] for att in self.pending_attachments]
        
        return f"{len(files)} file(s): {', '.join(files)} ({total_size:,} chars)"


def get_api_key() -> Optional[str]:
    """Return API key from environment with a sensible fallback."""
    return os.getenv("PPLX_API_KEY") or os.getenv("PERPLEXITY_API_KEY")


def _write_or_update_line(path: str, match_prefix: str, new_line: str) -> bool:
    try:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        content = ""
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
        lines = content.splitlines()
        updated = False
        out_lines = []
        for ln in lines:
            if ln.strip().startswith(match_prefix):
                if not updated:
                    out_lines.append(new_line)
                    updated = True
                # skip existing matching line(s)
            else:
                out_lines.append(ln)
        if not updated:
            if out_lines and out_lines[-1].strip() != "":
                out_lines.append("")
            out_lines.append(new_line)
        with open(path, "w", encoding="utf-8") as f:
            f.write("\n".join(out_lines) + "\n")
        return True
    except Exception:
        return False


def persist_api_key_unix(key: str) -> Optional[str]:
    shell = os.getenv("SHELL", "")
    home = os.path.expanduser("~")
    is_darwin = platform.system().lower() == "darwin"
    # Choose profile files in priority order
    candidates: List[Tuple[str, str, str]] = []  # (path, match_prefix, line)
    # fish shell
    if "fish" in shell:
        cfg = os.path.join(home, ".config", "fish", "config.fish")
        candidates.append((cfg, "set -x PPLX_API_KEY", f"set -x PPLX_API_KEY \"{key}\""))
    # zsh
    if "zsh" in shell or (not shell and is_darwin):
        rc = os.path.join(home, ".zshrc")
        candidates.append((rc, "export PPLX_API_KEY=", f"export PPLX_API_KEY=\"{key}\""))
    # bash
    if "bash" in shell or not candidates:
        bashrc = os.path.join(home, ".bashrc")
        bash_profile = os.path.join(home, ".bash_profile")
        # macOS default interactive login shell uses .bash_profile
        if is_darwin:
            candidates.append((bash_profile, "export PPLX_API_KEY=", f"export PPLX_API_KEY=\"{key}\""))
            candidates.append((bashrc, "export PPLX_API_KEY=", f"export PPLX_API_KEY=\"{key}\""))
        else:
            candidates.append((bashrc, "export PPLX_API_KEY=", f"export PPLX_API_KEY=\"{key}\""))
            candidates.append((bash_profile, "export PPLX_API_KEY=", f"export PPLX_API_KEY=\"{key}\""))
    # Fallback
    profile = os.path.join(home, ".profile")
    candidates.append((profile, "export PPLX_API_KEY=", f"export PPLX_API_KEY=\"{key}\""))

    # Try to write to the first viable candidate
    for path, match_prefix, line in candidates:
        ok = _write_or_update_line(path, match_prefix, line)
        if ok:
            return path
    return None


def persist_api_key_windows(key: str) -> bool:
    try:
        # Use setx to set for the current user (persists across sessions)
        completed = subprocess.run(["setx", "PPLX_API_KEY", key], capture_output=True, text=True, shell=False)
        return completed.returncode == 0
    except Exception:
        return False


# --- Simple settings persistence ---
ASKP_DATA_DIR = os.path.join(os.path.expanduser("~"), ".askp")
LEGACY_CONFIG_PATH = os.path.join(os.path.expanduser("~"), ".askp_config.json")
CONFIG_PATH = os.path.join(ASKP_DATA_DIR, "askp_config.json")

PREF_KEYS = [
    "model",
    "academic",
    "domain",
    "recency",
    "citations",
    "usage",
    "stream",
    "attach_truncation_limit",
    "session_enabled",
    "session_name",
    "async_meta",
    "md_render",
    "theme",
    "echo_user",
    "text",  # default one-shot output mode (True=text, False=json)
    "persistent_system",  # persistent system message
    "disable_table_rendering",  # bypass table rendering entirely
]


def load_prefs() -> Optional[Dict[str, Any]]:
    try:
        # Migrate legacy config if needed
        if os.path.exists(LEGACY_CONFIG_PATH) and not os.path.exists(CONFIG_PATH):
            try:
                os.makedirs(ASKP_DATA_DIR, exist_ok=True)
                shutil.move(LEGACY_CONFIG_PATH, CONFIG_PATH)
            except Exception:
                pass
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, dict):
                return data
    except Exception:
        pass
    return None

# --- Session utilities ---
PROJECT_SESSION_DIRNAME = ASKP_DATA_DIR
SESSIONS_SUBDIR = "sessions"

def _project_session_dir() -> str:
    # Sessions are centralized per-user: stored under ~/.askp/sessions
    base = os.path.join(PROJECT_SESSION_DIRNAME, SESSIONS_SUBDIR)
    try:
        os.makedirs(base, exist_ok=True)
    except Exception:
        pass
    return base


def _sanitize_session_name(name: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in name.strip())
    return safe or "session"


def _session_file_for(name: str) -> str:
    return os.path.join(_project_session_dir(), f"{_sanitize_session_name(name)}.jsonl")


def _list_sessions() -> List[Tuple[str, float, int, Optional[float]]]:
    # returns list of (name, mtime, line_count, duration_minutes)
    out: List[Tuple[str, float, int, Optional[float]]] = []
    d = _project_session_dir()
    try:
        for fn in os.listdir(d):
            if not fn.endswith(".jsonl"):
                continue
            path = os.path.join(d, fn)
            try:
                first_time = None
                last_time = None
                n = 0
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    for line in f:
                        n += 1
                        try:
                            obj = json.loads(line.strip())
                            if "timestamp" in obj:
                                ts = obj["timestamp"]
                                if first_time is None:
                                    first_time = ts
                                last_time = ts
                        except Exception:
                            continue
                duration = None
                if first_time is not None and last_time is not None and last_time > first_time:
                    duration = (last_time - first_time) / 60.0  # minutes
            except Exception:
                n = 0
                duration = None
            out.append((fn[:-6], os.path.getmtime(path), n, duration))
    except Exception:
        pass
    out.sort(key=lambda x: x[1], reverse=True)
    return out


def _load_session_messages(name: str) -> List[Dict[str, Any]]:
    path = _session_file_for(name)
    msgs: List[Dict[str, Any]] = []
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    obj = json.loads(line)
                    if isinstance(obj, dict) and "role" in obj and "content" in obj:
                        msgs.append(obj)
                except Exception:
                    continue
    except FileNotFoundError:
        pass
    except Exception:
        pass
    return msgs


def _default_session_name() -> str:
    import datetime as _dt
    ts = _dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    return f"session-{ts}"


def _short_time_ago(from_ts: float) -> str:
    try:
        delta = max(0, time.time() - from_ts)
        if delta < 60:
            return f"{int(delta)}s ago"
        minutes = int(delta // 60)
        if minutes < 60:
            return f"{minutes}m ago"
        hours = int(delta // 3600)
        if hours < 24:
            return f"{hours}h ago"
        days = int(delta // 86400)
        if days < 30:
            return f"{days}d ago"
        months = int(days // 30)
        if months < 12:
            return f"{months}mo ago"
        years = int(months // 12)
        return f"{years}y ago"
    except Exception:
        return "just now"


def save_prefs_from_args(args: argparse.Namespace) -> bool:
    try:
        data: Dict[str, Any] = {}
        for k in PREF_KEYS:
            if hasattr(args, k):
                data[k] = getattr(args, k)
        os.makedirs(ASKP_DATA_DIR, exist_ok=True)
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return True
    except Exception:
        return False


def apply_prefs_to_args(prefs: Dict[str, Any], args: argparse.Namespace) -> None:
    for k in PREF_KEYS:
        if k in prefs:
            setattr(args, k, prefs[k])
            if k == "json_schema":
                # argparse uses dest json_schema for --json-schema, ensure it’s set
                setattr(args, "json_schema", prefs[k])
            if k == "attach_truncation_limit":
                try:
                    val = int(prefs[k])
                    setattr(args, "attach_truncation_limit", val)
                except Exception:
                    pass
            if k == "stream":
                # mark that stream preference came from persisted prefs
                setattr(args, "_stream_pref_loaded", True)


def extract_content(data: Dict[str, Any]) -> Optional[str]:
    try:
        # Handle both sync and async shapes. Async responses often wrap content under data["response"].
        root = data
        if isinstance(data, dict) and isinstance(data.get("response"), dict):
            root = data["response"]
        content = root["choices"][0]["message"]["content"]
        if isinstance(content, list):
            # If content is a list of content objects, extract text parts
            text_parts = []
            for item in content:
                if isinstance(item, dict) and item.get("type") == "text":
                    text_parts.append(item.get("text", ""))
            return "".join(text_parts)
        else:
            # If content is a string (for older models or simple text responses)
            return content
    except Exception:
        return None


def extract_delta_token(obj: Dict[str, Any]) -> str:
    try:
        ch = obj.get("choices", [{}])[0]
        if "delta" in ch and isinstance(ch["delta"], dict):
            return ch["delta"].get("content", "") or ""
        # some providers may stream full message chunks
        if "message" in ch and isinstance(ch["message"], dict):
            return ch["message"].get("content", "") or ""
    except Exception:
        pass
    return ""


# --- Formatting Infrastructure ---

class FormattingTheme:
    """Color definitions and style constants for consistent UI formatting."""
    def __init__(self, theme_cfg: Optional[Dict[str, str]] = None):
        self._apply_theme(theme_cfg or {})
    
    def _apply_theme(self, cfg: Dict[str, str]) -> None:
        # Map names -> Fore colors
        name_to_color = {
            "black": getattr(Fore, "BLACK", ""),
            "red": Fore.RED,
            "green": Fore.GREEN,
            "yellow": Fore.YELLOW,
            "blue": Fore.BLUE if hasattr(Fore, "BLUE") else "",
            "magenta": Fore.MAGENTA if hasattr(Fore, "MAGENTA") else "",
            "cyan": Fore.CYAN,
            "white": Fore.WHITE if hasattr(Fore, "WHITE") else "",
            "bright_black": getattr(Fore, "LIGHTBLACK_EX", ""),
            "bright_red": Fore.LIGHTRED_EX if hasattr(Fore, "LIGHTRED_EX") else Fore.RED,
            "bright_green": Fore.LIGHTGREEN_EX if hasattr(Fore, "LIGHTGREEN_EX") else Fore.GREEN,
            "bright_yellow": Fore.LIGHTYELLOW_EX if hasattr(Fore, "LIGHTYELLOW_EX") else Fore.YELLOW,
            "bright_blue": Fore.LIGHTBLUE_EX if hasattr(Fore, "LIGHTBLUE_EX") else (Fore.BLUE if hasattr(Fore, "BLUE") else ""),
            "bright_magenta": Fore.LIGHTMAGENTA_EX if hasattr(Fore, "LIGHTMAGENTA_EX") else (Fore.MAGENTA if hasattr(Fore, "MAGENTA") else ""),
            "bright_cyan": Fore.LIGHTCYAN_EX if hasattr(Fore, "LIGHTCYAN_EX") else Fore.CYAN,
            "bright_white": Fore.LIGHTWHITE_EX if hasattr(Fore, "LIGHTWHITE_EX") else (Fore.WHITE if hasattr(Fore, "WHITE") else ""),
        }
        def col(name: str, fallback: str) -> str:
            try:
                key = cfg.get(name)
            except Exception:
                key = None
            if key and key in name_to_color:
                return name_to_color[key]
            return fallback
        # Color definitions with fallbacks
        self.USER_COLOR = col("user", Fore.CYAN) + Style.BRIGHT
        self.AI_COLOR = col("ai", Fore.GREEN)
        self.CITATION_COLOR = col("citations", Fore.YELLOW)
        self.METADATA_COLOR = col("metadata", Fore.BLUE if hasattr(Fore, "BLUE") else Fore.CYAN)
        self.ERROR_COLOR = col("error", Fore.RED)
        self.SUCCESS_COLOR = col("success", Fore.GREEN)
        self.WARNING_COLOR = col("warning", Fore.YELLOW)
        self.PROMPT_COLOR = col("user", Fore.CYAN)
        self.SECTION_COLOR = col("section", (Fore.MAGENTA if hasattr(Fore, "MAGENTA") else Fore.CYAN)) + Style.BRIGHT
        # Symbols and decorators
        self.USER_PREFIX = "You"
        self.AI_PREFIX = "AI"
        self.BULLET_POINT = "•"
        self.INDENT = "  "
        self.RESET = Style.RESET_ALL
    
    @property
    def SECTION_SEPARATOR(self) -> str:
        try:
            width = shutil.get_terminal_size().columns
            return "─" * width
        except Exception:
            return "─" * 80  # Fallback


class OutputFormatter:
    """Centralized formatting system for consistent UI presentation."""
    
    def __init__(self, use_colors: bool = True, theme_cfg: Optional[Dict[str, str]] = None, md_render: bool = True):
        self.use_colors = use_colors
        self.theme = FormattingTheme(theme_cfg or {})
        self.md_render = md_render
        self._disable_tables = False  # Allow table rendering by default
    
    def _get_terminal_width(self) -> int:
        """Get terminal width with fallback and padding."""
        try:
            terminal_width = shutil.get_terminal_size().columns
        except Exception:
            terminal_width = 80  # Fallback width
        
        # Apply terminal padding like Gemini CLI (TERMINAL_PADDING_X = 8)
        TERMINAL_PADDING_X = 8
        return max(20, terminal_width - TERMINAL_PADDING_X)  # Ensure minimum width
    
    def _wrap_text(self, text: str, width: Optional[int] = None, indent: int = 0) -> str:
        """Wrap text to specified width with proper word boundaries."""
        if width is None:
            width = self._get_terminal_width()
        
        if not text.strip():
            return text
        
        # Handle indentation
        indent_str = " " * indent
        effective_width = max(20, width - indent)
        
        lines = []
        for paragraph in text.split('\n'):
            if not paragraph.strip():
                lines.append('')
            else:
                wrapped = textwrap.fill(
                    paragraph,
                    width=effective_width,
                    initial_indent=indent_str,
                    subsequent_indent=indent_str,
                    break_long_words=False,
                    break_on_hyphens=False
                )
                lines.append(wrapped)
        
        return '\n'.join(lines)
    
    def _cleanup_broken_tables(self, text: str) -> str:
        """Aggressively detect and reformat broken tables into readable list format.
        
        This catches tables that are:
        - Poorly formatted by the AI
        - Broken during terminal rendering
        - Have misaligned columns
        - Have truncated headers
        """
        if not text or len(text) > 50000:  # Skip very large content
            return text
        
        lines = text.splitlines()
        result_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Detect potential table start - look for patterns that indicate tabular data
            if self._looks_like_broken_table_start(line, lines[i:i+5]):
                # Try to extract and reformat the table
                table_lines, consumed = self._extract_broken_table(lines[i:])
                if table_lines and consumed > 1:
                    # Successfully detected a broken table
                    reformatted = self._reformat_broken_table(table_lines)
                    result_lines.extend(reformatted)
                    i += consumed
                    continue
            
            result_lines.append(lines[i])
            i += 1
        
        return '\n'.join(result_lines)
    
    def _looks_like_broken_table_start(self, line: str, next_lines: List[str]) -> bool:
        """Detect if this looks like the start of a broken table."""
        if not line:
            return False
        
        # Look for common broken table patterns
        indicators = [
            # Headers with separators
            ('|' in line and any('─' in nl or '---' in nl for nl in next_lines[:3])),
            # Headers followed by data with rank/numbers
            (any(keyword in line.lower() for keyword in ['rank', 'entity', 'country', 'name', 'item']) 
             and any(nl.strip() and (nl.strip()[0].isdigit() or '|' in nl) for nl in next_lines[:3])),
            # Lines that look like broken table headers
            (line.count('|') > 1 and len(line) > 30),
            # Headers with cut-off text (common in broken tables)
            ('...' in line and '|' in line),
        ]
        
        return any(indicators)
    
    def _extract_broken_table(self, lines: List[str]) -> Tuple[List[str], int]:
        """Extract lines that appear to be part of a broken table."""
        table_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue
            
            # Stop conditions - end of table
            if (i > 0 and 
                not any(c in line for c in ['|', '─', '%', 'million', 'billion', 'tons']) and
                not line[0].isdigit() and
                not any(keyword in line.lower() for keyword in ['total', 'others', 'world'])):
                break
            
            # Include lines that look like table content
            if (line and (
                '|' in line or 
                '─' in line or 
                line[0].isdigit() or
                any(keyword in line.lower() for keyword in ['rank', 'entity', 'emissions', 'country', 'total', 'others'])
            )):
                table_lines.append(line)
            
            i += 1
            if i > 20:  # Limit to prevent runaway detection
                break
        
        return table_lines, i
    
    def _reformat_broken_table(self, table_lines: List[str]) -> List[str]:
        """Reformat broken table lines into a clean, readable list format."""
        if not table_lines:
            return []
        
        # Try to extract structured data from the broken table
        data_rows = []
        headers = []
        
        # First, try to identify headers and data
        for line in table_lines:
            if '─' in line:  # Skip separator lines
                continue
            
            # Clean the line
            clean_line = re.sub(r'\s+', ' ', line.strip())
            clean_line = re.sub(r'\|+', '|', clean_line)  # Collapse multiple pipes
            
            # Skip if it's mostly separators or empty
            if not clean_line or len(clean_line.replace('|', '').replace(' ', '')) < 3:
                continue
            
            # Try to extract meaningful data
            if '|' in clean_line:
                # Parse pipe-separated data
                parts = [p.strip() for p in clean_line.split('|') if p.strip()]
                if parts:
                    if not headers and any(keyword in clean_line.lower() for keyword in ['rank', 'entity', 'country', 'emissions']):
                        headers = parts
                    else:
                        data_rows.append(parts)
            else:
                # Try to parse space-separated data (common in broken tables)
                # Look for patterns like: "1 China 14,093 27%"
                match = re.match(r'^(\d+)\s+([^0-9]+?)\s+([\d,]+(?:\.\d+)?)\s*.*?(\d+(?:\.\d+)?%)', clean_line)
                if match:
                    rank, entity, value, percentage = match.groups()
                    data_rows.append([rank.strip(), entity.strip(), value.strip(), percentage.strip()])
                else:
                    # Try other patterns
                    parts = clean_line.split()
                    if len(parts) >= 3 and parts[0].isdigit():
                        # Likely a data row
                        data_rows.append(parts)
        
        # Generate clean output
        result = []
        
        if data_rows:
            use_colors = self.use_colors
            bullet = self.theme.BULLET_POINT if self.theme.BULLET_POINT else "•"
            bullet_colored = f"{self.theme.CITATION_COLOR}{bullet}{self.theme.RESET}" if use_colors else bullet
            
            # Add a header
            if headers:
                header_text = "Table Data"
            else:
                header_text = "Extracted Table Data"
            
            if use_colors:
                header_colored = f"{self.theme.SECTION_COLOR}{header_text}{self.theme.RESET}"
            else:
                header_colored = header_text
            
            result.append(f"\n{header_colored}")
            result.append("─" * min(len(header_text), 50))
            
            # Format each data row
            for i, row in enumerate(data_rows[:10]):  # Limit to 10 rows to prevent spam
                if len(row) >= 2:
                    # Format as: • Primary: Details
                    primary = row[0] if row[0] else f"Item {i+1}"
                    if len(row) > 1:
                        secondary = row[1]
                        
                        # Try to identify what this data represents
                        details = []
                        for j, item in enumerate(row[2:], 2):
                            if item and item.strip():
                                if '%' in item:
                                    details.append(f"Percentage: {item}")
                                elif any(char.isdigit() for char in item):
                                    if ',' in item or 'million' in item.lower() or 'billion' in item.lower():
                                        details.append(f"Value: {item}")
                                    else:
                                        details.append(f"Data: {item}")
                                else:
                                    details.append(item)
                        
                        if use_colors:
                            primary_colored = f"{self.theme.SECTION_COLOR}{primary}{self.theme.RESET}"
                            secondary_colored = f"{secondary}"
                        else:
                            primary_colored = primary
                            secondary_colored = secondary
                        
                        # Build the line
                        line_parts = [f"{bullet_colored} {primary_colored}: {secondary_colored}"]
                        
                        # Add details with indentation
                        for detail in details:
                            if use_colors:
                                detail_colored = f"  {self.theme.METADATA_COLOR}├─{self.theme.RESET} {detail}"
                            else:
                                detail_colored = f"  ├─ {detail}"
                            line_parts.append(detail_colored)
                        
                        result.extend(line_parts)
            
            result.append("")  # Add spacing after table
        
        return result
    
    def _apply_color(self, text: str, color: str) -> str:
        """Apply color if colors are enabled, otherwise return plain text."""
        if not self.use_colors:
            return text
        return f"{color}{text}{self.theme.RESET}"
    
    def format_user_input(self, text: str) -> str:
        """Format user input with consistent styling."""
        prefix = self._apply_color(f"{self.theme.USER_PREFIX}>", self.theme.USER_COLOR)
        return f"{prefix} {text}"
    
    def format_ai_response(self, text: str) -> str:
        """Format AI response with proper text wrapping."""
        prefix = self._apply_color(f"{self.theme.AI_PREFIX}>", self.theme.AI_COLOR)
        rendered = self._render_basic_markdown(text) if getattr(self, 'md_render', True) else text
        
        # Wrap text properly to prevent word cutting
        wrapped_text = self._wrap_text(rendered, indent=len(f"{self.theme.AI_PREFIX}> "))
        
        # Apply prefix only to the first line, maintain alignment for subsequent lines
        lines = wrapped_text.split('\n')
        if lines:
            lines[0] = f"{prefix} {lines[0].lstrip()}"
        
        return '\n'.join(lines)
    
    def _render_basic_markdown(self, text: str) -> str:
        """Enhanced Markdown rendering for terminals with improved table support.
        
        Supports:
        - Headings: #, ##, ### at line start
        - Bold: **text**
        - Italics: *text* or _text_
        - Inline code: `code`
        - Bullets: -, *, + at line start -> •
        - Links: [text](url) -> text (url)
        - Tables: | col1 | col2 | with proper alignment and borders
        
        Notes:
        - Skips rendering for very large strings to avoid performance cost
        - Enhanced table rendering with proper column width calculation
        - Can be disabled entirely with _disable_tables flag
        """
        try:
            if not text:
                return text
            # Avoid heavy processing on very large content
            if len(text) > 200000:
                return text

            # Check if table rendering is disabled
            if hasattr(self, '_disable_tables') and self._disable_tables:
                # Still do basic markdown but skip table processing
                return self._render_basic_markdown_no_tables(text)

            # AGGRESSIVE TABLE CLEANUP FIRST - catch broken tables before markdown processing
            text = self._cleanup_broken_tables(text)

            # Compile patterns once per instance
            if not hasattr(self, "_md_compiled"):
                self._md_compiled = True
                self._re_bold = re.compile(r"\*\*(.+?)\*\*")
                # Match italics but not bold (avoid ** by negative lookarounds)
                self._re_ital_star = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
                self._re_ital_us = re.compile(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)")  # No underscore stripping
                self._re_code = re.compile(r"`([^`]+)`")
                self._re_link = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
                self._re_heading = re.compile(r"^(\s{0,3})(#{1,3})\s+(.+)$")
                self._re_bullet = re.compile(r"^(\s*)([-*+])\s+")
                # More robust table detection patterns
                self._re_table_row = re.compile(r"^\s*\|(.+?)\|\s*$")
                self._re_table_sep = re.compile(r"^\s*\|[\s\-:]+\|\s*$")
                # Also detect tables without outer pipes
                self._re_table_row_alt = re.compile(r"^[^|]*\|[^|]*(\|[^|]*)*$")

            # Process line-by-line for headings/bullets/tables; apply inline after
            lines = text.splitlines()
            out_lines = []
            use_colors = self.use_colors
            bullet_sym = self.theme.BULLET_POINT if self.theme.BULLET_POINT else "•"
            
            # Table processing state
            in_table = False
            table_rows = []
            table_widths = []
            
            def _parse_table_cells(raw_content: str) -> List[str]:
                """Parse table cells from raw content, handling various edge cases."""
                if not raw_content.strip():
                    return []
                
                cells = []
                # Split by | and handle empty cells properly
                raw_cells = raw_content.split('|')
                
                for cell in raw_cells:
                    # Clean and preserve all cells, even empty ones
                    cleaned = cell.strip()
                    cells.append(cleaned)
                
                # Only remove truly empty cells from start/end if they're artifacts
                # But be more conservative - only remove if clearly padding
                while len(cells) > 1 and cells[0] == '':
                    cells.pop(0)
                while len(cells) > 1 and cells[-1] == '':
                    cells.pop()
                
                # Filter out cells that are just whitespace or separators
                meaningful_cells = []
                for cell in cells:
                    # Skip cells that are just dashes (separator artifacts)
                    if cell and not re.match(r'^[\s\-:]+$', cell):
                        meaningful_cells.append(cell)
                    elif cell == '':  # Keep empty cells for structure
                        meaningful_cells.append(cell)
                
                return meaningful_cells if meaningful_cells else []
            
            def _flush_table():
                """Flush accumulated table rows with smart formatting based on available space."""
                nonlocal table_rows, table_widths, in_table
                if not table_rows:
                    return
                
                # Always use list format to avoid rendering issues
                # Filter out separator rows and validate table structure
                data_rows = []
                for row in table_rows:
                    if not (len(row) == 1 and '---' in row[0]):
                        # Only include rows with meaningful content
                        if row and any(cell.strip() for cell in row):
                            data_rows.append(row)
                
                if not data_rows:
                    table_rows = []
                    table_widths = []
                    in_table = False
                    return
                
                # Always convert to list format for better readability
                header_row = data_rows[0] if data_rows else []
                data_row_list = data_rows[1:] if len(data_rows) > 1 else []
                
                bullet = self.theme.BULLET_POINT if self.theme.BULLET_POINT else "•"
                bullet_colored = f"{self.theme.CITATION_COLOR}{bullet}{self.theme.RESET}" if use_colors else bullet
                
                for row_idx, row in enumerate(data_row_list):
                    # Create a readable entry for each data row
                    if len(header_row) > 0 and len(row) > 0:
                        # Primary field (first column) with proper formatting
                        primary_field = row[0].strip()
                        if primary_field:
                            if use_colors:
                                primary_field = f"{self.theme.SECTION_COLOR}{primary_field}{self.theme.RESET}"
                            
                            # Build the list item with primary field
                            list_parts = [f"{bullet_colored} **{primary_field}**"]
                            
                            # Add other fields as supplementary information
                            additional_info = []
                            for i in range(1, min(len(header_row), len(row))):
                                value = row[i].strip() if i < len(row) else ""
                                if value:
                                    additional_info.append(value)
                            
                            if additional_info:
                                list_parts[0] += ": " + " | ".join(additional_info)
                            
                            # Output the formatted list item
                            out_lines.append(list_parts[0])
                
                # Reset table state
                table_rows = []
                table_widths = []
                in_table = False
            
            for ln in lines:
                # Check for table rows first - try multiple patterns
                table_match = self._re_table_row.match(ln)
                table_match_alt = None  # disable alt table detection to avoid false positives
                sep_match = self._re_table_sep.match(ln)
                
                # Determine if this is a table row
                is_table_row = table_match or table_match_alt or sep_match
                
                if is_table_row:
                    if not in_table:
                        in_table = True
                    
                    if table_match:
                        # Standard table format: | cell | cell |
                        raw_content = table_match.group(1)
                        cells = _parse_table_cells(raw_content)
                        if cells:
                            table_rows.append(cells)
                    elif table_match_alt:
                        # Alternative format: cell | cell | cell (no outer pipes)
                        cells = _parse_table_cells(ln.strip())
                        if cells:
                            table_rows.append(cells)
                    elif sep_match:
                        # Separator row (---)
                        table_rows.append(['---'])
                    continue
                else:
                    # Not a table row - flush any pending table
                    if in_table:
                        _flush_table()
                
                # Headings (#, ##, ###)
                mh = self._re_heading.match(ln)
                if mh:
                    indent, hashes, title = mh.groups()
                    level = len(hashes)
                    if use_colors:
                        # Use bright and section color for headings
                        title_fmt = f"{Style.BRIGHT}{self.theme.SECTION_COLOR}{title}{self.theme.RESET}"
                    else:
                        title_fmt = title
                    # Simple visual underline for H1/H2 when colors available
                    if level == 1 and use_colors:
                        out_lines.append(f"{indent}{title_fmt}")
                        out_lines.append(f"{indent}{'=' * max(3, min(len(title), 80))}")
                    elif level == 2 and use_colors:
                        out_lines.append(f"{indent}{title_fmt}")
                        out_lines.append(f"{indent}{'-' * max(3, min(len(title), 80))}")
                    else:
                        out_lines.append(f"{indent}{title_fmt}")
                    continue
                # Bullets: -, *, +  -> • with slight color
                mb = self._re_bullet.match(ln)
                if mb:
                    sp, _ch = mb.groups()
                    bullet = bullet_sym
                    if use_colors:
                        bullet = f"{self.theme.CITATION_COLOR}{bullet_sym}{self.theme.RESET}"
                    ln = self._re_bullet.sub(f"{sp}{bullet} ", ln, count=1)
                out_lines.append(ln)
            
            # Flush any remaining table at end of text
            if in_table:
                _flush_table()

            s = "\n".join(out_lines)

            # Inline elements
            def _bold_sub(m):
                inner = m.group(1)
                if use_colors:
                    return f"{Style.BRIGHT}{inner}{self.theme.RESET}"
                return inner
            s = self._re_bold.sub(_bold_sub, s)

            def _ital_sub(m):
                inner = m.group(1)
                if use_colors:
                    return f"{Style.BRIGHT}{inner}{self.theme.RESET}"
                return inner
            s = self._re_ital_star.sub(_ital_sub, s)
            # s = self._re_ital_us.sub(_ital_sub, s)  # Disable underscore italics to preserve underscores

            def _code_sub(m):
                inner = m.group(1)
                if use_colors:
                    return f"{self.theme.METADATA_COLOR}{inner}{self.theme.RESET}"
                return inner
            s = self._re_code.sub(_code_sub, s)

            def _link_sub(m):
                txt, url = m.group(1), m.group(2)
                if use_colors:
                    return f"{Style.BRIGHT}{txt}{self.theme.RESET} ({self.theme.METADATA_COLOR}{url}{self.theme.RESET})"
                return f"{txt} ({url})"
            s = self._re_link.sub(_link_sub, s)

            return s
        except Exception:
            return text
    
    def _render_basic_markdown_no_tables(self, text: str) -> str:
        """Simplified markdown rendering that skips table processing entirely."""
        try:
            # Just do basic inline formatting without table processing
            if not hasattr(self, "_md_compiled_simple"):
                self._md_compiled_simple = True
                self._re_bold_simple = re.compile(r"\*\*(.+?)\*\*")
                self._re_ital_star_simple = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
                self._re_ital_us_simple = re.compile(r"(?<!_)_(?!_)(.+?)(?<!_)_(?!_)")
                self._re_code_simple = re.compile(r"`([^`]+)`")
                self._re_link_simple = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
                self._re_heading_simple = re.compile(r"^(\s{0,3})(#{1,3})\s+(.+)$")
                self._re_bullet_simple = re.compile(r"^(\s*)([-*+])\s+")

            lines = text.splitlines()
            out_lines = []
            use_colors = self.use_colors
            bullet_sym = self.theme.BULLET_POINT if self.theme.BULLET_POINT else "•"
            
            for ln in lines:
                # Skip any table-like lines entirely
                if '|' in ln and ln.strip().startswith('|') and ln.strip().endswith('|'):
                    # Just output as plain text, no special formatting
                    out_lines.append(ln)
                    continue
                
                # Headings (#, ##, ###)
                mh = self._re_heading_simple.match(ln)
                if mh:
                    indent, hashes, title = mh.groups()
                    if use_colors:
                        title_fmt = f"{Style.BRIGHT}{self.theme.SECTION_COLOR}{title}{self.theme.RESET}"
                    else:
                        title_fmt = title
                    out_lines.append(f"{indent}{title_fmt}")
                    continue
                
                # Bullets: -, *, +  -> •
                mb = self._re_bullet_simple.match(ln)
                if mb:
                    sp, _ch = mb.groups()
                    bullet = bullet_sym
                    if use_colors:
                        bullet = f"{self.theme.CITATION_COLOR}{bullet_sym}{self.theme.RESET}"
                    ln = self._re_bullet_simple.sub(f"{sp}{bullet} ", ln, count=1)
                
                out_lines.append(ln)

            s = "\n".join(out_lines)

            # Inline elements
            def _bold_sub(m):
                inner = m.group(1)
                if use_colors:
                    return f"{Style.BRIGHT}{inner}{self.theme.RESET}"
                return inner
            s = self._re_bold_simple.sub(_bold_sub, s)

            def _ital_sub(m):
                inner = m.group(1)
                if use_colors:
                    return f"{Style.BRIGHT}{inner}{self.theme.RESET}"
                return inner
            s = self._re_ital_star_simple.sub(_ital_sub, s)
            s = self._re_ital_us_simple.sub(_ital_sub, s)

            def _code_sub(m):
                inner = m.group(1)
                if use_colors:
                    return f"{self.theme.METADATA_COLOR}{inner}{self.theme.RESET}"
                return inner
            s = self._re_code_simple.sub(_code_sub, s)

            def _link_sub(m):
                txt, url = m.group(1), m.group(2)
                if use_colors:
                    return f"{Style.BRIGHT}{txt}{self.theme.RESET} ({self.theme.METADATA_COLOR}{url}{self.theme.RESET})"
                return f"{txt} ({url})"
            s = self._re_link_simple.sub(_link_sub, s)

            return s
        except Exception:
            return text
    
    def format_citations(self, citations: List[Any]) -> str:
        """Format citations with improved visual hierarchy."""
        if not citations:
            return ""
        
        lines = []
        header = self._apply_color("Citations", self.theme.SECTION_COLOR)
        lines.append(f"\n{header}")
        
        for idx, c in enumerate(citations, start=1):
            # Support both plain URL strings and dict objects
            if isinstance(c, str):
                url = c
                title = None
            elif isinstance(c, dict):
                url = c.get("url") or c.get("source") or c.get("link") or c.get("href")
                title = c.get("title") or c.get("name")
                if not url:
                    # Fallback to simple string representation for performance
                    url = str(c)
            else:
                url = str(c)
                title = None
            
            # Format citation with subdued colors and proper wrapping
            citation_text = self._apply_color(f"[{idx}]", self.theme.CITATION_COLOR)
            if title:
                citation_line = f"{title} - {url}"
            else:
                citation_line = url
            
            # Wrap citation text to prevent line breaks in URLs
            wrapped_citation = self._wrap_text(
                citation_line,
                indent=len(f"{self.theme.INDENT}{citation_text} ")
            )
            
            # Apply prefix to first line only
            citation_lines = wrapped_citation.split('\n')
            if citation_lines:
                citation_lines[0] = f"{self.theme.INDENT}{citation_text} {citation_lines[0].lstrip()}"
                lines.extend(citation_lines)
        
        return "\n".join(lines)
    
    def format_usage(self, usage: Dict[str, Any]) -> str:
        """Format usage statistics with compact, visually separated format."""
        if not usage:
            return ""
        
        lines = []
        header = self._apply_color("Tokens", self.theme.SECTION_COLOR)
        lines.append(f"\n{header}")
        
        for k, v in usage.items():
            key_text = self._apply_color(k, self.theme.METADATA_COLOR)
            lines.append(f"{self.theme.INDENT}{key_text}: {v}")
        
        return "\n".join(lines)
    
    def format_section_header(self, title: str) -> str:
        """Format section headers with visual emphasis."""
        return self._apply_color(title, self.theme.SECTION_COLOR)
    
    def format_command_help(self, commands: Dict[str, str]) -> str:
        """Format command help with proper alignment and colors."""
        lines = []
        max_cmd_len = max(len(cmd) for cmd in commands.keys()) if commands else 0
        
        for cmd, desc in commands.items():
            cmd_colored = self._apply_color(cmd.ljust(max_cmd_len), self.theme.USER_COLOR)
            lines.append(f"{self.theme.INDENT}{cmd_colored}  {desc}")
        
        return "\n".join(lines)
    
    def format_status_message(self, message: str, status_type: str = "info") -> str:
        """Format status messages with appropriate colors."""
        color_map = {
            "success": self.theme.SUCCESS_COLOR,
            "error": self.theme.ERROR_COLOR,
            "warning": self.theme.WARNING_COLOR,
            "info": self.theme.METADATA_COLOR
        }
        color = color_map.get(status_type, self.theme.METADATA_COLOR)
        return self._apply_color(message, color)

    def format_json_output(self, data: Dict[str, Any]) -> str:
        """Format JSON output with indentation and syntax highlighting (if colors enabled)."""
        try:
            json_str = json.dumps(data, indent=2)
            # Skip expensive regex formatting for large JSON responses (>10KB)
            if self.use_colors and len(json_str) < 10000:
                json_str = re.sub(r'(".*?")(?=:)', r'{}\1{}'.format(self.theme.METADATA_COLOR, self.theme.RESET), json_str) # Keys
                json_str = re.sub(r'(": ")(.*?)(?=")', r'\1{}{}{}'.format(self.theme.AI_COLOR, r'\2', self.theme.RESET), json_str) # String values
                json_str = re.sub(r'(\b\d+\b)', r'{}\1{}'.format(self.theme.CITATION_COLOR, self.theme.RESET), json_str) # Numbers
                json_str = re.sub(r'(\btrue\b|\bfalse\b|\bnull\b)', r'{}\1{}'.format(self.theme.USER_COLOR, self.theme.RESET), json_str) # Booleans and null
            return json_str
        except Exception:
            return str(data) # Fallback to plain string if JSON formatting fails


# --- Layout Utilities ---

def add_spacing(lines: int = 1) -> str:
    """Add vertical spacing."""
    return "\n" * lines


def create_border(text: str, style: str = "simple") -> str:
    """Create a border around text."""
    if style == "simple":
        border_char = "─"
        width = len(text) + 4
        border = border_char * width
        return f"{border}\n  {text}  \n{border}"
    return text


def indent_text(text: str, levels: int = 1) -> str:
    """Indent text by specified levels."""
    indent = FormattingTheme().INDENT * levels
    return "\n".join(f"{indent}{line}" for line in text.split("\n"))


def wrap_section(content: str, title: Optional[str] = None) -> str:
    """Wrap content in a section with optional title."""
    lines = []
    if title:
        lines.append(FormattingTheme().SECTION_SEPARATOR)
        lines.append(title)
        lines.append(FormattingTheme().SECTION_SEPARATOR)
    lines.append(content)
    return "\n".join(lines)


# --- Legacy Functions (for backward compatibility) ---

def print_usage(formatter: OutputFormatter, usage: Dict[str, Any]) -> None:
    formatted = formatter.format_usage(usage)
    if formatted:
        print(formatted)


def print_citations(formatter: OutputFormatter, citations: List[Any]) -> None:
    formatted = formatter.format_citations(citations)
    if formatted:
        print(formatted)


def make_payload_from_messages(
    messages: List[Dict[str, Any]], args: argparse.Namespace
) -> Dict[str, Any]:
    def fix_message_alternation(msgs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        # Ensures messages alternate between user/tool and assistant
        if not msgs:
            return msgs
        fixed = []
        last_role = None
        for msg in msgs:
            role = msg.get("role")
            if last_role == role and role in ("user", "tool"):
                # Insert assistant placeholder to fix alternation
                fixed.append({"role": "assistant", "content": "(placeholder)"})
            fixed.append(msg)
            last_role = role
        return fixed

    processed_messages = []
    for msg in fix_message_alternation(messages):
        if isinstance(msg.get("content"), str):
            processed_messages.append(
                {
                    "role": msg["role"],
                    "content": [{"type": "text", "text": msg["content"]}],
                }
            )
        else:
            processed_messages.append(msg)

    payload: Dict[str, Any] = {
        "model": args.model,
        "messages": processed_messages,
    }
    if args.stream:
        payload["stream"] = True
    if getattr(args, "max_tokens", None):
        payload["max_tokens"] = args.max_tokens
    if getattr(args, "temperature", None) is not None:
        payload["temperature"] = args.temperature
    if getattr(args, "top_p", None) is not None:
        payload["top_p"] = args.top_p
    if args.academic:
        payload["search_filter"] = "academic"
    if args.domain:
        payload["search_domain_filter"] = [args.domain]
    if args.recency:
        payload["search_recency_filter"] = args.recency
    if args.json_schema:
        payload["response_format"] = {
            "type": "json_schema",
            "json_schema": {
                "schema": {
                    "type": "object",
                    "properties": {
                        "startups": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "company_name": {"type": "string"},
                                    "funding_amount": {"type": "string"},
                                    "focus_area": {"type": "string"},
                                },
                                "required": [
                                    "company_name",
                                    "funding_amount",
                                    "focus_area",
                                ],
                            },
                        }
                    },
                    "required": ["startups"],
                }
            },
        }
    return payload


def chat_loop(
    args: argparse.Namespace,
    headers: Dict[str, str],
    config: ApiConfig,
    messages: Optional[List[Dict[str, Any]]] = None,
) -> int:
    last_async_id: Optional[str] = None

    # Initialize formatter for output formatting
    theme_cfg = getattr(args, 'theme', None) if hasattr(args, 'theme') else None
    use_colors = True  # Default to using colors
    md_render = getattr(args, 'md_render', True) if hasattr(args, 'md_render') else True
    formatter = OutputFormatter(use_colors=use_colors, theme_cfg=theme_cfg, md_render=md_render)
    
    # Set table rendering preference
    disable_tables = getattr(args, 'disable_table_rendering', False)
    formatter._disable_tables = disable_tables
    
    # Apply persistent table rendering setting
    disable_tables = bool(getattr(args, 'disable_table_rendering', False))
    formatter._disable_tables = disable_tables

    # Initialize attachment state for pre-prompt document handling
    attachment_state = AttachmentState()

    # Session state for this chat
    current_session_enabled: bool = not getattr(args, "no_session", False)
    current_session_name: Optional[str] = getattr(args, "session_name", None)
    current_session_file: Optional[str] = _session_file_for(current_session_name) if current_session_name else None

    def _suggest_session_name() -> str:
        # Only suggest a session name if there is a user message
        import datetime as _dt
        ts = _dt.datetime.now().strftime("%Y%m%d-%H%M%S")
        title: str = ""
        try:
            for m in (messages or []):
                if m.get("role") == "user":
                    c = m.get("content")
                    if isinstance(c, str) and c.strip():
                        title = c.strip()
                        break
                    elif isinstance(c, list):
                        for part in c:
                            if isinstance(part, dict) and part.get("type") == "text" and part.get("text", "").strip():
                                title = part.get("text").strip()
                                break
                        if title:
                            break
        except Exception:
            title = ""
        if not title:
            return ""  # No user message, no session name
        first_line = title.splitlines()[0].strip()
        if len(first_line) > 60:
            first_line = first_line[:60]
        slug = "".join(ch if ch.isalnum() or ch in ("-", "_") else "-" for ch in first_line)
        try:
            import re as _re
            slug = _re.sub(r"-+", "-", slug).strip("-") or "session"
        except Exception:
            slug = slug.strip("-") or "session"
        return f"{slug}-{ts}"

    def _session_write(msg: Dict[str, Any]) -> None:
        if not current_session_enabled:
            return
        if msg.get("role") != "user":
            return
        nonlocal current_session_file
        try:
            if current_session_file is None:
                # auto-generate name if not set
                nonlocal current_session_name
                current_session_name = _suggest_session_name()
                if not current_session_name:
                    return  # No user message, do not create file
                current_session_file = _session_file_for(current_session_name)
            os.makedirs(os.path.dirname(current_session_file), exist_ok=True)
            with open(current_session_file, "a", encoding="utf-8") as f:
                f.write(json.dumps(msg, ensure_ascii=False) + "\n")
        except Exception:
            pass

    def _session_write_all(ms: List[Dict[str, Any]]) -> None:
        if not current_session_enabled:
            return
        user_msgs = [m for m in ms if m.get("role") == "user"]
        if not user_msgs:
            return  # No user messages, do not create file
        nonlocal current_session_file
        try:
            if current_session_file is None:
                nonlocal current_session_name
                current_session_name = _suggest_session_name()
                if not current_session_name:
                    return  # No user message, do not create file
                current_session_file = _session_file_for(current_session_name)
            os.makedirs(os.path.dirname(current_session_file), exist_ok=True)
            with open(current_session_file, "w", encoding="utf-8") as f:
                for m in ms:
                    f.write(json.dumps(m, ensure_ascii=False) + "\n")
        except Exception:
            pass

    def _session_load(name: str) -> List[Dict[str, Any]]:
        return _load_session_messages(name)

    # If session is enabled and name provided at startup, preload it (accept history as well)
    if current_session_enabled and current_session_name:
        try:
            prior = _session_load(current_session_name)
            if prior:
                if messages is None:
                    messages = []
                messages.extend(prior)
        except Exception:
            pass

    # Inject persistent system message if configured and not already present
    persistent_sys = getattr(args, "persistent_system", None)
    if persistent_sys and isinstance(persistent_sys, str) and persistent_sys.strip():
        if messages is None:
            messages = []
        # Only add if there's no system message already, or if the existing one doesn't contain the persistent message
        has_persistent = False
        if messages and messages[0].get("role") == "system":
            existing_content = messages[0].get("content", "")
            if persistent_sys.strip() in existing_content:
                has_persistent = True
        if not has_persistent:
            if messages and messages[0].get("role") == "system":
                # Prepend to existing system message
                existing = messages[0]["content"]
                messages[0]["content"] = f"{persistent_sys.strip()}\n\n{existing}"
            else:
                # Insert as new system message
                messages.insert(0, {"role": "system", "content": persistent_sys.strip()})

    # Only persist session if there is at least one user message (after system message injection)
    if current_session_enabled and messages and any(m.get("role") == "user" for m in messages):
        _session_write_all(messages)

    # Inject persistent system message if configured and not already present
    persistent_sys = getattr(args, "persistent_system", None)
    if persistent_sys and isinstance(persistent_sys, str) and persistent_sys.strip():
        if messages is None:
            messages = []
        # Only add if there's no system message already, or if the existing one doesn't contain the persistent message
        has_persistent = False
        if messages and messages[0].get("role") == "system":
            existing_content = messages[0].get("content", "")
            if persistent_sys.strip() in existing_content:
                has_persistent = True
        if not has_persistent:
            if messages and messages[0].get("role") == "system":
                # Prepend to existing system message
                existing = messages[0]["content"]
                messages[0]["content"] = f"{persistent_sys.strip()}\n\n{existing}"
            else:
                # Insert as new system message
                messages.insert(0, {"role": "system", "content": persistent_sys.strip()})

    def fix_message_alternation(msgs: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        # Ensures messages alternate between user/tool and assistant
        if not msgs:
            return msgs
        fixed = []
        last_role = None
        for msg in msgs:
            role = msg.get("role")
            if last_role == role and role in ("user", "tool"):
                fixed.append({"role": "assistant", "content": "(placeholder)"})
            fixed.append(msg)
            last_role = role
        return fixed

    def fix_messages_in_place():
        nonlocal messages
        messages = fix_message_alternation(messages)

    def build_async_request(msgs: List[Dict[str, Any]]) -> Dict[str, Any]:
        # Ensure model is async-compatible
        model_val = args.model
        if model_val != "sonar-deep-research":
            print(formatter.format_status_message(
                "Note: async API only supports 'sonar-deep-research'; using it for this request. Use /model to change.", "warning"
            ))
            model_val = "sonar-deep-research"
        # Convert messages to multimodal format consistent with sync path
        processed_messages: List[Dict[str, Any]] = []
        for msg in fix_message_alternation(msgs or []):
            if isinstance(msg.get("content"), str):
                processed_messages.append({
                    "role": msg["role"],
                    "content": [{"type": "text", "text": msg["content"]}],
                })
            else:
                processed_messages.append(msg)
        # Build request
        req: Dict[str, Any] = {
            "model": model_val,
            "messages": processed_messages,
        }
        if getattr(args, "max_tokens", None):
            req["max_tokens"] = args.max_tokens
        if getattr(args, "temperature", None) is not None:
            req["temperature"] = args.temperature
        if getattr(args, "top_p", None) is not None:
            req["top_p"] = args.top_p
        if getattr(args, "search_mode", None):
            req["search_mode"] = args.search_mode
        if getattr(args, "reasoning_effort", None):
            req["reasoning_effort"] = args.reasoning_effort
        if getattr(args, "return_images", None):
            req["return_images"] = args.return_images
        if getattr(args, "return_related_questions", None):
            req["return_related_questions"] = args.return_related_questions
        if args.academic:
            req["search_filter"] = "academic"
        if args.domain:
            req["search_domain_filter"] = [args.domain]
        if args.recency:
            req["search_recency_filter"] = args.recency
        return {"request": req}

    def async_submit(msgs: List[Dict[str, Any]]) -> Optional[str]:
        try:
            body = build_async_request(msgs)
            r = requests.post(ASYNC_API_URL, headers=headers, json=body)
            try:
                r.raise_for_status()
            except requests.exceptions.HTTPError as e:
                err_text = None
                try:
                    err_text = r.text
                except Exception:
                    pass
                print(formatter.format_status_message(f"Async submit failed: {e}", "error"))
                if err_text:
                    try:
                        err_json = json.loads(err_text)
                        print(formatter.format_json_output(err_json))
                    except Exception:
                        print(err_text)
                return None
            data = r.json()
            rid = data.get("id") or data.get("request_id") or data.get("requestId")
            return rid
        except Exception as e:
            print(formatter.format_status_message(f"Async submit failed: {e}", "error"))
            return None

    def async_get(request_id: str) -> Optional[Dict[str, Any]]:
        try:
            url = f"{ASYNC_API_URL}/{request_id}"
            r = requests.get(url, headers=headers)
            r.raise_for_status()
            return r.json()
        except Exception as e:
            print(formatter.format_status_message(f"Async get failed: {e}", "error"))
            return None

    def async_list() -> Optional[Dict[str, Any]]:
        try:
            r = requests.get(ASYNC_API_URL, headers=headers)
            r.raise_for_status()
            return r.json()
        except Exception as e:
            print(formatter.format_status_message(f"Async list failed: {e}", "error"))
            return None

    def async_wait(
        request_id: str, interval: float = 1.0, timeout: float = 300.0
    ) -> Optional[Dict[str, Any]]:
        start = time.time()
        while time.time() - start < timeout:
            data = async_get(request_id)
            if not data:
                return None
            status = (data.get("status") or data.get("state") or "").upper()
            if status in ("COMPLETED", "SUCCEEDED"):
                return data
            if status in ("FAILED", "ERROR"):
                print(formatter.format_status_message("Async job failed.", "error"))
                return data
            time.sleep(interval)
        print(formatter.format_status_message("Async wait timed out.", "warning"))
        return None

    colorama_init(autoreset=True)
    welcome_header = formatter.format_section_header("askp Interactive Chat")
    print(welcome_header)
    print("Type /help for commands. Ctrl+C or /exit to quit.")
    if bool(getattr(args, "json_schema", False)):
        print(formatter.format_status_message("Structured JSON mode is ON (example schema). Type /jsonschema off to disable.", "warning"))
    print(f"{formatter.theme.SECTION_SEPARATOR}\n")
    if messages is None:
        messages = []
    buffer_lines: List[str] = []
    multiline = False

    slash_commands = {
        "/colors": "Toggle color output: /colors on|off",
        "/help": "Show this help",
        "/reset": "Reset the conversation (preserves /system)",
        "/new": "Start a new conversation (preserves /system)",
        "/model": "Switch model: /model <name>",
        "/models": "Choose a model from a list",
        "/system": "Set system prompt: /system <text> | /system show | /system persistent <text> | add | edit",
        "/academic": "Toggle academic filter: /academic on|off",
        "/domain": "Set/clear domain filter: /domain <host> | /domain (to clear)",
        "/recency": "Set recency filter: /recency day|week|month|off",
        "/jsonschema": "Toggle structured JSON (example schema; not persisted): /jsonschema on|off",
        "/citations": "Toggle printing citations (non-stream): /citations on|off",
        "/usage": "Toggle printing usage (non-stream): /usage on|off",
        "/asyncmeta": "Toggle showing async metadata (usage/citations) by default: /asyncmeta on|off",
        "/stream": "Toggle streaming: /stream on|off",
        "/mdrender": "Toggle Markdown rendering in AI output: /mdrender on|off",
        "/text": "Toggle one-shot default output mode (text vs json): /text on|off",
        "/echo": "Toggle echoing your input before the response: /echo on|off",
        "/theme": "Customize colors: /theme (interactive) or /theme reset",
        "/tableformat": "Add table formatting instructions to system message: /tableformat [on|off|show]",
        "/notables": "Disable table rendering (show raw markdown): /notables [on|off]",

        "/attach": 'Attach a local file or directory: /attach [path] [preprompt|summarize|full|all] [--as-user] [--max-files N] [--pattern ".py,.md"] [--include-hidden] (defaults to preprompt mode)',
        "/ocr": "Extract text from image file(s) via OCR. Select interactively or use 'all': /ocr [path|file] [all] [--as-user] [--max-chars N] [--max-files N] [--pattern \".png,.jpg\"] [--include-hidden]",
        "/async": "Async ops: /async submit [prompt]|list|get <id>|wait <id>",
        "/attachlimit": "Set truncation limit for /attach --as-user: /attachlimit <N>",
        "/attachments": "Show pending attachments or clear them: /attachments [clear]",
        "/save": "Save conversation to a file (JSONL): /save <path>",
        "/copy": "Copy last AI response to clipboard",
        "/settings": "Show current settings",
        "/prune": "Summarize chat and restart with the summary as context: /prune [N words] (default 200)",
        "/compact": "Alias for /prune",
        "/session": "Toggle or manage project sessions (default: on): /session [on|off|name <n>|show]",
        "/sessions": "List/load sessions: /sessions [list]|open <name|#>|merge <name|#>",
        "/clear": "Clear screen and start new conversation",
        "/exit": "Exit the chat",
        "/quit": "Exit the chat",
    }
    completer: Any = None
    session: Any = None
    history_file = os.path.join(ASKP_DATA_DIR, "askp_history")
    # Ensure .askp directory exists before using history file
    try:
        os.makedirs(ASKP_DATA_DIR, exist_ok=True)
    except Exception:
        pass
    if PROMPT_TOOLKIT:

        class SlashCompleter(PTCompleter):  # type: ignore[misc]
            def get_completions(self, document, complete_event):
                text = document.text_before_cursor
                if not text.startswith("/"):
                    return
                
                # Define subcommands for various commands
                subcommands = {
                    "/system": ["show", "clear", "persistent", "persistent add", "persistent edit", "persistent clear"],
                    "/async": ["submit", "list", "get", "wait"],
                    "/session": ["on", "off", "show", "name"],
                    "/sessions": ["list", "open", "merge"],
                    "/academic": ["on", "off"],
                    "/domain": ["clear"],
                    "/recency": ["day", "week", "month", "off"],
                    "/citations": ["on", "off"],
                    "/usage": ["on", "off"],
                    "/stream": ["on", "off"],
                    "/colors": ["on", "off"],
                    "/mdrender": ["on", "off"],
                    "/text": ["on", "off"],
                    "/echo": ["on", "off"],
                    "/asyncmeta": ["on", "off"],
                    "/jsonschema": ["on", "off"],
                    "/theme": ["reset"],
                    "/tableformat": ["on", "off", "show"],
                    "/notables": ["on", "off"],
                    "/model": AVAILABLE_MODELS,
                    "/attach": ["summarize", "full", "all", "--as-user", "--max-files", "--pattern", "--include-hidden"],
                    "/ocr": ["all", "--as-user", "--max-chars", "--max-files", "--pattern", "--include-hidden"],
                }
                
                # Parse the current input
                parts = text.split()
                if len(parts) == 0:
                    return
                    
                base_cmd = parts[0]
                
                # If we're still typing the base command
                if len(parts) == 1 and not text.endswith(" "):
                    word = document.get_word_before_cursor(WORD=True)
                    query = (word if word else "/").lower()
                    for cmd in sorted(slash_commands.keys()):
                        if query == "/" or query.lstrip("/") in cmd.lstrip("/"):
                            display_meta = slash_commands.get(cmd, "")
                            start_pos = -len(word) if word else 0
                            if PTCompletion is not None:
                                yield PTCompletion(
                                    cmd,
                                    start_position=start_pos,
                                    display=cmd,
                                    display_meta=display_meta,
                                )
                    return
                
                # If we have a base command and are looking for subcommands
                if base_cmd in subcommands:
                    available_subs = subcommands[base_cmd]
                    
                    # Get the current word being typed (subcommand)
                    if len(parts) > 1 and not text.endswith(" "):
                        # Currently typing a subcommand
                        current_sub = parts[-1].lower()
                        start_pos = -len(parts[-1])
                    else:
                        # Just finished base command, show all subcommands
                        current_sub = ""
                        start_pos = 0
                    
                    # Filter and yield matching subcommands
                    for sub in available_subs:
                        if not current_sub or current_sub in sub.lower():
                            # Create helpful descriptions for common subcommands
                            descriptions = {
                                "show": "Display current value",
                                "clear": "Clear/remove current value", 
                                "on": "Enable/turn on",
                                "off": "Disable/turn off",
                                "persistent": "Set persistent system message",
                                "submit": "Submit async request",
                                "list": "List available items",
                                "get": "Get specific item",
                                "wait": "Wait for completion",
                                "open": "Open/load item",
                                "merge": "Merge with current",
                                "name": "Set name",
                                "day": "Filter to last day",
                                "week": "Filter to last week", 
                                "month": "Filter to last month",
                                "reset": "Reset to defaults",
                                "summarize": "Summarize content",
                                "full": "Include full content",
                                "all": "Select all items",
                                "--as-user": "Insert as user message",
                                "--max-files": "Limit number of files",
                                "--pattern": "File pattern filter",
                                "--include-hidden": "Include hidden files",
                                "--max-chars": "Character limit"
                            }
                            
                            display_meta = descriptions.get(sub, f"{base_cmd} {sub}")
                            
                            if PTCompletion is not None:
                                yield PTCompletion(
                                    sub,
                                    start_position=start_pos,
                                    display=sub,
                                    display_meta=display_meta,
                                )

        completer = SlashCompleter()
        try:
            assert PT_PromptSession is not None
            if PT_FileHistory is not None:
                session = PT_PromptSession(history=PT_FileHistory(history_file))
            else:
                session = PT_PromptSession()
        except Exception:
            assert PT_PromptSession is not None
            session = PT_PromptSession()
        # Ensure suggestion menu opens automatically when starting with '/'
        if session and PT_get_app is not None and PT_get_app():
            try:
                # Some terminals require this to open completions while typing
                if hasattr(session, "app") and hasattr(
                    session.app, "complete_while_typing"
                ):
                    session.app.complete_while_typing = True  # type: ignore[attr-defined]
            except Exception:
                pass

    # default truncation limit for /attach --as-user
    # Ensure a valid integer default even if the parser created the attribute with None
    if (
        not isinstance(getattr(args, "attach_truncation_limit", None), int)
        or getattr(args, "attach_truncation_limit", 0) <= 0
    ):
        args.attach_truncation_limit = 8000

    prefill = ""
    while True:
        try:
            if session is not None:
                # Use prompt_toolkit HTML formatting to avoid raw ANSI codes
                if not multiline:
                    display_prompt = (
                        PT_HTML(
                            "<ansicyan>You></ansicyan> <ansiyellow>(type / for commands)</ansiyellow> "
                        )
                        if PT_HTML
                        else "You> (type / for commands) "
                    )
                else:
                    display_prompt = (
                        PT_HTML("<ansicyan>...></ansicyan> ") if PT_HTML else "...> "
                    )
                assert session is not None
                line = session.prompt(
                    display_prompt,
                    completer=completer,
                    complete_while_typing=True,
                    default=prefill,
                )
            else:
                hint = f" {formatter.theme.METADATA_COLOR}(type / for commands){formatter.theme.RESET}" if PROMPT_TOOLKIT else ""
                if not multiline:
                    prompt = f"{formatter.theme.USER_COLOR}❯{formatter.theme.RESET} {hint}"
                else:
                    prompt = f"{formatter.theme.USER_COLOR}│{formatter.theme.RESET} "
                line = input(prompt)
        except (EOFError, KeyboardInterrupt):
            print()  # newline
            return 0
        if line is None:
            continue
        prefill = ""
        # Multiline toggles
        if line.strip() == "" and multiline and buffer_lines:
            user_input = "\n".join(buffer_lines).strip()
            buffer_lines.clear()
            multiline = False
        elif line.strip().endswith("\\"):
            multiline = True
            buffer_lines.append(line.rstrip("\\").rstrip())
            continue
        else:
            user_input = line.strip()
        if not user_input:
            continue
        if user_input.startswith("/"):
            raw = user_input[1:].strip()
            cmd = raw.split()[0] if raw else ""
            args_rest = raw[len(cmd) :].strip() if cmd else ""
            if cmd in ("exit", "quit", "q"):
                return 0
            if cmd in ("help", "h"):
                commands = {
                    "/ocr [path|file] [all] [--as-user] [--max-chars N] [--max-files N] [--pattern \".png,.jpg\"] [--include-hidden]": "OCR image(s) and print; optionally insert as user",
                    "/help": "Show this help",
                    "/new|/reset": "New conversation (preserve /system)",
                    "/model <n>": "Switch model (sonar|sonar-pro|sonar-reasoning|sonar-reasoning-pro)",
                    "/models": "Select model from a menu",
                    "/system <text> | /system show | /system persistent <text> | add | edit": "Set/view system prompt or manage persistent system message",
                    "/academic on|off": "Toggle academic search filter",
                    "/domain [host]": "Set/clear domain filter (no arg clears)",
                    "/recency <v>": "day|week|month|off",
                    "/jsonschema on|off": "Toggle structured JSON (example schema; not persisted)",
                    "/citations on|off": "Toggle printing citations (non-stream)",
                    "/usage on|off": "Toggle printing usage (non-stream)",
                    "/asyncmeta on|off": "Toggle showing async metadata (usage/citations) by default",
                    "/stream on|off": "Toggle streaming",
                    "/mdrender on|off": "Toggle Markdown rendering in AI output",
                    "/echo on|off": "Toggle echoing your input before the response",
                    "/theme": "Customize colors (interactive) or '/theme reset'",
                    "/tableformat on|off|show": "Manage table formatting instructions in system message",
                    "/notables on|off": "Disable table rendering (show raw markdown)",
                    "/mdrender on|off": "Toggle Markdown rendering in AI output",
                    "/attach [path] [summarize|full|all] [--as-user] [--max-files N] [--pattern \".py,.md\"] [--include-hidden]": "Attach a file or directory (txt/md/pdf/docx/html; images via OCR). Default: summarize. If [path] is omitted, uses current directory.",
                    "/async submit [prompt]|list|get <id>|wait <id>": "Async API helpers",
                    "/attachlimit <N>": "Set truncation limit for /attach --as-user",
                    "/save <path>": "Save conversation to JSONL",
                    "/copy": "Copy last AI response to clipboard",
                    "/settings": "Show current settings",
                    "/session [on|off|name <n>|show]": "Toggle or manage project session persistence (default: on)",
                    "/sessions [list]|open <name|#>|merge <name|#>": "List or load previous sessions",
                    "/prune [N]": "Summarize chat and restart with the summary as context (default N=200)",
                    "/clear": "Clear screen and start new conversation",
                    "/exit|/quit": "Exit the chat"
                }
                
                help_header = formatter.format_section_header("Available Commands")
                help_content = formatter.format_command_help(commands)
                print(help_header)
                print(help_content)
                continue
            if cmd == "reset" or cmd == "new":
                # preserve any leading system prompt
                sys_msg = (
                    messages[0]
                    if messages and messages[0].get("role") == "system"
                    else None
                )
                messages.clear()
                if sys_msg:
                    messages.append(sys_msg)
                print(formatter.format_status_message("Started a new conversation.", "success"))
                _session_write_all(messages)
                continue
            if cmd == "clear":
                # Clear terminal screen
                try:
                    os.system("cls" if os.name == "nt" else "clear")
                except Exception:
                    pass
                # Clear conversation history (preserve system prompt like /new and /reset)
                sys_msg = (
                    messages[0]
                    if messages and messages[0].get("role") == "system"
                    else None
                )
                messages.clear()
                if sys_msg:
                    messages.append(sys_msg)
                print(formatter.format_status_message("Cleared screen and started a new conversation.", "success"))
                _session_write_all(messages)
                continue
            if cmd == "model":
                parts = args_rest.split()
                if len(parts) == 0 or parts[0] == "":
                    prefill = "/model "
                elif len(parts) == 1 and parts[0] in AVAILABLE_MODELS:
                    args.model = parts[0]
                    print(formatter.format_status_message(f"Model set to {args.model}", "success"))
                    save_prefs_from_args(args)
                else:
                    print(formatter.format_status_message(f"Usage: /model <name> where <name> in {AVAILABLE_MODELS}", "info"))
                continue
            if cmd == "models":
                # Terminal-based selector: list models with index and prompt for choice
                models_header = formatter.format_section_header("Available Models")
                print(models_header)
                print(f"{formatter.theme.SECTION_SEPARATOR}")
                for idx, m in enumerate(AVAILABLE_MODELS, start=1):
                    print(f"  {idx}) {m}")
                print(f"{formatter.theme.SECTION_SEPARATOR}")
                try:
                    choice = input("Select model by number: ").strip()
                    if not choice:
                        continue
                    sel = int(choice)
                    if 1 <= sel <= len(AVAILABLE_MODELS):
                        args.model = AVAILABLE_MODELS[sel - 1]
                        print(formatter.format_status_message(f"Model set to {args.model}", "success"))
                    else:
                        print(formatter.format_status_message("Invalid selection.", "error"))
                except Exception:
                    print(formatter.format_status_message("Invalid input.", "error"))
                continue
            if cmd == "system":
                # Manage system prompt: show, clear, persistent, or set
                sub = args_rest.strip()
                # Show current and prefill for editing when no text is provided
                if sub == "" or sub.lower() in ("show", "view"):
                    cur = (
                        messages[0]["content"]
                        if messages and messages[0].get("role") == "system"
                        else None
                    )
                    persistent_sys = getattr(args, "persistent_system", None)
                    if cur or persistent_sys:
                        system_header = formatter.format_section_header("Current System Prompt")
                        print(system_header)
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        if persistent_sys:
                            print(formatter.format_status_message("Persistent system message:", "info"))
                            print(persistent_sys)
                            if cur and persistent_sys.strip() not in cur:
                                print(f"\n{formatter.format_status_message('Session system message:', 'info')}")
                                print(cur)
                            elif cur:
                                print(f"\n{formatter.format_status_message('(Session system message is combined with persistent message above)', 'info')}")
                        elif cur:
                            print(cur)
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        print(formatter.format_status_message(
                            "(Tip: edit the line prefilled below to update; type '/system clear' to remove session message; '/system persistent clear' to remove persistent)", "info"
                        ))
                        prefill = "/system " + (cur if cur else "")
                    else:
                        print(formatter.format_status_message("No system prompt set.", "info"))
                        prefill = "/system "
                    continue
                if sub.lower() in ("clear", "off", "remove"):
                    if messages and messages[0].get("role") == "system":
                        messages.pop(0)
                        print(formatter.format_status_message("Session system prompt cleared.", "success"))
                    else:
                        print(formatter.format_status_message("No session system prompt to clear.", "warning"))
                    continue
                if sub.lower().startswith("persistent"):
                    # Handle persistent system message
                    persistent_args = sub[len("persistent"):].strip()
                    if persistent_args.lower() in ("clear", "off", "remove", ""):
                        if persistent_args.lower() in ("clear", "off", "remove"):
                            setattr(args, "persistent_system", None)
                            save_prefs_from_args(args)
                            print(formatter.format_status_message("Persistent system message cleared.", "success"))
                        else:
                            # Show current persistent message
                            current_persistent = getattr(args, "persistent_system", None)
                            if current_persistent:
                                print(formatter.format_status_message("Current persistent system message:", "info"))
                                print(current_persistent)
                                prefill = f"/system persistent {current_persistent}"
                            else:
                                print(formatter.format_status_message("No persistent system message set.", "info"))
                                prefill = "/system persistent "
                    elif persistent_args.lower() == "edit":
                        # Edit persistent system message in external editor
                        
                        current_persistent = getattr(args, "persistent_system", None) or ""
                        
                        # Create temporary file
                        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
                            f.write(current_persistent)
                            temp_file = f.name
                        
                        try:
                            # Get editor from environment or use sensible defaults
                            editor = os.environ.get('EDITOR') or os.environ.get('VISUAL')
                            if not editor:
                                # Try common editors - use simple existence check
                                if os.name == 'nt':  # Windows
                                    # On Windows, notepad is always available
                                    editor = 'notepad'
                                else:
                                    # On Unix-like systems, try common editors
                                    for candidate in ['nano', 'vim', 'vi']:
                                        try:
                                            subprocess.run(['which', candidate], capture_output=True, check=True, timeout=1)
                                            editor = candidate
                                            break
                                        except (subprocess.CalledProcessError, FileNotFoundError, subprocess.TimeoutExpired):
                                            continue
                            
                            if not editor:
                                print(formatter.format_status_message("No editor found. Set $EDITOR environment variable.", "error"))
                                continue
                            
                            print(formatter.format_status_message(f"Opening editor: {editor}", "info"))
                            
                            # Open editor
                            result = subprocess.run([editor, temp_file])
                            
                            if result.returncode == 0:
                                # Read the edited content
                                with open(temp_file, 'r') as f:
                                    edited_content = f.read().strip()
                                
                                if edited_content != current_persistent:
                                    if edited_content:
                                        setattr(args, "persistent_system", edited_content)
                                        save_prefs_from_args(args)
                                        print(formatter.format_status_message("Persistent system message updated.", "success"))
                                    else:
                                        setattr(args, "persistent_system", None)
                                        save_prefs_from_args(args)
                                        print(formatter.format_status_message("Persistent system message cleared.", "success"))
                                else:
                                    print(formatter.format_status_message("No changes made.", "info"))
                            else:
                                print(formatter.format_status_message("Editor exited with error. No changes saved.", "warning"))
                        
                        finally:
                            # Clean up temporary file
                            try:
                                os.unlink(temp_file)
                            except OSError:
                                pass
                    
                    elif persistent_args.lower().startswith("add "):
                        # Add to persistent system message
                        add_text = persistent_args[4:].strip()  # Remove "add " prefix
                        if not add_text:
                            print(formatter.format_status_message("No text provided to add.", "warning"))
                            continue
                        
                        current_persistent = getattr(args, "persistent_system", None)
                        if current_persistent:
                            # Append with double newline separator
                            new_persistent = f"{current_persistent}\n\n{add_text}"
                        else:
                            # First persistent message
                            new_persistent = add_text
                        
                        setattr(args, "persistent_system", new_persistent)
                        save_prefs_from_args(args)
                        print(formatter.format_status_message("Text added to persistent system message.", "success"))
                        print(formatter.format_status_message("This will be automatically added to all new conversations.", "info"))
                    
                    else:
                        # Set persistent system message (original behavior)
                        setattr(args, "persistent_system", persistent_args)
                        save_prefs_from_args(args)
                        print(formatter.format_status_message("Persistent system message saved.", "success"))
                        print(formatter.format_status_message("This will be automatically added to all new conversations.", "info"))
                    continue
                # Otherwise, set/update system prompt with provided text
                system_text = args_rest
                if messages and messages[0].get("role") == "system":
                    messages[0]["content"] = system_text
                    print(formatter.format_status_message("System prompt updated.", "success"))
                else:
                    messages.insert(0, {"role": "system", "content": system_text})
                    print(formatter.format_status_message("System prompt set.", "success"))
                _session_write_all(messages)
                continue
            if cmd == "ocr":
                # /ocr [path|file] [all] [--as-user] [--max-chars N] [--max-files N] [--pattern ".png,.jpg"] [--include-hidden]
                try:
                    parts = shlex.split(args_rest)
                except Exception:
                    parts = args_rest.split()
                as_user = False
                max_chars: Optional[int] = None
                max_files: Optional[int] = None
                include_hidden = False
                patterns: Optional[Set[str]] = None
                core: List[str] = []
                i = 0
                while i < len(parts):
                    p = parts[i]
                    if p == "--as-user":
                        as_user = True
                        i += 1
                        continue
                    if p == "--max-chars":
                        try:
                            max_chars = int(parts[i + 1])
                            i += 2
                            continue
                        except Exception:
                            print(formatter.format_status_message("Invalid --max-chars value; expecting integer.", "error"))
                            i += 1
                            continue
                    if p == "--max-files":
                        try:
                            max_files = int(parts[i + 1])
                            i += 2
                            continue
                        except Exception:
                            print(formatter.format_status_message("Invalid --max-files value; expecting integer.", "error"))
                            i += 1
                            continue
                    if p == "--pattern":
                        try:
                            raw = parts[i + 1]
                            patterns = {
                                x.strip().lower() for x in raw.split(",") if x.strip()
                            }
                            i += 2
                            continue
                        except Exception:
                            print(formatter.format_status_message(
                                "Invalid --pattern value; expecting comma-separated extensions like .png,.jpg", "error"
                            ))
                            i += 1
                            continue
                    if p == "--include-hidden":
                        include_hidden = True
                        i += 1
                        continue
                    if p.startswith("-"):
                        i += 1
                        continue
                    core.append(p)
                    i += 1
                # Parse positional tokens: [path|file] [all]
                select_all = False
                path_token: Optional[str] = None
                for tok in core:
                    tl = tok.lower()
                    if tl == "all":
                        select_all = True
                    elif path_token is None:
                        path_token = tok
                if path_token is None:
                    path_token = "."
                target_path = os.path.abspath(path_token)
                try:
                    from PIL import Image
                    import pytesseract
                except Exception as e:
                    print(formatter.format_status_message(f"OCR dependencies missing: {e}", "error"))
                    continue
                # If a directory, scan for supported images and select
                if os.path.isdir(target_path):
                    supported_ext = {".png", ".jpg", ".jpeg", ".tif", ".tiff"}
                    files: List[str] = []
                    for root, _, fnames in os.walk(target_path):
                        for fn in fnames:
                            rb = os.path.basename(os.path.abspath(root))
                            if not include_hidden and (
                                fn.startswith(".") or rb.startswith(".")
                            ):
                                continue
                            extc = os.path.splitext(fn)[1].lower()
                            if patterns is not None and extc not in patterns:
                                continue
                            if extc in supported_ext:
                                files.append(os.path.join(root, fn))
                    if not files:
                        print(formatter.format_status_message("No images found for OCR in directory.", "warning"))
                        continue
                    limit = (
                        max_files
                        if isinstance(max_files, int) and max_files > 0
                        else 50
                    )
                    shown = files[:limit]
                    selected_files: Optional[List[str]] = None
                    if select_all:
                        targets = shown
                    else:
                        if PROMPT_TOOLKIT and len(shown) > 1:
                            print(formatter.format_status_message(
                                f"Found {len(files)} images. Top choice is [ALL]. Use Up/Down to navigate, Enter to toggle, Space to confirm. Press Esc to cancel. Press Space with none selected (or [ALL] selected) to OCR all shown.", "info"
                            ))
                            try:
                                from prompt_toolkit.widgets import CheckboxList, Frame
                                from prompt_toolkit.layout.containers import HSplit
                                from prompt_toolkit.layout import Layout
                                from prompt_toolkit.application import Application
                                from prompt_toolkit.key_binding import KeyBindings

                                rel_items = [("__ALL__", "[ALL] (select all)")]
                                for fp in shown:
                                    rel = os.path.relpath(fp, target_path)
                                    rel_items.append((fp, rel))
                                checklist = CheckboxList(values=rel_items)
                                container = HSplit(
                                    [
                                        Frame(
                                            checklist,
                                            title=f"Select images (showing first {len(shown)} of {len(files)})",
                                        )
                                    ]
                                )
                                kb = KeyBindings()

                                @kb.add("enter", eager=True)
                                def _(event):  # type: ignore[func-annotations]
                                    try:
                                        idx = getattr(checklist, "_selected_index", None)
                                        if idx is None:
                                            return
                                        val = checklist.values[idx][0]
                                        cur = set(checklist.current_values)
                                        if val in cur:
                                            cur.remove(val)
                                        else:
                                            cur.add(val)
                                        checklist.current_values = list(cur)
                                        event.app.invalidate()
                                    except Exception:
                                        pass

                                @kb.add("space", eager=True)
                                def __(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=True)

                                @kb.add("escape", eager=True)
                                def ___(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=None)

                                app = Application(
                                    layout=Layout(container),
                                    key_bindings=kb,
                                    full_screen=False,
                                )
                                ret = app.run()
                                if ret is None:
                                    print(formatter.format_status_message("Selection canceled.", "warning"))
                                    continue
                                cur = list(checklist.current_values)
                                if not cur or "__ALL__" in cur:
                                    selected_files = None
                                else:
                                    selected_files = [v for v in cur if v != "__ALL__"]
                            except Exception:
                                selected_files = None
                        else:
                            print(formatter.format_status_message(
                                f"Found {len(files)} images. Showing first {len(shown)}. Enter comma-separated indices (1-based) to select, or press Enter for all shown. Type 'c' to cancel.", "info"
                            ))
                            try:
                                for i, fp in enumerate(shown, 1):
                                    print(f"  {i:2d}) {os.path.relpath(fp, target_path)}")
                                line = input(
                                    "Select images (e.g., 1,3,5) or Enter for all (or 'c' to cancel): "
                                ).strip()
                                if line.lower() in ("c", "q", "cancel", "quit"):
                                    print(formatter.format_status_message("Selection canceled.", "warning"))
                                    continue
                                if line:
                                    picks: List[str] = []
                                    for tok in line.split(","):
                                        tok = tok.strip()
                                        if tok.isdigit():
                                            idx = int(tok)
                                            if 1 <= idx <= len(shown):
                                                picks.append(shown[idx - 1])
                                    if picks:
                                        selected_files = picks
                            except Exception:
                                selected_files = None
                        targets = selected_files if selected_files else shown
                    # OCR each target
                    combined_blocks: List[str] = []
                    print(formatter.format_status_message(f"Running OCR on {len(targets)} image(s)...", "info"))
                    for fp in targets:
                        try:
                            text = pytesseract.image_to_string(Image.open(fp)) or ""
                        except Exception as e:
                            print(formatter.format_status_message(f"OCR failed for {fp}: {e}", "error"))
                            text = ""
                        trimmed = text.strip()
                        if max_chars is not None and max_chars > 0:
                            trimmed = trimmed[:max_chars]
                        header = f"### {os.path.relpath(fp, target_path)}"
                        if trimmed:
                            ocr_header = formatter.format_section_header("OCR Text Extracted")
                            print(ocr_header)
                            print(header)
                            print(f"{formatter.theme.SECTION_SEPARATOR}")
                            print(trimmed)
                            print(f"{formatter.theme.SECTION_SEPARATOR}")
                            combined_blocks.append(f"{header}\n{trimmed}")
                        else:
                            print(formatter.format_status_message(f"No text detected by OCR in {fp}.", "warning"))
                    if as_user and combined_blocks:
                        limit = getattr(args, "attach_truncation_limit", 8000)
                        lim = limit if isinstance(limit, int) and limit > 0 else 8000
                        # Clip each block to lim to keep message reasonable
                        clipped_blocks = []
                        for blk in combined_blocks:
                            clipped_blocks.append(
                                blk[:lim]
                                + ("\n... [truncated]" if len(blk) > lim else "")
                            )
                        messages.append(
                            {
                                "role": "user",
                                "content": f"[OCR from directory {os.path.basename(target_path)}]\n\n"
                                + "\n\n".join(clipped_blocks),
                            }
                        )
                        print(formatter.format_status_message("Inserted OCR text as a user message.", "success"))
                        _session_write({
                            "role": "user",
                            "content": f"[OCR from directory {os.path.basename(target_path)}]\n\n" + "\n\n".join(clipped_blocks),
                        })
                    continue
                # If a file, OCR single image
                try:
                    text = pytesseract.image_to_string(Image.open(target_path))
                except FileNotFoundError:
                    print(formatter.format_status_message("File not found.", "error"))
                    continue
                except Exception as e:
                    print(formatter.format_status_message(f"OCR failed: {e}", "error"))
                    continue
                if not text.strip():
                    print(formatter.format_status_message("No text detected by OCR.", "warning"))
                    continue
                ocr_header = formatter.format_section_header("OCR Text Extracted")
                print(ocr_header)
                print(f"{formatter.theme.SECTION_SEPARATOR}")
                trimmed = text.strip()
                if max_chars is not None and max_chars > 0:
                    trimmed = trimmed[:max_chars]
                print(trimmed)
                print(f"{formatter.theme.SECTION_SEPARATOR}")
                if as_user:
                    limit = getattr(args, "attach_truncation_limit", 8000)
                    lim = limit if isinstance(limit, int) and limit > 0 else 8000
                    clipped = trimmed[:lim] + (
                        "\n... [truncated]" if len(trimmed) > lim else ""
                    )
                    messages.append(
                        {
                            "role": "user",
                            "content": f"[OCR from {os.path.basename(target_path)}]\n\n{clipped}",
                        }
                    )
                    print(formatter.format_status_message("Inserted OCR text as a user message.", "success"))
                    _session_write({
                        "role": "user",
                        "content": f"[OCR from {os.path.basename(target_path)}]\n\n{clipped}",
                    })
                continue
            if cmd == "attach":

                # /attach [path] [summarize|full]
                try:
                    parts = shlex.split(args_rest)
                except Exception:
                    parts = args_rest.split()
                # parse flags and positional args
                as_user = False
                max_files = None
                include_hidden = False
                debug_scan = False
                patterns: Optional[Set[str]] = None
                core: List[str] = []
                i = 0
                while i < len(parts):
                    p = parts[i]
                    if p == "--as-user":
                        as_user = True
                        i += 1
                        continue
                    if p == "--debug":
                        debug_scan = True
                        i += 1
                        continue
                    if p == "--include-hidden":

                        include_hidden = True
                        i += 1
                        continue
                    if p == "--pattern":
                        try:
                            raw = parts[i + 1]
                            patterns = {
                                x.strip().lower() for x in raw.split(",") if x.strip()
                            }
                            i += 2
                            continue
                        except Exception:
                            print(
                                "Invalid --pattern value; expecting comma-separated extensions like .py,.md"
                            )
                            i += 1
                            continue
                    if p == "--max-files":
                        try:
                            max_files = int(parts[i + 1])
                            i += 2
                            continue
                        except Exception:
                            print(formatter.format_status_message("Invalid --max-files value; expecting integer.", "error"))
                            i += 1
                            continue
                    if p.startswith("-"):
                        i += 1
                        continue
                    core.append(p)
                    i += 1
                # Parse positional args: [path] [preprompt|summarize|full|all] in any order; default path='.' and mode='preprompt'
                select_all = False
                path_token = None
                mode = "preprompt"  # Changed default to preprompt
                if core:
                    for tok in core:
                        tl = tok.lower()
                        if tl in ("preprompt", "summarize", "full"):
                            mode = tl
                        elif tl == "all":
                            select_all = True
                        elif path_token is None:
                            path_token = tok
                        else:
                            # extra tokens ignored
                            pass
                if path_token is None:
                    path_token = "."
                path = os.path.abspath(path_token)
                try:
                    if path and os.path.isdir(path):
                        # Attach directory: gather supported files and summarize all
                        supported_ext = {
                            ".txt",
                            ".md",
                            ".pdf",
                            ".docx",
                            ".html",
                            ".htm",
                            ".png",
                            ".jpg",
                            ".jpeg",
                            ".tif",
                            ".tiff",
                            ".py",
                            ".json",
                            ".yaml",
                            ".yml",
                            ".toml",
                            ".ini",
                            ".cfg",
                            ".csv",
                            ".tsv",
                            ".sh",
                            ".ps1",
                            ".bat",
                            ".cmd",
                            ".ipynb",
                        }
                        files = []
                        scanned = 0
                        skipped_hidden = 0
                        skipped_pattern = 0
                        skipped_ext = 0
                        for root, _, fnames in os.walk(path):
                            for fn in fnames:
                                scanned += 1
                                rb = os.path.basename(os.path.abspath(root))
                                if not include_hidden and (
                                    fn.startswith(".") or rb.startswith(".")
                                ):
                                    skipped_hidden += 1
                                    continue
                                extc = os.path.splitext(fn)[1].lower()
                                if patterns is not None and extc not in patterns:
                                    skipped_pattern += 1
                                    continue
                                if extc in supported_ext:
                                    files.append(os.path.join(root, fn))
                                else:
                                    skipped_ext += 1
                        if not files:
                            print(formatter.format_status_message("No supported files found in directory.", "warning"))
                            if debug_scan:
                                print(
                                    f"Scanned: {scanned}, skipped_hidden: {skipped_hidden}, skipped_pattern: {skipped_pattern}, skipped_ext: {skipped_ext}"
                                )
                                # Show up to 10 sample filenames in the top-level dir for quick inspection
                                try:
                                    samples = os.listdir(path)
                                    print(formatter.format_status_message("Sample entries in directory:", "info"))
                                    for s in samples[:10]:
                                        print(" -", s)
                                except Exception:
                                    pass
                            continue
                        limit = (
                            max_files
                            if isinstance(max_files, int) and max_files > 0
                            else 50
                        )
                        shown = files[:limit]
                        selected_files = None
                        # Interactive multi-select if prompt_toolkit is available
                        if select_all:
                            targets = shown
                        else:
                            if PROMPT_TOOLKIT and len(shown) > 1:
                                print(
                                    f"Found {len(files)} files. Top choice is [ALL]. Use Up/Down to navigate, Enter to toggle, Space to confirm. Press Esc to cancel. Press Space with none selected (or [ALL] selected) to attach all shown."
                                )
                                try:
                                    from prompt_toolkit.widgets import (
                                        CheckboxList,
                                        Frame,
                                    )
                                    from prompt_toolkit.layout.containers import HSplit
                                    from prompt_toolkit.layout import Layout
                                    from prompt_toolkit.application import Application
                                    from prompt_toolkit.key_binding import KeyBindings

                                    # Build a checkbox list with relative paths for display
                                    # Include a top-level "ALL" option so users can select all at once.
                                    rel_items = [("__ALL__", "[ALL] (select all)")]
                                    for fp in shown:
                                        rel = os.path.relpath(fp, path)
                                        rel_items.append((fp, rel))
                                    checklist = CheckboxList(values=rel_items)
                                    container = HSplit(
                                        [
                                            Frame(
                                                checklist,
                                                title=f"Select files (showing first {len(shown)} of {len(files)})",
                                            )
                                        ]
                                    )
                                    kb = KeyBindings()
                                    # Swap keys: Enter toggles selection, Space confirms/apply.
                                    @kb.add("enter", eager=True)
                                    def _(event):  # type: ignore[func-annotations]
                                        try:
                                            idx = getattr(checklist, "_selected_index", None)
                                            if idx is None:
                                                return
                                            val = checklist.values[idx][0]
                                            cur = set(checklist.current_values)
                                            if val in cur:
                                                cur.remove(val)
                                            else:
                                                cur.add(val)
                                            checklist.current_values = list(cur)
                                            event.app.invalidate()
                                        except Exception:
                                            pass

                                    @kb.add("space", eager=True)
                                    def __(event):  # type: ignore[func-annotations]
                                        event.app.exit(result=True)

                                    @kb.add("escape", eager=True)
                                    def ___(event):  # type: ignore[func-annotations]
                                        event.app.exit(result=None)

                                    app = Application(
                                        layout=Layout(container),
                                        key_bindings=kb,
                                        full_screen=False,
                                    )
                                    ret = app.run()
                                    if ret is None:
                                        print(formatter.format_status_message("Selection canceled.", "warning"))
                                        continue
                                    # If user pressed Enter without toggling, current_values may be empty; treat as None -> all shown
                                    cur = list(checklist.current_values)
                                    if not cur or "__ALL__" in cur:
                                        selected_files = None  # means all shown
                                    else:
                                        # Filter out the special ALL token if present alongside explicit picks
                                        selected_files = [
                                            v for v in cur if v != "__ALL__"
                                        ]
                                except Exception:
                                    selected_files = None
                            else:
                                # Fallback simple selection via input
                                print(
                                    f"Found {len(files)} files. Showing first {len(shown)}. Enter comma-separated indices (1-based) to select, or press Enter for all shown. Type 'c' to cancel."
                                )
                                try:
                                    for i, fp in enumerate(shown, 1):
                                        print(f"  {i:2d}) {os.path.relpath(fp, path)}")
                                    line = input(
                                        "Select files (e.g., 1,3,5) or Enter for all (or 'c' to cancel): "
                                    ).strip()
                                    if line.lower() in ("c", "q", "cancel", "quit"):
                                        print(formatter.format_status_message("Selection canceled.", "warning"))
                                        continue
                                    if line:
                                        picks = []
                                        for tok in line.split(","):
                                            tok = tok.strip()
                                            if tok.isdigit():
                                                idx = int(tok)
                                                if 1 <= idx <= len(shown):
                                                    picks.append(shown[idx - 1])
                                        if picks:
                                            selected_files = picks
                                except Exception:
                                    selected_files = None
                            targets = selected_files if selected_files else shown
                        if mode == "preprompt":
                            # New preprompt mode: store files for next user message
                            total_chars = 0
                            for fp in targets:
                                try:
                                    ext2 = os.path.splitext(fp)[1].lower()
                                    text2 = None
                                    if ext2 in (".txt", ".md"):
                                        with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                                            text2 = f.read()
                                    elif ext2 == ".pdf":
                                        from pypdf import PdfReader
                                        reader = PdfReader(fp)
                                        seg = []
                                        for pg in reader.pages:
                                            try:
                                                seg.append(pg.extract_text() or "")
                                            except Exception:
                                                pass
                                        text2 = "\n".join(seg)
                                    elif ext2 == ".docx":
                                        from docx import Document
                                        doc = Document(fp)
                                        seg = [p.text for p in doc.paragraphs]
                                        text2 = "\n".join(seg)
                                    elif ext2 in (".html", ".htm"):
                                        try:
                                            from bs4 import BeautifulSoup
                                            with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                                                soup = BeautifulSoup(f, "html.parser")
                                            for tag in soup(["script", "style"]):
                                                tag.extract()
                                            text2 = " ".join(soup.stripped_strings)
                                        except Exception:
                                            import re
                                            with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                                                html = f.read()
                                            t = re.sub(r"<script[\s\S]*?</script>", " ", html, flags=re.IGNORECASE)
                                            t = re.sub(r"<style[\s\S]*?</style>", " ", t, flags=re.IGNORECASE)
                                            t = re.sub(r"<[^>]+>", " ", t)
                                            t = re.sub(r"\s+", " ", t)
                                            text2 = t.strip()
                                    elif ext2 in (".py", ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".csv", ".tsv", ".sh", ".ps1", ".bat", ".cmd"):
                                        with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                                            text2 = f.read()
                                    elif ext2 == ".ipynb":
                                        try:
                                            import json as _json
                                            with open(fp, "r", encoding="utf-8", errors="ignore") as f:
                                                nb = _json.load(f)
                                            cells = nb.get("cells", [])
                                            parts = []
                                            for c in cells:
                                                if c.get("cell_type") in ("markdown", "code"):
                                                    src = c.get("source", [])
                                                    if isinstance(src, list):
                                                        parts.append("".join(src))
                                                    else:
                                                        parts.append(str(src))
                                            text2 = "\n".join(parts)
                                        except Exception:
                                            pass
                                    elif ext2 in (".png", ".jpg", ".jpeg", ".tif", ".tiff"):
                                        from PIL import Image
                                        import pytesseract
                                        text2 = pytesseract.image_to_string(Image.open(fp))
                                    
                                    if text2:
                                        rel = os.path.relpath(fp, path)
                                        attachment_state.add_attachment(rel, text2)
                                        total_chars += len(text2)
                                except Exception as e:
                                    print(formatter.format_status_message(f"Skip {fp}: {e}", "warning"))
                            
                            if attachment_state.has_attachments():
                                print(formatter.format_status_message(f"Attached {len(targets)} file(s) for next message ({total_chars:,} chars total)", "success"))
                                if attachment_state.get_total_size() > attachment_state.max_total_size:
                                    print(formatter.format_status_message(f"Warning: Total size ({attachment_state.get_total_size():,} chars) exceeds limit ({attachment_state.max_total_size:,}). Content will be truncated.", "warning"))
                            else:
                                print(formatter.format_status_message("No text extracted from directory.", "warning"))
                            continue
                        
                        # Original summarization mode
                        print(formatter.format_status_message(f"Extracting and summarizing {len(targets)} file(s)...", "info"))
                        all_chunks = []
                        for fp in targets:  # safety cap
                            try:
                                ext2 = os.path.splitext(fp)[1].lower()
                                text2 = None
                                if ext2 in (".txt", ".md"):
                                    with open(
                                        fp, "r", encoding="utf-8", errors="ignore"
                                    ) as f:
                                        text2 = f.read()
                                elif ext2 == ".pdf":
                                    from pypdf import PdfReader

                                    reader = PdfReader(fp)
                                    seg = []
                                    for pg in reader.pages:
                                        try:
                                            seg.append(pg.extract_text() or "")
                                        except Exception:
                                            pass
                                    text2 = "\n".join(seg)
                                elif ext2 == ".docx":
                                    from docx import Document

                                    doc = Document(fp)
                                    seg = [p.text for p in doc.paragraphs]
                                    text2 = "\n".join(seg)
                                elif ext2 in (".html", ".htm"):
                                    try:
                                        from bs4 import BeautifulSoup

                                        with open(
                                            fp, "r", encoding="utf-8", errors="ignore"
                                        ) as f:
                                            soup = BeautifulSoup(f, "html.parser")
                                        for tag in soup(["script", "style"]):
                                            tag.extract()
                                        text2 = " ".join(soup.stripped_strings)
                                    except Exception:
                                        import re

                                        with open(
                                            fp, "r", encoding="utf-8", errors="ignore"
                                        ) as f:
                                            html = f.read()
                                        t = re.sub(
                                            r"<script[\s\S]*?</script>",
                                            " ",
                                            html,
                                            flags=re.IGNORECASE,
                                        )
                                        t = re.sub(
                                            r"<style[\s\S]*?</style>",
                                            " ",
                                            t,
                                            flags=re.IGNORECASE,
                                        )
                                        t = re.sub(r"<[^>]+>", " ", t)
                                        t = re.sub(r"\s+", " ", t)
                                        text2 = t.strip()
                                elif ext2 in (
                                    ".py",
                                    ".json",
                                    ".yaml",
                                    ".yml",
                                    ".toml",
                                    ".ini",
                                    ".cfg",
                                    ".csv",
                                    ".tsv",
                                    ".sh",
                                    ".ps1",
                                    ".bat",
                                    ".cmd",
                                ):
                                    with open(
                                        fp, "r", encoding="utf-8", errors="ignore"
                                    ) as f:
                                        text2 = f.read()
                                elif ext2 == ".ipynb":
                                    try:
                                        import json as _json

                                        with open(
                                            fp, "r", encoding="utf-8", errors="ignore"
                                        ) as f:
                                            nb = _json.load(f)
                                        cells = nb.get("cells", [])
                                        parts = []
                                        for c in cells:
                                            if c.get("cell_type") in (
                                                "markdown",
                                                "code",
                                            ):
                                                src = c.get("source", [])
                                                if isinstance(src, list):
                                                    parts.append("".join(src))
                                                else:
                                                    parts.append(str(src))
                                        text2 = "\n".join(parts)
                                    except Exception:
                                        pass
                                elif ext2 in (".png", ".jpg", ".jpeg", ".tif", ".tiff"):
                                    from PIL import Image
                                    import pytesseract

                                    text2 = pytesseract.image_to_string(Image.open(fp))
                                if text2:
                                    rel = os.path.relpath(fp, path)
                                    all_chunks.append(f"### {rel}\n{text2}")
                            except Exception as e:
                                print(formatter.format_status_message(f"Skip {fp}: {e}", "warning"))
                        if not all_chunks:
                            print(formatter.format_status_message("No text extracted from directory.", "warning"))
                            continue
                        combined = "\n\n".join(all_chunks)
                        print(formatter.format_status_message("Summarizing directory content...", "info"))
                        summarize_messages = list(messages)
                        summarize_messages.append(
                            {
                                "role": "user",
                                "content": f"Summarize the following set of files into a single concise brief (<=400 words). Use bullet points and include file names where relevant. Reply only with the summary.\n\n{combined[:200000]}",
                            }
                        )
                        payload = make_payload_from_messages(summarize_messages, args)
                        payload.pop("stream", None)
                        try:
                            r = requests.post(
                                config.api_url, headers=headers, json=payload
                            )
                            r.raise_for_status()
                            d = r.json()
                            summary = extract_content(d) or ""
                            if not summary:
                                print(formatter.format_status_message("Failed to produce summary.", "error"))
                            else:
                                messages.append(
                                    {
                                        "role": "system",
                                        "content": f"Attachment summary (directory {os.path.basename(path)}): {summary}",
                                    }
                                )
                                print(formatter.format_status_message(
                                    "Directory summarized and added as system context.", "success"
                                ))
                        except Exception as e:
                            print(formatter.format_status_message(f"Summarize failed: {e}", "error"))
                        continue
                    ext = os.path.splitext(path)[1].lower()
                    content_text = None
                    content_parts: List[Dict[str, Any]] = []
                    ocr_text: Optional[str] = None
                    if ext in (".txt", ".md"):
                        with open(path, "r", encoding="utf-8", errors="ignore") as f:
                            content_text = f.read()
                    elif ext == ".pdf":
                        try:
                            from pypdf import PdfReader

                            reader = PdfReader(path)
                            txt = []
                            for page in reader.pages:
                                try:
                                    txt.append(page.extract_text() or "")
                                except Exception:
                                    pass
                            content_text = "\n".join(txt)
                        except Exception as e:
                            print(formatter.format_status_message(f"PDF read failed: {e}", "error"))
                            content_text = None
                    elif ext == ".docx":
                        try:
                            from docx import Document

                            doc = Document(path)
                            paras = []
                            for p in doc.paragraphs:
                                paras.append(p.text)
                            content_text = "\n".join(paras)
                        except Exception as e:
                            print(formatter.format_status_message(f"DOCX read failed: {e}", "error"))
                            content_text = None
                    elif ext in (".html", ".htm"):
                        try:
                            try:
                                from bs4 import BeautifulSoup  # optional

                                with open(
                                    path, "r", encoding="utf-8", errors="ignore"
                                ) as f:
                                    soup = BeautifulSoup(f, "html.parser")
                                # remove script/style
                                for tag in soup(["script", "style"]):
                                    tag.extract()
                                content_text = " ".join(soup.stripped_strings)
                            except Exception:
                                import re

                                with open(
                                    path, "r", encoding="utf-8", errors="ignore"
                                ) as f:
                                    html = f.read()
                                text = re.sub(
                                    r"<script[\s\S]*?</script>",
                                    " ",
                                    html,
                                    flags=re.IGNORECASE,
                                )
                                text = re.sub(
                                    r"<style[\s\S]*?</style>",
                                    " ",
                                    text,
                                    flags=re.IGNORECASE,
                                )
                                text = re.sub(r"<[^>]+>", " ", text)
                                text = re.sub(r"\s+", " ", text)
                                content_text = text.strip()
                        except Exception as e:
                            print(formatter.format_status_message(f"HTML read failed: {e}", "error"))
                            content_text = None
                    elif ext in (".png", ".jpg", ".jpeg", ".tif", ".tiff"):
                        try:
                            # 1) Embed the image as base64 for models that support multimodal input
                            with open(path, "rb") as f:
                                image_data = f.read()
                            base64_image = base64.b64encode(image_data).decode("utf-8")
                            mime_type = (
                                f"image/{ext[1:]}" if ext != ".jpg" else "image/jpeg"
                            )
                            content_object = {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                },
                            }
                            content_parts = [content_object]
                            # 2) Also OCR the image to produce text for text-only models or summarization
                            try:
                                from PIL import Image
                                import pytesseract

                                ocr_text = (
                                    pytesseract.image_to_string(Image.open(path)) or ""
                                )
                            except Exception as ocr_e:
                                print(formatter.format_status_message(f"OCR failed (non-fatal): {ocr_e}", "error"))
                                ocr_text = ""
                        except Exception as e:
                            print(formatter.format_status_message(f"Image encoding failed: {e}", "error"))
                            content_parts = []
                    else:
                        print(
                            "Unsupported file type. Supported: .txt .md .pdf .docx .png .jpg .jpeg .tif .tiff .html .htm"
                        )
                        continue
                    if (
                        not content_parts
                        and not content_text
                        and not (ocr_text and ocr_text.strip())
                    ):
                        print(formatter.format_status_message("No text extracted from file.", "warning"))
                        continue
                    # Decide how to inject: preprompt by default, summarize if requested, or direct if --as-user or full mode
                    if mode == "preprompt":
                        # New preprompt mode: store file for next user message
                        if content_text:
                            attachment_state.add_attachment(os.path.basename(path), content_text)
                            print(formatter.format_status_message(f"Attached {os.path.basename(path)} for next message ({len(content_text):,} chars)", "success"))
                            if attachment_state.get_total_size() > attachment_state.max_total_size:
                                print(formatter.format_status_message(f"Warning: Total size ({attachment_state.get_total_size():,} chars) exceeds limit ({attachment_state.max_total_size:,}). Content will be truncated.", "warning"))
                        elif content_parts:
                            # For images, store OCR text if available
                            if ocr_text and ocr_text.strip():
                                attachment_state.add_attachment(os.path.basename(path), ocr_text)
                                print(formatter.format_status_message(f"Attached OCR text from {os.path.basename(path)} for next message ({len(ocr_text):,} chars)", "success"))
                            else:
                                print(formatter.format_status_message(f"Image {os.path.basename(path)} attached but no OCR text available for preprompt mode.", "warning"))
                        continue
                    elif as_user or (
                        mode == "full"
                        and (
                            (content_text and len(content_text) < 8000) or content_parts
                        )
                    ):
                        final_content_for_message = []
                        if content_parts:  # Image attachment
                            final_content_for_message.extend(content_parts)
                            # Include OCR text alongside the image, if any
                            if ocr_text and ocr_text.strip():
                                max_chars_val = getattr(
                                    args, "attach_truncation_limit", 8000
                                )
                                max_chars = (
                                    max_chars_val
                                    if isinstance(max_chars_val, int)
                                    and max_chars_val > 0
                                    else 8000
                                )
                                clipped = ocr_text[:max_chars] + (
                                    "\n... [truncated]"
                                    if len(ocr_text) > max_chars
                                    else ""
                                )
                                final_content_for_message.append(
                                    {
                                        "type": "text",
                                        "text": f"[OCR from {os.path.basename(path)}]\n\n{clipped}",
                                    }
                                )
                            else:
                                # Fallback descriptive text if OCR is empty
                                final_content_for_message.append(
                                    {
                                        "type": "text",
                                        "text": f"[Attached image from {os.path.basename(path)}]",
                                    }
                                )
                            print(formatter.format_status_message(f"Attached image from {path} as a user message.", "success"))
                        elif content_text:  # Text-based attachment
                            max_chars_val = getattr(
                                args, "attach_truncation_limit", 8000
                            )
                            max_chars = (
                                max_chars_val
                                if isinstance(max_chars_val, int) and max_chars_val > 0
                                else 8000
                            )
                            clipped = content_text[:max_chars] + (
                                "\n... [truncated]"
                                if len(content_text) > max_chars
                                else ""
                            )
                            final_content_for_message.append(
                                {
                                    "type": "text",
                                    "text": f"[Attached content from {os.path.basename(path)}]\n\n{clipped}",
                                }
                            )
                            attach_status = formatter.format_status_message(
                                f"Attached content from {path} as a user message{' (truncated)' if len(content_text) > max_chars else ''}.", 
                                "success"
                            )
                            print(attach_status)

                        if final_content_for_message:
                            messages.append(
                                {"role": "user", "content": final_content_for_message}
                            )
                        continue
                    else:  # Summarize content via model and attach summary as system note
                        summarize_header = formatter.format_section_header("Processing Attachment")
                        print(summarize_header)
                        print("Summarizing attachment...")
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        summarize_messages = list(messages)
                        # For summarization, prefer OCR text for images; fall back to placeholder if OCR empty
                        content_to_summarize = ""
                        if content_parts:
                            if ocr_text and ocr_text.strip():
                                content_to_summarize = ocr_text
                            else:
                                content_to_summarize = (
                                    f"[Image attached: {os.path.basename(path)}]"
                                )
                        elif content_text:
                            content_to_summarize = content_text

                        summarize_messages.append(
                            {
                                "role": "user",
                                "content": f"Summarize the following content into <=200 words focusing on the most relevant facts and context. Reply only with the summary.\n\nCONTENT BEGIN\n{content_to_summarize}\nCONTENT END",
                            }
                        )
                        payload = make_payload_from_messages(summarize_messages, args)
                        payload.pop("stream", None)
                        try:
                            r = requests.post(
                                config.api_url, headers=headers, json=payload
                            )
                            r.raise_for_status()
                            d = r.json()
                            summary = extract_content(d) or ""
                            if not summary:
                                print(
                                    "Failed to produce summary; attaching truncated text instead."
                                )
                                messages.append(
                                    {
                                        "role": "user",
                                        "content": f"[Attached content from {os.path.basename(path)}]\n\n{content_to_summarize[:4000]}",
                                    }
                                )
                            else:
                                messages.append(
                                    {
                                        "role": "system",
                                        "content": f"Attachment summary ({os.path.basename(path)}): {summary}",
                                    }
                                )
                                print(
                                    "Attachment summarized and added as system context."
                                )
                        except Exception as e:
                            print(formatter.format_status_message(f"Summarize failed: {e}", "error"))
                            messages.append(
                                {
                                    "role": "user",
                                    "content": f"[Attached content from {os.path.basename(path)}]\n\n{content_to_summarize[:4000]}",
                                }
                            )
                            _session_write({
                                "role": "user",
                                "content": f"[Attached content from {os.path.basename(path)}]\n\n{content_to_summarize[:4000]}",
                            })
                    continue
                except FileNotFoundError:
                    print(formatter.format_status_message("File not found.", "error"))
                    continue
                except Exception as e:
                    print(formatter.format_status_message(f"Attach failed: {e}", "error"))
                    continue
            if cmd == "copy":

                try:
                    import pyperclip  # type: ignore[import-not-found]
                except Exception:
                    print(formatter.format_status_message("pyperclip not installed. Run: pip install pyperclip", "warning"))
                    continue
                # find last assistant message
                last = None
                for m in reversed(messages):
                    if m.get("role") == "assistant":
                        last = m.get("content")
                        break
                if not last:
                    print(formatter.format_status_message("No assistant response to copy.", "info"))
                else:
                    try:
                        pyperclip.copy(last)
                        print(formatter.format_status_message("Copied last response to clipboard.", "success"))
                    except Exception as e:
                        print(formatter.format_status_message(f"Copy failed: {e}", "error"))
                continue
            if cmd == "async":
                sub = args_rest.split()
                if not sub or sub[0] in ("help", "?"):
                    print(
                        "/async submit [prompt] | /async list | /async get <id> | /async wait <id>"
                    )
                elif sub[0] == "submit":
                    # Allow optional inline prompt: /async submit <prompt>
                    inline_prompt = args_rest[len("submit"):].strip() if len(args_rest) > len("submit") else ""
                    msgs_for_job: List[Dict[str, Any]] = list(messages) if messages else []
                    if inline_prompt:
                        msgs_for_job = list(msgs_for_job)  # copy to avoid mutating chat history
                        msgs_for_job.append({"role": "user", "content": inline_prompt})
                    if not msgs_for_job:
                        print(formatter.format_status_message("Nothing to submit. Type a prompt first or use '/async submit <prompt>'.", "warning"))
                    else:
                        rid = async_submit(msgs_for_job)
                        if rid:
                            last_async_id = rid
                            submit_status = formatter.format_status_message(f"Submitted async request: {rid}", "success")
                            print(submit_status)
                elif sub[0] == "list":
                    data = async_list()
                    list_header = formatter.format_section_header("Async Requests")
                    print(list_header)
                    print(f"{formatter.theme.SECTION_SEPARATOR}")
                    print(formatter.format_json_output(data))
                    print(f"{formatter.theme.SECTION_SEPARATOR}")
                elif sub[0] == "get":
                    def _pick_async_id() -> Optional[str]:
                        listing = async_list() or {}
                        items = []
                        # Extract items robustly
                        if isinstance(listing, dict):
                            cand = listing.get("data") or listing.get("items") or listing.get("requests")
                            if isinstance(cand, list):
                                items = cand
                        elif isinstance(listing, list):
                            items = listing
                        # Build choices (id, label)
                        choices = []
                        for it in items:
                            rid = it.get("id") or it.get("request_id") or it.get("requestId")
                            status = (it.get("status") or it.get("state") or "").upper()
                            # Try to show a summary from the first user message
                            summary = ""
                            try:
                                req = it.get("request") or {}
                                msgs = req.get("messages") or []
                                if msgs:
                                    m0 = msgs[-1]
                                    c = m0.get("content")
                                    if isinstance(c, list) and c:
                                        for part in c:
                                            if isinstance(part, dict) and part.get("type") == "text":
                                                summary = part.get("text") or summary
                                                break
                                    elif isinstance(c, str):
                                        summary = c
                            except Exception:
                                pass
                            label = f"{rid}  [{status}]  {summary[:60]}" if rid else None
                            if rid and label:
                                choices.append((rid, label))
                        if not choices:
                            print(formatter.format_status_message("No async requests to select.", "info"))
                            return None
                        if PROMPT_TOOLKIT:
                            try:
                                from prompt_toolkit.widgets import RadioList, Frame
                                from prompt_toolkit.layout.containers import HSplit
                                from prompt_toolkit.layout import Layout
                                from prompt_toolkit.application import Application
                                from prompt_toolkit.key_binding import KeyBindings
                                radio = RadioList(values=choices)
                                container = HSplit([Frame(radio, title="Select an async request")])
                                kb = KeyBindings()
                                @kb.add("enter", eager=True)
                                def _(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=radio.current_value)
                                @kb.add("escape", eager=True)
                                def _(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=None)
                                app = Application(layout=Layout(container), key_bindings=kb, full_screen=False)
                                ret = app.run()
                                return ret
                            except Exception:
                                pass
                        # Fallback simple numeric selection
                        print(formatter.format_status_message("Select an async request:", "info"))
                        for i, (_, label) in enumerate(choices, 1):
                            print(f"  {i:2d}) {label}")
                        ans = input("Enter number (or 'c' to cancel): ").strip()
                        if ans.lower() in ("c", "q", "cancel", "quit"):
                            return None
                        if ans.isdigit():
                            idx = int(ans)
                            if 1 <= idx <= len(choices):
                                return choices[idx - 1][0]
                        print(formatter.format_status_message("Invalid selection.", "error"))
                        return None
                    target_id = sub[1] if len(sub) >= 2 else (last_async_id or _pick_async_id())
                    if not target_id:
                        print(formatter.format_status_message("No async id selected.", "warning"))
                    else:
                        data = async_get(target_id)
                        if target_id:
                            last_async_id = target_id
                        get_header = formatter.format_section_header(f"Async Request {target_id}")
                        print(get_header)
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        # Prefer text if available
                        if isinstance(data, dict):
                            content_txt = extract_content(data)
                        else:
                            content_txt = None
                        if content_txt:
                            formatted_response = formatter.format_ai_response(content_txt)
                            print(formatted_response)
                            # Prefer citations/usage from async 'response' wrapper if present
                            root = data.get("response") if isinstance(data, dict) and isinstance(data.get("response"), dict) else data
                            show_meta = getattr(args, "async_meta", False)
                            citations_val = root.get("citations") if isinstance(root, dict) else None
                            usage_val = root.get("usage") if isinstance(root, dict) else None
                            if (show_meta or args.citations) and isinstance(citations_val, list):
                                print_citations(formatter, citations_val)
                            if (show_meta or args.usage) and isinstance(usage_val, dict):
                                print_usage(formatter, usage_val)
                        else:
                            print(formatter.format_json_output(data))
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                elif sub[0] == "wait":
                    def _pick_async_id() -> Optional[str]:
                        # Reuse get selector
                        listing = async_list() or {}
                        items = []
                        if isinstance(listing, dict):
                            cand = listing.get("data") or listing.get("items") or listing.get("requests")
                            if isinstance(cand, list):
                                items = cand
                        elif isinstance(listing, list):
                            items = listing
                        choices = []
                        for it in items:
                            rid = it.get("id") or it.get("request_id") or it.get("requestId")
                            status = (it.get("status") or it.get("state") or "").upper()
                            summary = ""
                            try:
                                req = it.get("request") or {}
                                msgs = req.get("messages") or []
                                if msgs:
                                    m0 = msgs[-1]
                                    c = m0.get("content")
                                    if isinstance(c, list) and c:
                                        for part in c:
                                            if isinstance(part, dict) and part.get("type") == "text":
                                                summary = part.get("text") or summary
                                                break
                                    elif isinstance(c, str):
                                        summary = c
                            except Exception:
                                pass
                            label = f"{rid}  [{status}]  {summary[:60]}" if rid else None
                            if rid and label:
                                choices.append((rid, label))
                        if not choices:
                            print(formatter.format_status_message("No async requests to select.", "info"))
                            return None
                        if PROMPT_TOOLKIT:
                            try:
                                from prompt_toolkit.widgets import RadioList, Frame
                                from prompt_toolkit.layout.containers import HSplit
                                from prompt_toolkit.layout import Layout
                                from prompt_toolkit.application import Application
                                from prompt_toolkit.key_binding import KeyBindings
                                radio = RadioList(values=choices)
                                container = HSplit([Frame(radio, title="Select an async request to wait for")])
                                kb = KeyBindings()
                                @kb.add("enter", eager=True)
                                def _(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=radio.current_value)
                                @kb.add("escape", eager=True)
                                def _(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=None)
                                app = Application(layout=Layout(container), key_bindings=kb, full_screen=False)
                                ret = app.run()
                                return ret
                            except Exception:
                                pass
                        print(formatter.format_status_message("Select an async request:", "info"))
                        for i, (_, label) in enumerate(choices, 1):
                            print(f"  {i:2d}) {label}")
                        ans = input("Enter number (or 'c' to cancel): ").strip()
                        if ans.lower() in ("c", "q", "cancel", "quit"):
                            return None
                        if ans.isdigit():
                            idx = int(ans)
                            if 1 <= idx <= len(choices):
                                return choices[idx - 1][0]
                        print(formatter.format_status_message("Invalid selection.", "error"))
                        return None
                    target_id = sub[1] if len(sub) >= 2 else (last_async_id or _pick_async_id())
                    if not target_id:
                        print(formatter.format_status_message("No async id selected.", "warning"))
                    else:
                        data = async_wait(target_id)
                        if target_id:
                            last_async_id = target_id
                        wait_header = formatter.format_section_header(f"Async Result {target_id}")
                        print(wait_header)
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        if isinstance(data, dict):
                            content_txt = extract_content(data)
                        else:
                            content_txt = None
                        if content_txt:
                            formatted_response = formatter.format_ai_response(content_txt)
                            print(formatted_response)
                            root = data.get("response") if isinstance(data, dict) and isinstance(data.get("response"), dict) else data
                            show_meta = getattr(args, "async_meta", False)
                            citations_val = root.get("citations") if isinstance(root, dict) else None
                            usage_val = root.get("usage") if isinstance(root, dict) else None
                            if (show_meta or args.citations) and isinstance(citations_val, list):
                                print_citations(formatter, citations_val)
                            if (show_meta or args.usage) and isinstance(usage_val, dict):
                                print_usage(formatter, usage_val)
                        else:
                            print(formatter.format_json_output(data))
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                else:
                    print(
                        "Usage: /async submit | /async list | /async get <id> | /async wait <id>"
                    )
                continue
            if cmd == "save":
                path = args_rest if args_rest else "conversation.jsonl"
                try:
                    with open(path, "w", encoding="utf-8") as f:
                        for m in messages:
                            f.write(json.dumps(m, ensure_ascii=False) + "\n")
                    print(formatter.format_status_message(f"Saved to {path}", "success"))
                except Exception as e:
                    print(formatter.format_status_message(f"Failed to save: {e}", "error"))
                continue
            if cmd == "academic":
                val = args_rest.lower()
                if val in ("on", "off"):
                    args.academic = val == "on"
                elif val == "":
                    args.academic = not bool(getattr(args, "academic", False))
                else:
                    print(formatter.format_status_message("Usage: /academic on|off", "info"))
                    continue
                print(formatter.format_status_message(f"academic = {args.academic}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "domain":
                host = args_rest.strip()
                if host:
                    args.domain = host
                    print(formatter.format_status_message(f"domain = {args.domain}", "info"))
                else:
                    args.domain = None
                    print(formatter.format_status_message("Domain cleared.", "success"))
                save_prefs_from_args(args)
                continue
            if cmd == "recency":
                v = args_rest.strip().lower()
                if v in ("day", "week", "month"):
                    args.recency = v
                    print(formatter.format_status_message(f"recency = {args.recency}", "info"))
                elif v in ("off", "none", ""):
                    args.recency = None
                    print(formatter.format_status_message("Recency cleared.", "success"))
                else:
                    print(formatter.format_status_message("Usage: /recency day|week|month|off", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "attachlimit":
                try:
                    n = int(args_rest.strip())
                    if n <= 1000:
                        print(formatter.format_status_message("Attachlimit must be > 1000 characters to be useful.", "warning"))
                        continue
                    args.attach_truncation_limit = n
                    print(formatter.format_status_message(f"attachlimit set to {n}", "info"))
                except Exception:
                    print(formatter.format_status_message("Usage: /attachlimit <N>", "info"))
                continue
            if cmd == "attachments":
                sub = args_rest.strip().lower()
                if sub == "clear":
                    if attachment_state.has_attachments():
                        attachment_state.clear()
                        print(formatter.format_status_message("Cleared all pending attachments.", "success"))
                    else:
                        print(formatter.format_status_message("No pending attachments to clear.", "info"))
                else:
                    if attachment_state.has_attachments():
                        print(formatter.format_status_message(f"Pending attachments: {attachment_state.get_summary()}", "info"))
                        print(formatter.format_status_message("These will be included with your next message. Use '/attachments clear' to remove them.", "info"))
                    else:
                        print(formatter.format_status_message("No pending attachments.", "info"))
                continue
            if cmd == "jsonschema":

                val = args_rest.lower()
                if val in ("on", "off"):
                    args.json_schema = val == "on"
                elif val == "":
                    args.json_schema = not bool(getattr(args, "json_schema", False))
                else:
                    print(formatter.format_status_message("Usage: /jsonschema on|off", "info"))
                    continue
                print(formatter.format_status_message(f"jsonschema = {args.json_schema}", "info"))
                # Do not persist jsonschema; it's session-local to avoid surprising futures
                continue
            if cmd == "citations":
                val = args_rest.lower()
                if val in ("on", "off"):
                    args.citations = val == "on"
                elif val == "":
                    args.citations = not bool(getattr(args, "citations", False))
                else:
                    print(formatter.format_status_message("Usage: /citations on|off", "info"))
                    continue
                print(formatter.format_status_message(f"citations = {args.citations}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "usage":
                val = args_rest.lower()
                if val in ("on", "off"):
                    args.usage = val == "on"
                elif val == "":
                    args.usage = not bool(getattr(args, "usage", False))
                else:
                    print(formatter.format_status_message("Usage: /usage on|off", "info"))
                    continue
                print(formatter.format_status_message(f"usage = {args.usage}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "asyncmeta":
                val = args_rest.lower()
                if val in ("on", "off"):
                    setattr(args, "async_meta", val == "on")
                elif val == "":
                    setattr(args, "async_meta", not bool(getattr(args, "async_meta", False)))
                else:
                    print(formatter.format_status_message("Usage: /asyncmeta on|off", "info"))
                    continue
                print(formatter.format_status_message(f"async_meta = {getattr(args, 'async_meta', False)}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "stream":
                val = args_rest.lower()
                if val in ("on", "off"):
                    args.stream = val == "on"
                elif val == "":
                    args.stream = not bool(getattr(args, "stream", False))
                else:
                    print(formatter.format_status_message("Usage: /stream on|off", "info"))
                    continue
                print(formatter.format_status_message(f"stream = {args.stream}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "mdrender":
                val = args_rest.lower()
                if val in ("on", "off"):
                    setattr(args, "md_render", val == "on")
                elif val == "":
                    setattr(args, "md_render", not bool(getattr(args, "md_render", True)))
                else:
                    print(formatter.format_status_message("Usage: /mdrender on|off", "info"))
                    continue
                # Apply immediately to current formatter
                formatter.md_render = bool(getattr(args, "md_render", True))
                print(formatter.format_status_message(f"md_render = {getattr(args, 'md_render', True)}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "text":
                val = args_rest.lower()
                if val in ("on", "off"):
                    setattr(args, "text", val == "on")
                elif val == "":
                    setattr(args, "text", not bool(getattr(args, "text", False)))
                else:
                    print(formatter.format_status_message("Usage: /text on|off", "info"))
                    continue
                print(formatter.format_status_message(f"text (one-shot default) = {getattr(args, 'text', False)}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "theme":
                sub = args_rest.strip().lower()
                if sub == "reset":
                    setattr(args, "theme", {})
                    formatter.theme = FormattingTheme({})
                    print(formatter.format_status_message("Theme reset to defaults.", "success"))
                    save_prefs_from_args(args)
                    continue
                # Interactive theme chooser with colorblind-friendly recommendations
                current = getattr(args, "theme", {}) or {}
                fields = [
                    ("user", "User label"),
                    ("ai", "AI label"),
                    ("citations", "Citations"),
                    ("metadata", "Metadata/keys"),
                    ("error", "Errors"),
                    ("success", "Success"),
                    ("warning", "Warnings"),
                    ("section", "Section headers"),
                ]
                # Recommended first (colorblind-friendly emphasis): cyan, yellow, blue, white, magenta, bright variants
                recommended = [
                    "bright_cyan", "bright_yellow", "bright_blue", "bright_white", "bright_magenta",
                    "cyan", "yellow", "blue", "white", "magenta",
                ]
                others = [
                    "bright_black", "bright_green", "bright_red",
                    "black", "green", "red",
                ]
                palette_names = []
                for n in recommended + others:
                    if n not in palette_names:
                        palette_names.append(n)
                def cli_select(options: List[str], prompt_text: str, default_idx: int = 0) -> Optional[str]:
                    print(formatter.format_status_message(prompt_text, "info"))
                    for i, opt in enumerate(options, 1):
                        print(f"  {i:2d}) {opt}")
                    print("  c) cancel")
                    ans = input("Select: ").strip().lower()
                    if ans in ("c", "q", "cancel", "quit"): return None
                    if ans.isdigit():
                        idx = int(ans)
                        if 1 <= idx <= len(options):
                            return options[idx - 1]
                    return options[default_idx] if options else None
                def run_cli_editor() -> Optional[Dict[str, str]]:
                    theme_cfg = dict(current)
                    while True:
                        print(formatter.format_status_message("Select an item to edit:", "info"))
                        for i, (k, label) in enumerate(fields, 1):
                            curv = theme_cfg.get(k, current.get(k, "")) or "(default)"
                            print(f"  {i:2d}) {label}: {curv}")
                        print("  s) save and exit")
                        print("  r) reset to defaults")
                        print("  c) cancel")
                        sel = input("Choice: ").strip().lower()
                        if sel == "s":
                            return theme_cfg
                        if sel == "r":
                            return {}
                        if sel in ("c", "q", "cancel", "quit"):
                            return None
                        if sel.isdigit():
                            idx = int(sel)
                            if 1 <= idx <= len(fields):
                                key, label = fields[idx - 1]
                                default_name = theme_cfg.get(key, current.get(key, "cyan" if key in ("user",) else "green" if key == "ai" else "yellow" if key in ("citations", "warning") else "blue" if key == "metadata" else "magenta" if key == "section" else "white"))
                                choice = cli_select(palette_names + ["other..."], f"Color for {label} [{default_name}]", 0)
                                if choice is None:
                                    continue
                                if choice == "other...":
                                    raw = input("Type color name (e.g., bright_cyan): ").strip().lower()
                                    if raw:
                                        theme_cfg[key] = raw
                                else:
                                    theme_cfg[key] = choice
                    return None
                def run_pt_editor() -> Optional[Dict[str, str]]:
                    try:
                        from prompt_toolkit.widgets import RadioList, Frame
                        from prompt_toolkit.layout.containers import HSplit
                        from prompt_toolkit.layout import Layout
                        from prompt_toolkit.application import Application
                        from prompt_toolkit.key_binding import KeyBindings
                    except Exception:
                        return run_cli_editor()
                    theme_cfg = dict(current)
                    while True:
                        # Main menu
                        def color_preview(name: str) -> str:
                            # Rough preview by applying ANSI if available
                            try:
                                from colorama import Fore as _F, Style as _S
                                mapping = {
                                    'black': getattr(_F, 'BLACK', ''),
                                    'red': _F.RED,
                                    'green': _F.GREEN,
                                    'yellow': _F.YELLOW,
                                    'blue': getattr(_F, 'BLUE', ''),
                                    'magenta': getattr(_F, 'MAGENTA', ''),
                                    'cyan': _F.CYAN,
                                    'white': getattr(_F, 'WHITE', ''),
                                    'bright_black': getattr(_F, 'LIGHTBLACK_EX', ''),
                                    'bright_red': getattr(_F, 'LIGHTRED_EX', _F.RED),
                                    'bright_green': getattr(_F, 'LIGHTGREEN_EX', _F.GREEN),
                                    'bright_yellow': getattr(_F, 'LIGHTYELLOW_EX', _F.YELLOW),
                                    'bright_blue': getattr(_F, 'LIGHTBLUE_EX', getattr(_F, 'BLUE', '')),
                                    'bright_magenta': getattr(_F, 'LIGHTMAGENTA_EX', getattr(_F, 'MAGENTA', '')),
                                    'bright_cyan': getattr(_F, 'LIGHTCYAN_EX', _F.CYAN),
                                    'bright_white': getattr(_F, 'LIGHTWHITE_EX', getattr(_F, 'WHITE', '')),
                                }
                                col = mapping.get(name, '')
                                return f"{col}{name}{_S.RESET_ALL}" if col else name
                            except Exception:
                                return name
                        # Resolve effective color names for preview
                        defaults = {
                            'user': 'cyan', 'ai': 'green', 'citations': 'yellow', 'metadata': 'blue',
                            'error': 'red', 'success': 'green', 'warning': 'yellow', 'section': 'magenta'
                        }
                        def eff_color(key: str) -> str:
                            return theme_cfg.get(key, current.get(key, defaults.get(key, 'white')))
                        def color_label(key: str, label: str) -> Any:
                            name = eff_color(key)
                            txt = f"{label}: {name}"
                            if PT_ANSI and formatter.use_colors:
                                ansi_map = {
                                    'black': '\u001b[30m','red': '\u001b[31m','green': '\u001b[32m','yellow': '\u001b[33m','blue': '\u001b[34m','magenta': '\u001b[35m','cyan': '\u001b[36m','white': '\u001b[37m',
                                    'bright_black': '\u001b[90m','bright_red': '\u001b[91m','bright_green': '\u001b[92m','bright_yellow': '\u001b[93m','bright_blue': '\u001b[94m','bright_magenta': '\u001b[95m','bright_cyan': '\u001b[96m','bright_white': '\u001b[97m'
                                }
                                col = ansi_map.get(name, '')
                                if col:
                                    return PT_ANSI(f"{label}: {col}{name}\u001b[0m")
                            return txt
                        values = [(k, color_label(k, label)) for k, label in fields]
                        values += [("save", "Save and exit"), ("reset", "Reset to defaults"), ("cancel", "Cancel")]
                        radio = RadioList(values=values)
                        container = HSplit([Frame(radio, title="Theme editor (colorblind-friendly order)")])
                        kb = KeyBindings()
                        @kb.add("enter", eager=True)
                        def _(event):  # type: ignore[func-annotations]
                            event.app.exit(result=radio.current_value)
                        @kb.add("escape", eager=True)
                        def _(event):  # type: ignore[func-annotations]
                            print(formatter.format_status_message("Canceled.", "warning"))
                            event.app.exit(result="cancel")
                        app = Application(layout=Layout(container), key_bindings=kb, full_screen=False)
                        action = app.run()
                        if action in ("cancel", None):
                            return None
                        if action == "reset":
                            return {}
                        if action == "save":
                            return theme_cfg
                        # Choose color for a field
                        key = action
                        label = dict(fields)[key]
                        if PT_ANSI and formatter.use_colors:
                            color_values = [(n, PT_ANSI(color_preview(n))) for n in (palette_names + ["other..."])]
                        else:
                            color_values = [(n, color_preview(n)) for n in (palette_names + ["other..."])]
                        radio2 = RadioList(values=color_values)
                        container2 = HSplit([Frame(radio2, title=f"Color for {label}")])
                        kb2 = KeyBindings()
                        @kb2.add("enter", eager=True)
                        def __(event):  # type: ignore[func-annotations]
                            event.app.exit(result=radio2.current_value)
                        @kb2.add("escape", eager=True)
                        def ___(event):  # type: ignore[func-annotations]
                            print(formatter.format_status_message("Canceled.", "warning"))
                            event.app.exit(result=None)
                        app2 = Application(layout=Layout(container2), key_bindings=kb2, full_screen=False)
                        choice = app2.run()
                        if choice is None:
                            continue
                        if choice == "other...":
                            try:
                                raw = input("Type color name (e.g., bright_cyan): ").strip().lower()
                            except Exception:
                                raw = ""
                            if raw:
                                theme_cfg[key] = raw
                        else:
                            theme_cfg[key] = choice
                    return None
                new_cfg = run_pt_editor() if PROMPT_TOOLKIT else run_cli_editor()
                if new_cfg is None:
                    print(formatter.format_status_message("Theme change canceled.", "warning"))
                    continue
                setattr(args, "theme", new_cfg)
                formatter.theme = FormattingTheme(new_cfg)
                print(formatter.format_status_message("Theme updated.", "success"))
                save_prefs_from_args(args)
                continue
            if cmd == "colors":
                val = args_rest.lower()
                if val in ("on", "off"):
                    formatter.use_colors = val == "on"
                    print(formatter.format_status_message(f"Color output set to {val}.", "info"))
                    save_prefs_from_args(args)
                elif val == "":
                    # Toggle current state
                    formatter.use_colors = not formatter.use_colors
                    print(formatter.format_status_message(f"Color output toggled to {'on' if formatter.use_colors else 'off'}.", "info"))
                else:
                    print(formatter.format_status_message("Usage: /colors on|off", "info"))
                continue
            if cmd == "tableformat":
                sub = args_rest.strip().lower()
                table_format_instructions = (
                    "IMPORTANT: When providing tabular data, NEVER use markdown table syntax with pipes (|). "
                    "Instead, ALWAYS use a clean bulleted list format. "
                    "Example format:\n\n"
                    "**Global CO₂ Emissions 2019:**\n"
                    "• **China**: 14,093 million metric tons CO₂e (27% of global total)\n"
                    "• **United States**: 5,720 million metric tons CO₂e (11% of global total)\n"
                    "• **India**: 3,432 million metric tons CO₂e (6.6% of global total)\n"
                    "• **Russia**: 2,169 million metric tons CO₂e (4.2% of global total)\n\n"
                    "This bulleted list format is much more readable in terminal environments. "
                    "Never use pipe characters (|), dashes for separators, or traditional table markdown syntax."
                )
                
                if sub == "show":
                    # Show current state
                    current_persistent = getattr(args, "persistent_system", None)
                    has_table_format = current_persistent and table_format_instructions.strip() in current_persistent
                    print(formatter.format_status_message(f"Table formatting instructions: {'ON' if has_table_format else 'OFF'}", "info"))
                    if has_table_format:
                        print(f"\nCurrent table formatting instruction:")
                        print(table_format_instructions)
                    continue
                
                elif sub == "on":
                    # Add table formatting to persistent system message
                    current_persistent = getattr(args, "persistent_system", None) or ""
                    if table_format_instructions.strip() not in current_persistent:
                        if current_persistent:
                            new_persistent = f"{current_persistent}\n\n{table_format_instructions}"
                        else:
                            new_persistent = table_format_instructions
                        setattr(args, "persistent_system", new_persistent)
                        save_prefs_from_args(args)
                        print(formatter.format_status_message("Table formatting instructions added to persistent system message.", "success"))
                    else:
                        print(formatter.format_status_message("Table formatting instructions already present.", "info"))
                    continue
                
                elif sub == "off":
                    # Remove table formatting from persistent system message
                    current_persistent = getattr(args, "persistent_system", None)
                    if current_persistent and table_format_instructions.strip() in current_persistent:
                        new_persistent = current_persistent.replace(table_format_instructions, "").strip()
                        # Clean up extra newlines manually instead of using re.sub
                        while '\n\n\n' in new_persistent:
                            new_persistent = new_persistent.replace('\n\n\n', '\n\n')
                        new_persistent = new_persistent.strip()
                        if new_persistent:
                            setattr(args, "persistent_system", new_persistent)
                        else:
                            setattr(args, "persistent_system", None)
                        save_prefs_from_args(args)
                        print(formatter.format_status_message("Table formatting instructions removed from persistent system message.", "success"))
                    else:
                        print(formatter.format_status_message("Table formatting instructions not found in persistent system message.", "warning"))
                    continue
                
                elif sub == "":
                    # Toggle table formatting
                    current_persistent = getattr(args, "persistent_system", None)
                    has_table_format = current_persistent and table_format_instructions.strip() in current_persistent
                    
                    if has_table_format:
                        # Remove it
                        new_persistent = current_persistent.replace(table_format_instructions, "").strip()
                        # Clean up extra newlines manually instead of using re.sub
                        while '\n\n\n' in new_persistent:
                            new_persistent = new_persistent.replace('\n\n\n', '\n\n')
                        new_persistent = new_persistent.strip()
                        if new_persistent:
                            setattr(args, "persistent_system", new_persistent)
                        else:
                            setattr(args, "persistent_system", None)
                        save_prefs_from_args(args)
                        print(formatter.format_status_message("Table formatting instructions removed.", "success"))
                    else:
                        # Add it
                        if current_persistent:
                            new_persistent = f"{current_persistent}\n\n{table_format_instructions}"
                        else:
                            new_persistent = table_format_instructions
                        setattr(args, "persistent_system", new_persistent)
                        save_prefs_from_args(args)
                        print(formatter.format_status_message("Table formatting instructions added.", "success"))
                    continue
                
                else:
                    print(formatter.format_status_message("Usage: /tableformat [on|off|show]", "info"))
                    continue
            if cmd == "notables":
                val = args_rest.lower()
                if val in ("on", "off"):
                    setattr(args, "disable_table_rendering", val == "on")
                elif val == "":
                    current = bool(getattr(args, "disable_table_rendering", False))
                    setattr(args, "disable_table_rendering", not current)
                else:
                    print(formatter.format_status_message("Usage: /notables on|off", "info"))
                    continue
                
                # Update the formatter immediately
                disabled = bool(getattr(args, "disable_table_rendering", False))
                formatter._disable_tables = disabled
                print(formatter.format_status_message(f"Table rendering: {'DISABLED' if disabled else 'ENABLED'}", "info"))
                save_prefs_from_args(args)
                continue
            if cmd == "session":
                sub = args_rest.strip()
                # Toggle when called without args
                if sub == "":
                    if current_session_enabled:
                        current_session_enabled = False
                        print(formatter.format_status_message("Session disabled.", "info"))
                    else:
                        current_session_enabled = True
                        if current_session_name is None:
                            current_session_name = _suggest_session_name()
                            current_session_file = _session_file_for(current_session_name)
                        print(formatter.format_status_message(f"Session enabled (name={current_session_name}).", "success"))
                        if messages:
                            _session_write_all(messages)
                    continue
                if sub.lower() in ("show", "status"):
                    if current_session_enabled:
                        print(formatter.format_status_message(f"session = on (name={current_session_name or '(auto)'}, file={current_session_file or '(pending)'})", "info"))
                    else:
                        print(formatter.format_status_message("Session disabled.", "info"))
                    continue
                if sub.lower() == "on":
                    current_session_enabled = True
                    if current_session_name is None:
                        current_session_name = _suggest_session_name()
                        current_session_file = _session_file_for(current_session_name)
                    print(formatter.format_status_message(f"Session enabled (name={current_session_name}).", "success"))
                    if messages:
                        _session_write_all(messages)
                    continue
                if sub.lower() == "off":
                    current_session_enabled = False
                    print(formatter.format_status_message("Session disabled.", "info"))
                    continue
                if sub.lower().startswith("name"):
                    toks = sub.split(None, 1)
                    if len(toks) < 2 or not toks[1].strip():
                        print(formatter.format_status_message("Usage: /session name <name>", "info"))
                        continue
                    new_name = _sanitize_session_name(toks[1].strip())
                    current_session_name = new_name
                    current_session_file = _session_file_for(new_name)
                    current_session_enabled = True
                    print(formatter.format_status_message(f"Session name set to {new_name}.", "success"))
                    if messages:
                        _session_write_all(messages)
                    continue
                print(formatter.format_status_message("Usage: /session [on|off|name <n>|show]", "info"))
                continue
            if cmd == "sessions":
                parts = args_rest.split()
                sess = _list_sessions()
                if not parts or parts[0] in ("list", "ls"):
                    if not sess:
                        print(formatter.format_status_message("No sessions found in history (~/ .askp/sessions).", "info"))
                    else:
                        sessions_header = formatter.format_section_header("Available Sessions")
                        print(f"{sessions_header} (newest first)")
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        import datetime as _dt
                        for i, (nm, mt, nlines, duration) in enumerate(sess, 1):
                            t_ago = _short_time_ago(mt)
                            if duration is not None:
                                if duration < 1:
                                    duration_str = f"{int(duration * 60)}s"
                                elif duration < 60:
                                    duration_str = f"{duration:.1f}m"
                                else:
                                    duration_str = f"{duration/60:.1f}h"
                                print(f"  {i:2d}) {nm}  ({nlines} messages, {duration_str}, {t_ago})")
                            else:
                                print(f"  {i:2d}) {nm}  ({nlines} messages, {t_ago})")
                        print(f"\n{formatter.theme.SECTION_SEPARATOR}")
                        if PROMPT_TOOLKIT:
                            try:
                                from prompt_toolkit.widgets import CheckboxList, Frame, Label
                                from prompt_toolkit.layout.containers import HSplit, VSplit, Window
                                from prompt_toolkit.layout.controls import FormattedTextControl
                                from prompt_toolkit.layout.dimension import Dimension
                                from prompt_toolkit.layout import Layout
                                from prompt_toolkit.application import Application
                                from prompt_toolkit.key_binding import KeyBindings

                                # Build checkbox items with labels
                                values = []
                                for (nm, mt, nlines, duration) in sess:
                                    t_ago = _short_time_ago(mt)
                                    if duration is not None:
                                        if duration < 1:
                                            duration_str = f"{int(duration * 60)}s"
                                        elif duration < 60:
                                            duration_str = f"{duration:.1f}m"
                                        else:
                                            duration_str = f"{duration/60:.1f}h"
                                        label = f"{nm}  ({nlines} messages, {duration_str}, {t_ago})"
                                    else:
                                        label = f"{nm}  ({nlines} messages, {t_ago})"
                                    values.append((nm, label))

                                checklist = CheckboxList(values=values)
                                preview_text = FormattedTextControl(text="")
                                preview_window = Window(content=preview_text, wrap_lines=True, height=Dimension(preferred=12))
                                hint = Label(text="Hints: Up/Down to move • Enter toggles items • Space applies • Esc cancels")

                                container = VSplit([
                                    HSplit([
                                        Frame(checklist, title="Select sessions (multi-select)"),
                                        hint,
                                    ], width=Dimension(weight=1)),
                                    Frame(preview_window, title="Preview (focused session)", width=Dimension(weight=2)),
                                ])

                                kb = KeyBindings()

                                def _update_preview():
                                    try:
                                        idx = getattr(checklist, "_selected_index", None)
                                        if idx is None:
                                            preview_text.text = ""
                                            return
                                        nm = checklist.values[idx][0]
                                        msgs = _session_load(nm) or []
                                        max_msgs = 8
                                        start = max(0, len(msgs) - max_msgs)
                                        lines = []
                                        for m in msgs[start:]:
                                            role = m.get("role", "user")
                                            content = m.get("content", "")
                                            content = str(content)
                                            if len(content) > 500:
                                                content = content[:500] + "…"
                                            prefix = "U>" if role == "user" else "A>"
                                            lines.append(f"{prefix} {content}")
                                        preview_text.text = "\n".join(lines) or "(empty)"
                                    except Exception:
                                        preview_text.text = "(preview unavailable)"

                                # Initialize preview
                                _update_preview()

                                @kb.add("enter", eager=True)
                                def _(event):  # type: ignore[func-annotations]
                                    # Toggle item and update preview
                                    try:
                                        if event.app.layout.has_focus(checklist):
                                            idx = getattr(checklist, "_selected_index", None)
                                            if idx is None:
                                                return
                                            val = checklist.values[idx][0]
                                            cur = set(checklist.current_values)
                                            if val in cur:
                                                cur.remove(val)
                                            else:
                                                cur.add(val)
                                            checklist.current_values = list(cur)
                                            # Update preview to current focus
                                            _update_preview()
                                            event.app.invalidate()
                                    except Exception:
                                        pass

                                @kb.add("space", eager=True)
                                def __(event):  # type: ignore[func-annotations]
                                    # Apply: if exactly 1 selected -> open; if >1 -> merge all selected
                                    try:
                                        selected = list(checklist.current_values)
                                        event.app.exit(result=(selected, None))
                                    except Exception:
                                        event.app.exit(result=None)

                                @kb.add("escape", eager=True)
                                def ___(event):  # type: ignore[func-annotations]
                                    event.app.exit(result=None)

                                app = Application(layout=Layout(container), key_bindings=kb, full_screen=False)
                                ret = app.run()
                                if not ret:
                                    print(formatter.format_status_message("Cancelled.", "info"))
                                    continue
                                selected, _unused = ret
                                if not selected:
                                    print(formatter.format_status_message("Nothing selected.", "warning"))
                                    continue
                                # If exactly 1 selection, open it; if more, merge them
                                if len(selected) == 1:
                                    resolved_name = selected[0]
                                    loaded = _session_load(resolved_name)
                                    if not loaded:
                                        print(formatter.format_status_message("Selected session is empty or could not be loaded.", "warning"))
                                        continue
                                    messages.clear()
                                    messages.extend(loaded)
                                    fix_messages_in_place()
                                    current_session_enabled = True
                                    current_session_name = resolved_name
                                    current_session_file = _session_file_for(resolved_name)
                                    print(formatter.format_status_message(f"Opened session '{resolved_name}' with {len(loaded)} messages.", "success"))
                                else:
                                    # Merge in the order shown in the list
                                    ordering = [nm for (nm, _label) in values]
                                    sel_set = set(selected)
                                    merged_count = 0
                                    for nm in ordering:
                                        if nm in sel_set:
                                            loaded = _session_load(nm)
                                            if loaded:
                                                messages.extend(loaded)
                                                merged_count += len(loaded)
                                    current_session_enabled = True
                                    # Set current session to last selected
                                    current_session_name = selected[-1]
                                    current_session_file = _session_file_for(current_session_name)
                                    _session_write_all(messages)
                                    print(formatter.format_status_message(f"Merged {len(selected)} session(s) ({merged_count} messages) into current conversation.", "success"))
                                continue
                            except Exception:
                                # Fallback to CLI input
                                pass
                        # Fallback numeric input flow
                        selection_header = formatter.format_section_header("Selection Options")
                        print(selection_header)
                        print("  Enter number to open session")
                        print("  Type 'merge <number>' to merge with current conversation")
                        print("  Press Enter to cancel")
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        try:
                            selection = input("Selection: ").strip()
                            if not selection:
                                print(formatter.format_status_message("Cancelled.", "info"))
                                continue
                            merge_mode = False
                            if selection.startswith("merge "):
                                merge_mode = True
                                selection = selection[6:].strip()
                            if selection.isdigit():
                                idx = int(selection)
                                if 1 <= idx <= len(sess):
                                    resolved_name = sess[idx - 1][0]
                                    if merge_mode:
                                        loaded = _session_load(resolved_name)
                                        if not loaded:
                                            print(formatter.format_status_message("Selected session is empty or could not be loaded.", "warning"))
                                            continue
                                        messages.extend(loaded)
                                        fix_messages_in_place()
                                        current_session_enabled = True
                                        current_session_name = resolved_name
                                        current_session_file = _session_file_for(resolved_name)
                                        _session_write_all(messages)
                                        print(formatter.format_status_message(f"Merged session '{resolved_name}' ({len(loaded)} messages) into current conversation.", "success"))
                                    else:
                                        loaded = _session_load(resolved_name)
                                        if not loaded:
                                            print(formatter.format_status_message("Selected session is empty or could not be loaded.", "warning"))
                                            continue
                                        messages.clear()
                                        messages.extend(loaded)
                                        current_session_enabled = True
                                        current_session_name = resolved_name
                                        current_session_file = _session_file_for(resolved_name)
                                        print(formatter.format_status_message(f"Opened session '{resolved_name}' with {len(loaded)} messages.", "success"))
                                else:
                                    print(formatter.format_status_message("Invalid selection.", "error"))
                            else:
                                print(formatter.format_status_message("Invalid selection. Please enter a number.", "error"))
                        except (KeyboardInterrupt, EOFError):
                            print("\nCancelled.")
                    continue
                action = parts[0]
                target = parts[1] if len(parts) > 1 else None
                if not target:
                    print(formatter.format_status_message("Usage: /sessions open <name|#> | /sessions merge <name|#> | /sessions list", "info"))
                    continue
                # resolve target
                resolved_name = None
                if target.isdigit():
                    idx = int(target)
                    if 1 <= idx <= len(sess):
                        resolved_name = sess[idx - 1][0]
                else:
                    for (nm, _, _) in sess:
                        if nm == target:
                            resolved_name = nm
                            break
                if not resolved_name:
                    print(formatter.format_status_message("Session not found.", "error"))
                    continue
                if action == "open":
                    loaded = _session_load(resolved_name)
                    if not loaded:
                        print(formatter.format_status_message("Selected session is empty or could not be loaded.", "warning"))
                        continue
                    messages.clear()
                    messages.extend(loaded)
                    fix_messages_in_place()
                    current_session_enabled = True
                    current_session_name = resolved_name
                    current_session_file = _session_file_for(resolved_name)
                    print(formatter.format_status_message(f"Opened session '{resolved_name}' with {len(loaded)} messages.", "success"))
                    continue
                if action == "merge":
                    loaded = _session_load(resolved_name)
                    if not loaded:
                        print(formatter.format_status_message("Selected session is empty or could not be loaded.", "warning"))
                        continue
                    messages.extend(loaded)
                    fix_messages_in_place()
                    current_session_enabled = True
                    current_session_name = resolved_name
                    current_session_file = _session_file_for(resolved_name)
                    _session_write_all(messages)
                    print(formatter.format_status_message(f"Merged session '{resolved_name}' ({len(loaded)} messages) into current conversation.", "success"))
                    continue
                print(formatter.format_status_message("Usage: /sessions open <name|#> | /sessions merge <name|#> | /sessions list", "info"))
                continue
            if cmd == "settings":
                settings_header = formatter.format_section_header("Current Settings")
                print(settings_header)
                print(f"{formatter.theme.SECTION_SEPARATOR}")
                print(formatter.format_status_message(f"  model       = {args.model}", "info"))
                print(formatter.format_status_message(f"  academic    = {args.academic}", "info"))
                print(formatter.format_status_message(f"  domain      = {args.domain}", "info"))
                print(formatter.format_status_message(f"  recency     = {args.recency}", "info"))
                print(formatter.format_status_message(f"  jsonschema  = {args.json_schema}", "info"))
                print(formatter.format_status_message(f"  citations   = {args.citations}", "info"))
                print(formatter.format_status_message(f"  usage       = {args.usage}", "info"))
                print(formatter.format_status_message(f"  stream      = {args.stream}", "info"))
                print(formatter.format_status_message(f"  text        = {getattr(args, 'text', False)} (one-shot default output)", "info"))
                print(
                    f"  attachlimit = {getattr(args, 'attach_truncation_limit', 8000)}"
                )
                print(
                    f"  prefs file  = {CONFIG_PATH if os.path.exists(CONFIG_PATH) else '(none)'}"
                )
                print(
                    f"  session     = {'on' if current_session_enabled else 'off'} (name={current_session_name or '(auto)'})"
                )
                persistent_sys = getattr(args, "persistent_system", None)
                if persistent_sys:
                    print(formatter.format_status_message(f"  persistent  = {persistent_sys[:50]}{'...' if len(persistent_sys) > 50 else ''}", "info"))
                else:
                    print(formatter.format_status_message("  persistent  = (none)", "info"))
                print(f"{formatter.theme.SECTION_SEPARATOR}")
                continue
            if cmd in ("prune", "compact"):
                # Summarize the full conversation history and restart with the summary as system prompt
                if not messages:
                    print(formatter.format_status_message("Nothing to summarize.", "info"))
                    continue
                # Parse optional word limit and optional focus hint
                limit = 200
                focus_hint = ""
                try:
                    import shlex as _shlex
                    parts = _shlex.split(args_rest)
                except Exception:
                    parts = args_rest.split()
                # If first token is an integer, treat it as the word limit
                if parts and parts[0].isdigit():
                    try:
                        limit = max(50, min(1000, int(parts[0])))
                    except Exception:
                        limit = 200
                    parts = parts[1:]
                if parts:
                    focus_hint = " ".join(parts).strip()
                    # Normalize common phrasing like "focus on ..." or "focus: ..."
                    fl = focus_hint.lower()
                    if fl.startswith("focus on "):
                        focus_hint = focus_hint[10:].strip()
                    elif fl.startswith("focus "):
                        focus_hint = focus_hint[6:].strip()
                    elif fl.startswith("focus: "):
                        focus_hint = focus_hint[7:].strip()
                # Build summarization request using the entire message history plus an instruction
                summarize_messages = list(messages)
                instruction = (
                    f"Summarize the above conversation into a concise brief (<= {limit} words) "
                    f"capturing key decisions, facts, and context. Reply with only the summary."
                )
                if focus_hint:
                    instruction += f" Emphasize: {focus_hint}."
                summarize_messages.append({"role": "user", "content": instruction})
                payload = make_payload_from_messages(summarize_messages, args)
                # Ensure non-streaming request for prune/compact
                payload.pop("stream", None)
                try:
                    resp = requests.post(config.api_url, headers=headers, json=payload)
                    resp.raise_for_status()
                    data = resp.json()
                    summary = extract_content(data) or ""
                    if not summary:
                        print(formatter.format_status_message("Failed to summarize.", "error"))
                        continue
                    # Restart conversation with summary as system prompt
                    messages.clear()
                    messages.append(
                        {
                            "role": "system",
                            "content": f"Conversation summary (<= {limit} words){': ' + summary if summary else ''}",
                        }
                    )
                    print(formatter.format_status_message("Conversation compacted via summary. New system prompt set.", "success"))
                except Exception as e:
                    print(formatter.format_status_message(f"Prune failed: {e}", "error"))
                continue
            print(formatter.format_status_message("Unknown command. Type /help", "warning"))
            continue
        # Regular user message - check for pending attachments
        if attachment_state.has_attachments():
            # Build message with preprompt attachments
            attachment_summary = attachment_state.get_summary()
            enhanced_input = attachment_state.build_preprompt(user_input)
            messages.append({"role": "user", "content": enhanced_input})
            _session_write({"role": "user", "content": enhanced_input})
            print(formatter.format_status_message(f"Using {attachment_summary} with your message", "info"))
            # Clear attachments after use
            attachment_state.clear()
        else:
            messages.append({"role": "user", "content": user_input})
            _session_write({"role": "user", "content": user_input})
        payload = make_payload_from_messages(messages, args)
        logging.getLogger("askp").debug("Payload: %s", json.dumps(payload, indent=2))
        try:
            # Streaming path prints styled output as it arrives (md_render controls inline markdown).
            effective_stream = bool(args.stream)
            # Force non-streaming when structured JSON schema is enabled for cleaner output
            if getattr(args, "json_schema", False):
                effective_stream = False
            # Ensure payload matches streaming mode (Perplexity returns SSE when stream=true)
            if not effective_stream:
                payload.pop("stream", None)
            if effective_stream:
                with requests.post(
                    config.api_url, headers=headers, json=payload, stream=True
                ) as resp:
                    logging.getLogger("askp").debug("HTTP %s", resp.status_code)
                    resp.raise_for_status()
                    # Begin formatted streaming output
                    if getattr(args, "md_render", True):
                        answer_header = formatter.format_section_header("Answer")
                        print(answer_header)
                        print(f"{formatter.theme.SECTION_SEPARATOR}")
                        # For markdown mode, we'll handle the prefix in the flush function
                        prefix_printed = False
                    else:
                        print(f"{formatter.theme.AI_COLOR}AI>{formatter.theme.RESET} ", end="", flush=True)
                        prefix_printed = True
                    acc_tokens = []
                    acc_text = ""
                    last_formatted = ""
                    last_flush_time = time.time()
                    citations_acc = None
                    usage_acc = None
                    
                    def stream_flush():
                        """Format complete text and print only new content."""
                        nonlocal last_formatted, prefix_printed
                        
                        if not acc_text:
                            return
                        
                        # Format the complete accumulated text
                        if getattr(args, "md_render", True):
                            current_formatted = formatter.format_ai_response(acc_text)
                        else:
                            # Plain text with proper wrapping
                            prefix = f"{formatter.theme.AI_COLOR}AI>{formatter.theme.RESET}"
                            prefix_len = len("AI> ")
                            wrapped = formatter._wrap_text(acc_text, indent=prefix_len)
                            lines = wrapped.split('\n')
                            if lines:
                                lines[0] = f"{prefix} {lines[0].lstrip()}"
                                current_formatted = '\n'.join(lines)
                            else:
                                current_formatted = f"{prefix} "
                        
                        # Print only the new portion
                        if len(current_formatted) > len(last_formatted):
                            new_content = current_formatted[len(last_formatted):]
                            if new_content:
                                print(new_content, end="", flush=True)
                                last_formatted = current_formatted
                    for line in resp.iter_lines(decode_unicode=True):
                        if not line:
                            continue
                        # handle both SSE-like 'data:' and plain JSON lines
                        s = line
                        if s.startswith("data:"):
                            s = s[5:].strip()
                            if s == "[DONE]":
                                break
                        try:
                            obj = json.loads(s)
                        except Exception:
                            continue
                        # Accumulate any citations/usage metadata if present in stream chunk
                        if isinstance(obj, dict):
                            if citations_acc is None and isinstance(obj.get("citations"), list):
                                citations_acc = obj.get("citations")
                            if usage_acc is None and isinstance(obj.get("usage"), dict):
                                usage_acc = obj.get("usage")
                        token = extract_delta_token(obj)
                        if token:
                            acc_tokens.append(token)
                            acc_text += token
                            
                            # Check for complete sentences or natural break points
                            now = time.time()
                            time_since_last = now - last_flush_time
                            
                            should_flush_now = (
                                acc_text.endswith(('.', '!', '?')) and time_since_last > 0.1 or
                                acc_text.endswith(':') and time_since_last > 0.1 or
                                acc_text.endswith('\n\n') or
                                acc_text.endswith('\n') and time_since_last > 0.3 or
                                time_since_last > 1.0
                            )
                            
                            if should_flush_now:
                                stream_flush()
                                last_flush_time = now
                    
                    # Final flush of any remaining content
                    stream_flush()
                    if getattr(args, "md_render", True):
                        print(f"\n{formatter.theme.SECTION_SEPARATOR}")
                    print()
                    assistant_text = acc_text
                    if assistant_text:
                        messages.append(
                            {"role": "assistant", "content": assistant_text}
                        )
                        _session_write({"role": "assistant", "content": assistant_text})
                    # After stream completes, print citations/usage if requested and available
                    if args.citations and isinstance(citations_acc, list):
                        print_citations(formatter, citations_acc)
                    if args.usage and isinstance(usage_acc, dict):
                        print_usage(formatter, usage_acc)
                    # Add visual separator after complete response section
                    if (args.citations and isinstance(citations_acc, list)) or (args.usage and isinstance(usage_acc, dict)):
                        print(f"\n{formatter.theme.SECTION_SEPARATOR}")
                    print()  # Add spacing after AI response
            else:
                # Non-stream path; show spinner until response is ready and print formatted once
                stop_spinner = False
                spinner_thread = None
                def _spinner():
                    seq = ['-', '\\', '|', '/']
                    i = 0
                    while not stop_spinner:
                        try:
                            print(f"\r{formatter.theme.METADATA_COLOR}working {seq[i % len(seq)]}{formatter.theme.RESET}", end="", flush=True)
                        except Exception:
                            pass
                        time.sleep(0.1)
                        i += 1
                    # clear line
                    try:
                        print("\r" + " " * 60 + "\r", end="", flush=True)
                    except Exception:
                        pass
                spinner_thread = threading.Thread(target=_spinner, daemon=True)
                spinner_thread.start()
                try:
                    response = requests.post(config.api_url, headers=headers, json=payload)
                    logging.getLogger("askp").debug("HTTP %s", response.status_code)
                    response.raise_for_status()
                    data = response.json()
                finally:
                    if spinner_thread:
                        stop_spinner = True
                        spinner_thread.join(timeout=1.0)
                content = extract_content(data)
                if content is None:
                    print(formatter.format_json_output(data))
                else:
                    # Print formatted content depending on md_render
                    if getattr(args, "md_render", True):
                        formatted_response = formatter.format_ai_response(content)
                        print(formatted_response)
                    else:
                        print(content)
                    messages.append({"role": "assistant", "content": content})
                    _session_write({"role": "assistant", "content": content})
                    if args.citations and isinstance(data.get("citations"), list):
                        print_citations(formatter, data["citations"])  # printed after content
                    if args.usage and isinstance(data.get("usage"), dict):
                        print_usage(formatter, data["usage"])  # printed after content
                    # Add visual separator after complete response section
                    if (args.citations and isinstance(data.get("citations"), list)) or (args.usage and isinstance(data.get("usage"), dict)):
                        print(f"\n{formatter.theme.SECTION_SEPARATOR}")
                    print()  # Add spacing after AI response
        except requests.exceptions.HTTPError as e:
            sys.stderr.write(f"HTTP error: {e}\n")
            if e.response is not None:
                try:
                    sys.stderr.write(e.response.text + "\n")
                except Exception:
                    pass
        except Exception as e:
            sys.stderr.write(f"Unexpected error: {e}\n")
    return 0


# --- Argument Parser ---
parser = argparse.ArgumentParser(description="Ask Perplexity AI from CLI")
parser.add_argument(
    "--no-color", action="store_true", help="Disable colored output"
)
parser.add_argument(
    "query", nargs="?", help="Your question (omit to start interactive chat)"
)
parser.add_argument(
    "-a", "--academic", action="store_true", help="Use academic search filter"
)
parser.add_argument(
    "-d", "--domain", help="Restrict search to a specific domain (e.g. arxiv.org)"
)
parser.add_argument(
    "--recency", choices=["day", "week", "month"], help="Search recency filter"
)
parser.add_argument(
    "--json-schema",
    action="store_true",
    help="Enable structured JSON response (uses an example schema)",
)
parser.add_argument(
    "--stream", action="store_true", help="Stream tokens as they arrive"
)
parser.add_argument(
    "--no-stream",
    action="store_true",
    help="Disable streaming (interactive chat defaults to streaming)",
)
parser.add_argument(
    "-m",
    "--model",
    choices=AVAILABLE_MODELS,
    default=DEFAULT_MODEL,
    help=f"Model to use (default: {DEFAULT_MODEL})",
)
parser.add_argument(
    "-t",
    "--text",
    action="store_true",
    help="Print only the assistant text response instead of raw JSON",
)
parser.add_argument(
    "-u",
    "--usage",
    action="store_true",
    help="Show token usage summary if available (only for --text mode)",
)
parser.add_argument(
    "-c",
    "--citations",
    action="store_true",
    help="Show citations if available (only for --text mode)",
)
parser.add_argument(
    "--max-tokens", type=int, dest="max_tokens", help="Max tokens for response"
)
parser.add_argument(
    "--temperature", type=float, dest="temperature", help="Sampling temperature"
)
parser.add_argument("--top-p", type=float, dest="top_p", help="Top-p nucleus sampling")
parser.add_argument(
    "--reasoning-effort",
    choices=["low", "medium", "high"],
    dest="reasoning_effort",
    help="Reasoning effort (async only)",
)
parser.add_argument(
    "--search-mode",
    choices=["web", "none"],
    dest="search_mode",
    help="Search mode (async only)",
)
parser.add_argument(
    "--return-images",
    action="store_true",
    dest="return_images",
    help="Return images (async only)",
)
parser.add_argument(
    "--return-related-questions",
    action="store_true",
    dest="return_related_questions",
    help="Return related questions (async only)",
)
parser.add_argument(
    "--async",
    action="store_true",
    dest="async_mode",
    help="Use async API for single-shot",
)
parser.add_argument(
    "--wait",
    action="store_true",
    dest="wait_async",
    help="Wait for async result when using --async",
)
parser.add_argument("-v", "--verbose", action="store_true", help="Enable debug logging")
parser.add_argument(
    "--chat",
    nargs="?",
    const="",
    help="Start interactive chat; optional seed message if provided",
)
parser.add_argument(
    "--attach-limit",
    type=int,
    dest="attach_truncation_limit",
    help="Truncation limit for /attach --as-user",
)
# Session controls
parser.add_argument(
    "--session",
    action="store_true",
    dest="session_enabled",
    help="[Deprecated] Sessions are enabled by default; use --no-session or /session to disable",
)
parser.add_argument(
    "--no-session",
    action="store_true",
    dest="no_session",
    help="Disable project session persistence",
)
parser.add_argument(
    "--session-name",
    dest="session_name",
    help="Name of the project session file (without extension)",
)
def main():
    """Main entry point for the askp CLI."""
    args = parser.parse_args()

    # --- Setup logging ---
    logging.basicConfig(
    level=logging.DEBUG if args.verbose else logging.WARNING,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    )
    logger = logging.getLogger("askp")

    # --- Initialize formatter ---
    formatter = OutputFormatter()

    # --- API key ---
    API_KEY = get_api_key()
    if not API_KEY:
        # If interactive terminal, prompt for API key and offer to persist
        if sys.stdin.isatty():
            print(formatter.format_status_message("No Perplexity API key found in environment (PPLX_API_KEY or PERPLEXITY_API_KEY).", "error"))
            try:
                key_input = getpass.getpass("Enter your Perplexity API key (input hidden): ").strip()
            except Exception:
                key_input = input("Enter your Perplexity API key: ").strip()
            if not key_input:
                sys.stderr.write("No API key provided. Exiting.\n")
                sys.exit(1)
            # Set for current process so the run can proceed
            os.environ["PPLX_API_KEY"] = key_input
            API_KEY = key_input
            # Offer to persist
            ans = input("Persist this key for future sessions? [Y/n]: ").strip().lower()
            if ans in ("", "y", "yes"):
                if os.name == "nt":
                    ok = persist_api_key_windows(key_input)
                    if ok:
                        print(formatter.format_status_message("Saved PPLX_API_KEY to your user environment via setx. You may need to open a new terminal.", "info"))
                    else:
                        print(formatter.format_status_message("Failed to persist using setx. You can set it manually in PowerShell: $env:PPLX_API_KEY=\"...\"", "error"))
                else:
                    path = persist_api_key_unix(key_input)
                    if path:
                        print(formatter.format_status_message(f"Added export to {path}. Run 'source {path}' or open a new shell to apply.", "info"))
                    else:
                        print("Failed to write to a shell profile. You can manually add 'export PPLX_API_KEY=\"...\"' to your shell config.")
        else:
            sys.stderr.write(
                "Error: API key not found and input is not interactive. Set PPLX_API_KEY (preferred) or PERPLEXITY_API_KEY.\n"
            )
            sys.exit(1)

    # --- Build request or start interactive chat ---
    config = ApiConfig(api_key=API_KEY)
    headers = {
        "Authorization": f"Bearer {config.api_key}",
        "Content-Type": "application/json",
    }

    # Decide streaming defaults
    # Default: interactive chat streams unless --no-stream is passed; single-shot does not stream unless --stream is passed
    # Load persisted preferences before determining streaming default
    prefs = load_prefs()
    interactive_requested = (args.query is None) or (args.chat is not None)

    # Store original stream setting before applying preferences
    original_stream = args.stream

    if prefs:
        apply_prefs_to_args(prefs, args)

    # Override streaming preference for single-shot mode
    if not interactive_requested:
        # In single-shot mode, only stream if explicitly requested via --stream
        # Don't use saved streaming preference for single-shot
        args.stream = original_stream
    if (
        interactive_requested
        and not args.no_stream
        and not getattr(args, "_stream_pref_loaded", False)
    ):
        args.stream = True

    # Warn if running from the user's home directory (recommend project-specific dir)
    try:
        suppress = os.getenv("ASKP_SUPPRESS_HOME_WARNING")
        if not suppress:
            cwd = os.path.realpath(os.getcwd())
            home_dir = os.path.realpath(os.path.expanduser("~"))
            if cwd == home_dir and (sys.stdin.isatty() or interactive_requested):
                # Ensure ANSI colors work on Windows
                try:
                    colorama_init(autoreset=True)
                except Exception:
                    pass
                lines = [
                    "Home Directory",
                ]
                width = max(len(l) for l in lines)
                border = "+" + ("-" * (width + 2)) + "+"
                body = "\n".join([f"| {l.ljust(width)} |" for l in lines])
                msg = f"{border}\n{body}\n{border}\n"
                sys.stderr.write(Fore.YELLOW + Style.BRIGHT + msg + Style.RESET_ALL)
    except Exception:
        pass

    # Create formatter for both interactive and single-shot modes
    prefs = load_prefs()
    user_theme = (prefs or {}).get("theme") if prefs else None
    if user_theme and not isinstance(user_theme, dict):
        user_theme = None
    md_render_pref = (prefs or {}).get("md_render") if prefs else None
    if md_render_pref is not None:
        args.md_render = bool(md_render_pref)
    # Remove legacy preferences from older versions
    # Legacy prefs (styled_stream, final_style) are ignored going forward
    formatter = OutputFormatter(use_colors=sys.stdout.isatty() and not args.no_color, theme_cfg=user_theme or {}, md_render=getattr(args, "md_render", True))
    
    # Apply persistent table rendering setting
    disable_tables = bool(getattr(args, 'disable_table_rendering', False))
    formatter._disable_tables = disable_tables

    # Interactive chat mode when no query is provided or --chat is used
    if args.query is None or args.chat is not None:
        seed = None
        if args.chat is not None:
            seed = args.chat if args.chat != "" else None
        elif args.query is not None:
            seed = args.query
        seed_messages: Optional[List[Dict[str, Any]]] = None
        if seed:
            # Seed the conversation with the first user message and immediate response
            colorama_init(autoreset=True)
            seed_messages = []
            seed_messages.append({"role": "user", "content": seed})
            payload = make_payload_from_messages(seed_messages, args)
            try:
                response = requests.post(config.api_url, headers=headers, json=payload)
                response.raise_for_status()
                data = response.json()
                content = extract_content(data)
                if content is None:
                    print(formatter.format_json_output(data))
                else:
                    formatted_response = formatter.format_ai_response(content)
                    print(formatted_response)
                    seed_messages.append({"role": "assistant", "content": content})
            except Exception as e:
                sys.stderr.write(f"Unexpected error during seed: {e}\n")
                # fall through to chat loop anyway
        sys.exit(chat_loop(args, headers, config, messages=seed_messages))

    # Single-shot mode
    if args.async_mode:
        # Build async request body for single-shot
        model_val = args.model
        if model_val != "sonar-deep-research":
            sys.stderr.write(
                "Note: async API only supports 'sonar-deep-research'; overriding model for this request. Use -m sonar-deep-research.\n"
            )
            model_val = "sonar-deep-research"
        request_obj: Dict[str, Any] = {
            "model": model_val,
            "messages": [
                {
                    "role": "user",
                    "content": [{"type": "text", "text": args.query}],
                }
            ],
        }
        if getattr(args, "max_tokens", None):
            request_obj["max_tokens"] = args.max_tokens
        if getattr(args, "temperature", None) is not None:
            request_obj["temperature"] = args.temperature
        if getattr(args, "top_p", None) is not None:
            request_obj["top_p"] = args.top_p
        if getattr(args, "search_mode", None):
            request_obj["search_mode"] = args.search_mode
        if getattr(args, "reasoning_effort", None):
            request_obj["reasoning_effort"] = args.reasoning_effort
        if getattr(args, "return_images", None):
            request_obj["return_images"] = args.return_images
        if getattr(args, "return_related_questions", None):
            request_obj["return_related_questions"] = args.return_related_questions
        if args.academic:
            request_obj["search_filter"] = "academic"
        if args.domain:
            request_obj["search_domain_filter"] = [args.domain]
        if args.recency:
            request_obj["search_recency_filter"] = args.recency

        body = {"request": request_obj}

        try:
            logging.getLogger("askp").debug("Submitting async request to %s", ASYNC_API_URL)
            submit_resp = requests.post(ASYNC_API_URL, headers=headers, json=body)
            logging.getLogger("askp").debug("HTTP %s", submit_resp.status_code)
            submit_resp.raise_for_status()
            submit_data = submit_resp.json()
            request_id = (
                submit_data.get("id")
                or submit_data.get("request_id")
                or submit_data.get("requestId")
            )

            if args.wait_async and request_id:
                # Poll until completed or failed
                start = time.time()
                while True:
                    status_resp = requests.get(
                        f"{ASYNC_API_URL}/{request_id}", headers=headers
                    )
                    status_resp.raise_for_status()
                    result_data = status_resp.json()
                    status = (
                        result_data.get("status") or result_data.get("state") or ""
                    ).upper()
                    if status in ("COMPLETED", "SUCCEEDED"):
                        data = result_data
                        break
                    if status in ("FAILED", "ERROR"):
                        data = result_data
                        break
                    if time.time() - start > 300:
                        raise TimeoutError("Async wait timed out")
                    time.sleep(1.0)
            else:
                data = submit_data

            if args.text:
                content = extract_content(data)
                if content is None:
                    sys.stderr.write(
                        "No text content found in response. Printing raw JSON.\n"
                    )
                    print(formatter.format_json_output(data))
                else:
                    formatted_response = formatter.format_ai_response(content)
                    print(formatted_response)
                    if args.citations and isinstance(data.get("citations"), list):
                        print_citations(formatter, data["citations"])  # printed after content
                    if args.usage and isinstance(data.get("usage"), dict):
                        print_usage(formatter, data["usage"])  # printed after content
                    # Add visual separator after complete response section
                    if (args.citations and isinstance(data.get("citations"), list)) or (args.usage and isinstance(data.get("usage"), dict)):
                        print(f"\n{formatter.theme.SECTION_SEPARATOR}")
            else:
                print(formatter.format_json_output(data))

        except requests.exceptions.HTTPError as e:
            sys.stderr.write(f"HTTP error: {e}\n")
            if e.response is not None:
                try:
                    sys.stderr.write(e.response.text + "\n")
                except Exception:
                    pass
            sys.exit(1)
        except Exception as e:
            sys.stderr.write(f"Unexpected error: {e}\n")
            sys.exit(1)
    else:
        # Force async-compatible model for sync path? No, keep user's choice.
        payload = make_payload_from_messages(
            [{"role": "user", "content": args.query}], args
        )

        logging.getLogger("askp").debug("API URL: %s", config.api_url)
        logging.getLogger("askp").debug("Model: %s", args.model)
        logging.getLogger("askp").debug(
            "Filters - academic: %s, domain: %s, recency: %s",
            args.academic,
            args.domain,
            args.recency,
        )
        logging.getLogger("askp").debug("Payload: %s", json.dumps(payload, indent=2))

        # --- Send request ---
        try:
            response = requests.post(config.api_url, headers=headers, json=payload)
            logging.getLogger("askp").debug("HTTP %s", response.status_code)
            response.raise_for_status()
            data = response.json()

            if args.text:
                content = extract_content(data)
                if content is None:
                    sys.stderr.write(
                        "No text content found in response. Printing raw JSON.\n"
                    )
                    print(formatter.format_json_output(data))
                else:
                    formatted_response = formatter.format_ai_response(content)
                    print(formatted_response)
                    if args.citations and isinstance(data.get("citations"), list):
                        print_citations(formatter, data["citations"])  # printed after content
                    if args.usage and isinstance(data.get("usage"), dict):
                        print_usage(formatter, data["usage"])  # printed after content
                    # Add visual separator after complete response section
                    if (args.citations and isinstance(data.get("citations"), list)) or (args.usage and isinstance(data.get("usage"), dict)):
                        print(f"\n{formatter.theme.SECTION_SEPARATOR}")
            else:
                # Raw JSON for piping to tools like jq
                print(formatter.format_json_output(data))

        except requests.exceptions.HTTPError as e:
            sys.stderr.write(f"HTTP error: {e}\n")
            if e.response is not None:
                try:
                    sys.stderr.write(e.response.text + "\n")
                except Exception:
                    pass
            sys.exit(1)
        except Exception as e:
            sys.stderr.write(f"Unexpected error: {e}\n")
            sys.exit(1)

if __name__ == "__main__":
    main()
