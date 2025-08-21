from pathlib import Path
from typing import Optional

from docx import Document as DocxDocument
from docx.shared import Pt

from mitoolspro.document.document_structure import (
    BBox,
    Box,
    Char,
    Document,
    Image,
    Line,
    Page,
    Run,
)


def docx_to_document(docx_path: Path) -> Document:
    doc = Document()
    docx = DocxDocument(str(docx_path))
    # DOCX does not have explicit page size, so we use a default A4 size in points
    PAGE_WIDTH = 595.3  # 8.27 inch * 72 pt/inch
    PAGE_HEIGHT = 841.9  # 11.69 inch * 72 pt/inch
    page = Page(PAGE_WIDTH, PAGE_HEIGHT)

    for para in docx.paragraphs:
        # Use a dummy bbox for the box and line, as docx does not provide coordinates
        box_bbox = BBox(0, 0, PAGE_WIDTH, 20)  # Height is arbitrary per box
        box = Box(box_bbox)
        line_bbox = BBox(0, 0, PAGE_WIDTH, 20)  # Height is arbitrary per line
        line = Line(line_bbox)
        for run in para.runs:
            fontname = run.font.name or "Times New Roman"
            size = run.font.size.pt if run.font.size else 12.0
            for c in run.text:
                # Each char gets a dummy bbox
                char_bbox = BBox(0, 0, 0, 0)
                char = Char(c, fontname, size, char_bbox)
                # If the last run matches, append to it, else create new run
                if (
                    line.runs
                    and line.runs[-1].fontname == fontname
                    and abs(line.runs[-1].size - size) < 0.01
                ):
                    line.runs[-1].append_char(char)
                else:
                    new_run = Run(fontname, size)
                    new_run.append_char(char)
                    line.add_run(new_run)
        if line.runs:
            box.add_line(line)
        if box.elements:
            page.add_box(box)
    doc.add_page(page)
    return doc
