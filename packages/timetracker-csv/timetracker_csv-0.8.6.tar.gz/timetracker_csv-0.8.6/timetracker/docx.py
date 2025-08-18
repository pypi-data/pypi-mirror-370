"""Generate a Microsoft Word document containing a table of data"""

__copyright__ = 'Copyright (C) 2025-present, DV Klopfenstein, PhD. All rights reserved.'
__author__ = "DV Klopfenstein, PhD"


def write_invoice(fout_docx, billable):
    """Write an invoice into a Microsoft Word document"""
    if (obj := _get_worddoc()) is not None:
        obj.write_invoice(fout_docx, billable)
        return True
    return _msg_docx_not_installed('INVOICE', fout_docx)

def write_doc(fout_docx, time_formatted):
    """Write a report into a Microsoft Word document"""
    if (obj := _get_worddoc(time_formatted)) is not None:
        obj.write_doc(fout_docx)
        return True
    return _msg_docx_not_installed('REPORT', fout_docx)

def _msg_docx_not_installed(typ, fout_docx):
    print('INFO: `docx` not installed')
    print(f'INFO: CANNOT WRITE DOCX {typ}: {fout_docx}')
    print('INFO:  Do `pip install python-docx` to install module `docx`')
    return False

def _get_worddoc(time_formatted=None):
    """Get WordDoc instance if python-docx is installed"""

    try:
        # Installing `lxml`, which is required by `python-docx`, can be difficult to install
        # pylint: disable=import-outside-toplevel
        #https://python-docx.readthedocs.io/en/latest/
        from docx import Document
        from docx.shared import Inches

        class WordDoc:
            """Generate a Microsoft Word document containing a table of data"""
            # pylint: disable=too-few-public-methods

            wdct = {
                'Day': Inches(.18),
                'Date': Inches(.95),
                'Duration': Inches(.5),
                'Span': Inches(.6),
                'Total': Inches(.5),
                'Price': Inches(.5),
                'Due': Inches(.7),
                'Description': Inches(3),
            }

            def __init__(self, time_formatted=None):
                self.nttext = time_formatted

            def write_invoice(self, fout_docx, nts):
                """Write an invoice into a Microsoft Word document"""
                document = Document()
                document.add_heading('Invoice', 0)
                self._add_table_invoice(document, nts)
                document.add_page_break()
                document.save(fout_docx)

            def write_doc(self, fout_docx):
                """Write a report into a Microsoft Word document"""
                document = Document()
                document.add_heading('Document Title', 0)

                #p = document.add_paragraph('A plain paragraph having some ')
                #p.add_run('bold').bold = True
                #p.add_run(' and some ')
                #p.add_run('italic.').italic = True

                #document.add_heading('Heading, level 1', level=1)
                #document.add_paragraph('Intense quote', style='Intense Quote')

                #document.add_paragraph(
                #    'first item in unordered list', style='List Bullet'
                #)
                #document.add_paragraph(
                #    'first item in ordered list', style='List Number'
                #)

                #document.add_picture('monty-truth.png', width=Inches(1.25))

                self._add_table(document)
                document.add_page_break()
                document.save(fout_docx)

            def _get_headers(self):
                """Get the number of rows in the timetracking data (self.nttext must have data)"""
                return self.nttext[0]._fields

            def _get_nrows(self):
                """Get the number of rows in the timetracking data (self.nttext must have data)"""
                return len(self.nttext)

            def _get_ncols(self):
                """Get the number of rows in the timetracking data (self.nttext must have data)"""
                return len(self.nttext[0])

            def _add_table_invoice(self, doc, nts):
                """Add a table containing timetracking data to a Word document"""
                if not nts:
                    return
                nt1 = nts[0]
                print('FIRST NT:', nt1)
                table = doc.add_table(rows=1, cols=len(nt1), style='Table Grid')
                # https://python-docx.readthedocs.io/en/latest/api/enum/WdRowAlignment.html
                ##print('TTTTTTTTTTTTTTTTTTTTTT', dir(table))
                ##print('TTTTTTTTTTTTTTTTTTTTTT', table.style)
                ##print('TTTTTTTTTTTTTTTTTTTTTT', table.alignment)
                hdrs = nt1._fields
                self._add_table_headers(table, hdrs)
                for ntd in nts:
                    ##print('DDDDDDDDDDDDDDDD', ntd)
                    row_cells = table.add_row().cells
                    for hdr, cell, val in zip(hdrs, row_cells, list(ntd)):
                        ##pcell = cell.add_paragraph()
                        ##pcell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if hdr in self.wdct:
                            cell.width = self.wdct[hdr]
                        cell.text = val
                        ##print('CCCCCCCCCCCCCCCCCCCCCCCCCCCC', dir(pcell.style))

            def _add_table(self, doc):
                """Add a table containing timetracking data to a Word document"""
                if not self.nttext:
                    return
                table = doc.add_table(rows=1, cols=self._get_ncols(), style='Table Grid')
                hdrs = self._get_headers()
                self._add_table_headers(table, hdrs)
                for ntd in self.nttext:
                    row_cells = table.add_row().cells
                    for hdr, cell, val in zip(hdrs, row_cells, list(ntd)):
                        if hdr in self.wdct:
                            cell.width = self.wdct[hdr]
                        cell.text = val

            def _add_table_headers(self, table, hdrs):
                """Add a table header for timetracking data to a Word document"""
                for hdr, cell in zip(hdrs, table.rows[0].cells):
                    if hdr in self.wdct:
                        cell.width = self.wdct[hdr]
                    cell.text = hdr

        return WordDoc(time_formatted)

    except ModuleNotFoundError as err:
        print(f'ERROR: {err}')
    return None


# Copyright (C) 2025-present, DV Klopfenstein, PhD. All rights reserved.
