import docx
from unidoc_agent.base_tool import BaseTool
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

class DocxTool(BaseTool):
    def can_handle(self, file_path, mime_type):
        return mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def extract_content(self, file_path):
        doc = docx.Document(file_path)
        content = []

        def iter_block_items(parent):
            for child in parent.element.body.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield docx.table.Table(child, parent)

        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                if block.text.strip():
                    content.append(block.text.strip())
            elif isinstance(block, docx.table.Table):
                for row in block.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        content.append("\t".join(row_text))

        # Add headers/footers if needed
        try:
            for section in doc.sections:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        content.insert(0, "[Header] " + p.text.strip())
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        content.append("[Footer] " + p.text.strip())
        except Exception:
            pass

        return "\n".join(content)