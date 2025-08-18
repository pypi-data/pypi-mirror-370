import os
import html
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Optional

SUPPORTED_FORMATS = {
    "txt",
    "pdf",
    "epub",
    "mobi",
    "azw",
    "azw3",
    "srt",
    "vtt",
    "json",
    "md",
    "docx",
}


def infer_format_from_path(path: Optional[str]) -> Optional[str]:
    if not path:
        return None
    ext = Path(path).suffix.lower().lstrip(".")
    return ext if ext in SUPPORTED_FORMATS else None


def export_transcript(
    text: str,
    out_path: str,
    fmt: str,
    title: Optional[str] = None,
    author: Optional[str] = None,
    cover_image: Optional[str] = None,
    pdf_font: str = "Arial",
    pdf_font_size: int = 12,
    pdf_margin: int = 15,
    epub_css_file: Optional[str] = None,
    epub_css_text: Optional[str] = None,
    pdf_cover_fullpage: bool = False,
    pdf_first_page_cover_only: bool = False,
    pdf_page_size: str = "A4",
    pdf_orientation: str = "portrait",
    pdf_font_file: Optional[str] = None,
    # Optional rich metadata for advanced formats
    segments: Optional[list[dict]] = None,
    words: Optional[list[dict]] = None,
    metadata: Optional[dict] = None,
    # PDF/EPUB richness
    pdf_header: Optional[str] = None,
    pdf_footer: Optional[str] = None,
    auto_toc: bool = False,
    cover_image_bytes: Optional[bytes] = None,
    # DOCX options
    docx_cover_first: bool = False,
    docx_cover_width_inches: Optional[float] = None,
    docx_footer_text: Optional[str] = None,
    docx_footer_include_page_number: bool = True,
    # PDF attribution page
    pdf_append_attribution: bool = False,
    pdf_attribution_text: Optional[str] = None,
    # DOCX attribution page
    docx_attribution_text: Optional[str] = None,
) -> None:
    fmt = fmt.lower()
    if fmt not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {fmt}")

    Path(out_path).parent.mkdir(parents=True, exist_ok=True)

    if fmt == "txt":
        Path(out_path).write_text(text, encoding="utf-8")
        return

    if fmt == "pdf":
        _export_pdf(
            text,
            out_path,
            title=title,
            author=author,
            font=pdf_font,
            font_size=pdf_font_size,
            margin=pdf_margin,
            cover_path=cover_image,
            cover_image_bytes=cover_image_bytes,
            cover_fullpage=pdf_cover_fullpage,
            page_size=pdf_page_size,
            orientation=pdf_orientation,
            first_page_cover_only=pdf_first_page_cover_only,
            font_file=pdf_font_file,
            header_text=pdf_header,
            footer_text=pdf_footer,
            toc_segments=(segments if auto_toc else None),
            append_attribution=pdf_append_attribution,
            attribution_text=pdf_attribution_text,
        )
        return

    if fmt == "epub":
        _export_epub(
            text,
            out_path,
            title=title,
            author=author,
            cover_image=cover_image,
            cover_image_bytes=cover_image_bytes,
            css_file=epub_css_file,
            css_text=epub_css_text,
            segments=segments,
            metadata=metadata,
        )
        return

    if fmt == "srt":
        _export_srt(text, out_path, segments=segments)
        return

    if fmt == "vtt":
        _export_vtt(text, out_path, segments=segments)
        return

    if fmt == "json":
        _export_json(
            text,
            out_path,
            segments=segments,
            words=words,
            title=title,
            author=author,
            metadata=metadata,
        )
        return

    if fmt == "md":
        _export_md(text, out_path, title=title, author=author)
        return

    if fmt == "docx":
        _export_docx(
            text,
            out_path,
            title=title,
            author=author,
            cover_image=cover_image,
            cover_image_bytes=cover_image_bytes,
            cover_first=docx_cover_first,
            cover_width_inches=docx_cover_width_inches,
            attribution_text=None,
            footer_text=docx_footer_text,
            footer_page_number=bool(docx_footer_include_page_number),
        )
        return

    # Kindle family: convert via Calibre's ebook-convert from an EPUB
    _export_kindle(
        text,
        out_path,
        target_fmt=fmt,
        title=title,
        author=author,
        cover_image=cover_image,
        css_file=epub_css_file,
        css_text=epub_css_text,
    )


def _export_pdf(
    text: str,
    out_path: str,
    title: Optional[str],
    author: Optional[str],
    font: str = "Arial",
    font_size: int = 12,
    margin: int = 15,
    cover_path: Optional[str] = None,
    cover_image_bytes: Optional[bytes] = None,
    cover_fullpage: bool = False,
    first_page_cover_only: bool = False,
    page_size: str = "A4",
    orientation: str = "portrait",
    font_file: Optional[str] = None,
    header_text: Optional[str] = None,
    footer_text: Optional[str] = None,
    toc_segments: Optional[list[dict]] = None,
    append_attribution: bool = False,
    attribution_text: Optional[str] = None,
):
    try:
        from fpdf import FPDF  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "PDF export requires 'fpdf2'. Install with: pip install fpdf2"
        ) from e

    orient_flag = "P" if str(orientation).lower().startswith("p") else "L"
    try:

        class PDFDoc(FPDF):
            def header(self_inner):
                if header_text:
                    self_inner.set_font("Arial", size=9)
                    self_inner.cell(0, 8, header_text, 0, 1, "C")

            def footer(self_inner):
                if footer_text:
                    self_inner.set_y(-12)
                    self_inner.set_font("Arial", size=9)
                    self_inner.cell(
                        0, 10, f"{footer_text}  |  {self_inner.page_no()}", 0, 0, "C"
                    )

        pdf = PDFDoc(orientation=orient_flag, format=page_size)
    except Exception:
        pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=margin)
    pdf.add_page()
    pdf.set_title(title or "Transcript")
    if author:
        pdf.set_author(author)
    # Custom font embedding
    if font_file:
        ff = Path(font_file)
        if not ff.exists():
            raise FileNotFoundError(f"PDF font file not found: {font_file}")
        try:
            if hasattr(pdf, "add_font"):
                pdf.add_font("Embedded", "", str(ff), uni=True)  # type: ignore[arg-type]
                pdf.set_font("Embedded", size=font_size)
            else:
                pdf.set_font(font, size=font_size)
        except Exception:
            # Fallback to built-in font on any failure
            pdf.set_font(font, size=font_size)
    else:
        pdf.set_font(font, size=font_size)
    # Optional cover on first page
    if cover_path or cover_image_bytes:
        fd, tmp_cover = tempfile.mkstemp(prefix="cover_", suffix=".jpg")
        os.close(fd)
        try:
            if cover_path:
                p = Path(cover_path)
                if not p.exists():
                    raise FileNotFoundError(f"Cover image not found: {cover_path}")
                data = _prepare_cover_bytes(p)
            else:
                data = cover_image_bytes or b""
            Path(tmp_cover).write_bytes(data)
            if cover_fullpage:
                # Full-page: draw image across page width, minimal margins
                try:
                    page_w = pdf.w  # type: ignore[attr-defined]
                except Exception:
                    page_w = 210  # A4 width mm fallback
                pdf.image(tmp_cover, x=0, y=0, w=page_w)
                # Start transcript on a new page
                pdf.add_page()
            else:
                try:
                    epw = pdf.epw  # type: ignore[attr-defined]
                except Exception:
                    try:
                        epw = pdf.w - 2 * pdf.l_margin  # type: ignore[attr-defined]
                    except Exception:
                        epw = 150
                try:
                    x = pdf.l_margin  # type: ignore[attr-defined]
                except Exception:
                    x = 10
                pdf.image(tmp_cover, x=x, y=None, w=epw)
                pdf.ln(10)
                if first_page_cover_only:
                    pdf.add_page()
        finally:
            try:
                Path(tmp_cover).unlink(missing_ok=True)
            except Exception:
                pass
    # Auto TOC page when segments provided
    if toc_segments:
        pdf.add_page()
        pdf.set_font(font, size=font_size + 2)
        pdf.cell(0, 10, "Table of Contents", ln=1)
        pdf.set_font(font, size=font_size)
        for seg in toc_segments:
            start = _format_timestamp(seg.get("start", 0.0))
            spk = seg.get("speaker")
            label = f"[{start}] {spk + ': ' if spk else ''}{(seg.get('text', '')[:60]).strip()}"
            pdf.multi_cell(0, 6, label)

    # Basic wrapping: split on double newlines as paragraphs
    for para in text.split("\n\n"):
        for line in para.splitlines():
            pdf.multi_cell(0, 8, line)
        pdf.ln(4)
    if append_attribution and attribution_text:
        try:
            pdf.add_page()
            pdf.set_font(font, size=font_size + 2)
            pdf.cell(0, 10, "Attribution", ln=1, align="C")
            pdf.ln(4)
            pdf.set_font(font, size=font_size)
            pdf.multi_cell(0, 8, attribution_text, align="C")
        except Exception:
            pass
    pdf.output(out_path)


def _export_epub(
    text: str,
    out_path: str,
    title: Optional[str],
    author: Optional[str],
    cover_image: Optional[str] = None,
    cover_image_bytes: Optional[bytes] = None,
    css_file: Optional[str] = None,
    css_text: Optional[str] = None,
    segments: Optional[list[dict]] = None,
    metadata: Optional[dict] = None,
):
    try:
        from ebooklib import epub  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "EPUB export requires 'ebooklib'. Install with: pip install ebooklib"
        ) from e

    book = epub.EpubBook()
    book.set_title(title or "Transcript")
    if author:
        book.add_author(author)
    # Basic metadata helpful for KDP ingestion
    try:
        lang = None
        if metadata and isinstance(metadata, dict):
            lang = metadata.get("language") or metadata.get("lang")
            desc = metadata.get("description")
            subj = metadata.get("keywords") or metadata.get("subjects")
            if lang:
                try:
                    book.set_language(lang)
                except Exception:
                    pass
            if desc:
                try:
                    book.add_metadata("DC", "description", str(desc))
                except Exception:
                    pass
            if subj:
                try:
                    # join list/tuple or accept string
                    if isinstance(subj, (list, tuple)):
                        subj_val = ", ".join(map(str, subj))
                    else:
                        subj_val = str(subj)
                    book.add_metadata("DC", "subject", subj_val)
                except Exception:
                    pass
    except Exception:
        pass
    if cover_image:
        p = Path(cover_image)
        if not p.exists():
            raise FileNotFoundError(f"Cover image not found: {cover_image}")
        if p.suffix.lower() not in {".jpg", ".jpeg", ".png"}:
            raise ValueError("Cover image must be a .jpg, .jpeg or .png file")
        try:
            data = _prepare_cover_bytes(
                p
            )  # if Pillow available, resizes and re-encodes
        except Exception:
            data = p.read_bytes()
        book.set_cover(p.name, data)
    elif cover_image_bytes:
        book.set_cover("cover.jpg", cover_image_bytes)

    # Simple HTML content with optional embedded CSS
    combined_css = css_text or None
    if css_file:
        css_path = Path(css_file)
        if not css_path.exists():
            raise FileNotFoundError(f"EPUB CSS file not found: {css_file}")
        css_from_file = css_path.read_text(encoding="utf-8")
        combined_css = (combined_css or "") + css_from_file

    head = (
        "<head><meta charset='utf-8'/>"
        + (f"<style>{html.escape(combined_css)}</style>" if combined_css else "")
        + "</head>"
    )
    html_parts = [head, "<body>", "<h1>" + (title or "Transcript") + "</h1>"]
    for para in text.split("\n\n"):
        html_parts.append(
            "<p>" + "<br/>".join(html.escape(p) for p in para.splitlines()) + "</p>"
        )
    html_parts.append("</body>")
    content = "\n".join(html_parts)
    chapters = []
    if segments:
        # Split into small chapters by segments, every ~10 minutes or speaker changes
        cur_html = [head, "<body>"]
        cur_len = 0
        cur_idx = 1
        for seg in segments:
            start = _format_timestamp(seg.get("start", 0.0))
            spk = seg.get("speaker")
            text_seg = seg.get("text", "")
            cur_html.append(
                f"<h2 id='seg-{cur_idx}'>[{start}] {html.escape(spk + ': ' if spk else '')}</h2>"
            )
            cur_html.append("<p>" + html.escape(text_seg) + "</p>")
            cur_len += len(text_seg)
            if cur_len > 4000:  # rough size threshold
                ch = epub.EpubHtml(
                    title=f"Section {cur_idx}",
                    file_name=f"section_{cur_idx}.xhtml",
                    lang="en",
                )
                ch.content = "\n".join(cur_html + ["</body>"])
                book.add_item(ch)
                chapters.append(ch)
                cur_html = [head, "<body>"]
                cur_len = 0
                cur_idx += 1
        if cur_html and len(cur_html) > 2:
            ch = epub.EpubHtml(
                title=f"Section {cur_idx}",
                file_name=f"section_{cur_idx}.xhtml",
                lang="en",
            )
            ch.content = "\n".join(cur_html + ["</body>"])
            book.add_item(ch)
            chapters.append(ch)
    else:
        c = epub.EpubHtml(title="Transcript", file_name="transcript.xhtml", lang="en")
        c.content = content
        book.add_item(c)
        chapters.append(c)

    book.toc = tuple(chapters)
    # Add navigation items with compatibility across ebooklib versions
    try:
        book.add_item(epub.EpubNav())
    except Exception:
        try:
            book.add_item(epub.EpubNavi())
        except Exception:
            pass
    try:
        book.add_item(epub.EpubNcx())
    except Exception:
        try:
            book.add_item(epub.EpubNCX())
        except Exception:
            pass
    book.spine = ["nav", *chapters]
    epub.write_epub(out_path, book)


def _export_docx(
    text: str,
    out_path: str,
    title: Optional[str],
    author: Optional[str],
    cover_image: Optional[str] = None,
    cover_image_bytes: Optional[bytes] = None,
    cover_first: bool = False,
    cover_width_inches: Optional[float] = None,
    attribution_text: Optional[str] = None,
    footer_text: Optional[str] = None,
    footer_page_number: bool = True,
) -> None:
    try:
        import docx  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "DOCX export requires 'python-docx'. Install with: pip install python-docx or pip install podcast-transcriber[docx]"
        ) from e
    doc = docx.Document()

    # Helper to add cover page
    def _add_cover(tmp_path: str):
        doc.add_page_break()
        try:
            if cover_width_inches:
                doc.add_picture(tmp_path, width=Inches(float(cover_width_inches)))
            else:
                doc.add_picture(tmp_path, width=Inches(6))
            doc.add_page_break()
        except Exception:
            pass

    # Cover support: either file path or raw bytes written to a temp file
    if cover_image or cover_image_bytes:
        tmp_path = None
        try:
            if cover_image:
                p = Path(cover_image)
                if p.exists():
                    tmp_path = str(p)
            elif cover_image_bytes:
                fd, tmp = tempfile.mkstemp(prefix="cover_", suffix=".jpg")
                os.close(fd)
                Path(tmp).write_bytes(cover_image_bytes)  # type: ignore[arg-type]
                tmp_path = tmp
            # order: cover first or after title/author
            if tmp_path and cover_first:
                _add_cover(tmp_path)
        finally:
            if cover_image_bytes and tmp_path:
                try:
                    Path(tmp_path).unlink(missing_ok=True)
                except Exception:
                    pass
    # Title and author
    if title:
        doc.add_heading(title, level=0)
    if author:
        p = doc.add_paragraph()
        run = p.add_run(author)
        run.italic = True
    # Cover after title/author when not cover_first
    if (cover_image or cover_image_bytes) and not cover_first:
        tmp_path2 = None
        try:
            if cover_image:
                p = Path(cover_image)
                if p.exists():
                    tmp_path2 = str(p)
            elif cover_image_bytes:
                fd, tmp = tempfile.mkstemp(prefix="cover_", suffix=".jpg")
                os.close(fd)
                Path(tmp).write_bytes(cover_image_bytes)  # type: ignore[arg-type]
                tmp_path2 = tmp
            if tmp_path2:
                _add_cover(tmp_path2)
        finally:
            if cover_image_bytes and tmp_path2:
                try:
                    Path(tmp_path2).unlink(missing_ok=True)
                except Exception:
                    pass
    # Body
    for para in text.split("\n\n"):
        doc.add_paragraph(para.strip())
        doc.add_paragraph("")
    if attribution_text:
        try:
            doc.add_page_break()
            doc.add_heading("Attribution", level=1)
            doc.add_paragraph(attribution_text)
        except Exception:
            pass
    # Footer: optional text and page number field
    try:
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore

        for section in doc.sections:
            ftr = section.footer
            p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            if footer_text:
                p.add_run(str(footer_text) + "  ")
            if footer_page_number:
                r = p.add_run()
                r._r.append(OxmlElement("w:fldChar"))
                r._r[-1].set(qn("w:fldCharType"), "begin")
                instr = OxmlElement("w:instrText")
                instr.text = " PAGE "
                r._r.append(instr)
                r._r.append(OxmlElement("w:fldChar"))
                r._r[-1].set(qn("w:fldCharType"), "separate")
                r._r.append(OxmlElement("w:fldChar"))
                r._r[-1].set(qn("w:fldCharType"), "end")
    except Exception:
        pass
    doc.save(out_path)


def export_book(
    chapters: list[dict],
    out_path: str,
    fmt: str,
    title: Optional[str] = None,
    author: Optional[str] = None,
    cover_image: Optional[str] = None,
    cover_image_bytes: Optional[bytes] = None,
    metadata: Optional[dict] = None,
    epub_css_file: Optional[str] = None,
    epub_css_text: Optional[str] = None,
) -> None:
    """Export a multi-chapter book.

    chapters: list of {"title": str, "text": str}
    fmt: one of epub, docx, md, txt, pdf (basic)
    """
    fmt = fmt.lower()
    if fmt == "epub":
        try:
            from ebooklib import epub  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "EPUB export requires 'ebooklib'. Install with: pip install ebooklib"
            ) from e
        book = epub.EpubBook()
        book.set_title(title or "Podcast Book")
        if author:
            book.add_author(author)
        if metadata and isinstance(metadata, dict):
            try:
                lang = metadata.get("language") or metadata.get("lang")
                if lang:
                    book.set_language(lang)
                desc = metadata.get("description")
                if desc:
                    book.add_metadata("DC", "description", str(desc))
                subj = metadata.get("keywords") or metadata.get("subjects")
                if subj:
                    if isinstance(subj, (list, tuple)):
                        subj_val = ", ".join(map(str, subj))
                    else:
                        subj_val = str(subj)
                    book.add_metadata("DC", "subject", subj_val)
            except Exception:
                pass
        # Cover
        if cover_image:
            p = Path(cover_image)
            if p.exists():
                try:
                    data = _prepare_cover_bytes(p)
                except Exception:
                    data = p.read_bytes()
                book.set_cover(p.name, data)
        elif cover_image_bytes:
            book.set_cover("cover.jpg", cover_image_bytes)
        # CSS
        combined_css = epub_css_text or None
        if epub_css_file:
            css_path = Path(epub_css_file)
            if not css_path.exists():
                raise FileNotFoundError(f"EPUB CSS file not found: {epub_css_file}")
            css_from_file = css_path.read_text(encoding="utf-8")
            combined_css = (combined_css or "") + css_from_file
        head = (
            "<head><meta charset='utf-8'/>"
            + (f"<style>{html.escape(combined_css)}</style>" if combined_css else "")
            + "</head>"
        )
        epub_chapters = []
        for idx, ch in enumerate(chapters, start=1):
            node = epub.EpubHtml(
                title=str(ch.get("title") or f"Chapter {idx}"),
                file_name=f"chapter_{idx}.xhtml",
                lang="en",
            )
            parts = [
                head,
                "<body>",
                f"<h1>{html.escape(str(ch.get('title') or f'Chapter {idx}'))}</h1>",
            ]
            for para in str(ch.get("text", "")).split("\n\n"):
                parts.append(
                    "<p>"
                    + "<br/>".join(html.escape(p) for p in para.splitlines())
                    + "</p>"
                )
            parts.append("</body>")
            node.content = "\n".join(parts)
            book.add_item(node)
            epub_chapters.append(node)
        book.toc = tuple(epub_chapters)
        try:
            book.add_item(epub.EpubNav())
        except Exception:
            try:
                book.add_item(epub.EpubNavi())
            except Exception:
                pass
        try:
            book.add_item(epub.EpubNcx())
        except Exception:
            try:
                book.add_item(epub.EpubNCX())
            except Exception:
                pass
        book.spine = ["nav", *epub_chapters]
        epub.write_epub(out_path, book)
        return
    elif fmt == "docx":
        try:
            import docx  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "DOCX export requires 'python-docx'. Install with: pip install python-docx or pip install podcast-transcriber[docx]"
            ) from e
        d = docx.Document()
        if title:
            d.add_heading(title, level=0)
        if author:
            p = d.add_paragraph(author)
        if cover_image or cover_image_bytes:
            from docx.shared import Inches  # type: ignore

            tmp_path = None
            try:
                if cover_image:
                    p = Path(cover_image)
                    if p.exists():
                        tmp_path = str(p)
                elif cover_image_bytes:
                    fd, tmp = tempfile.mkstemp(prefix="cover_", suffix=".jpg")
                    os.close(fd)
                    Path(tmp).write_bytes(cover_image_bytes)  # type: ignore[arg-type]
                    tmp_path = tmp
                if tmp_path:
                    d.add_page_break()
                    try:
                        d.add_picture(tmp_path, width=Inches(6))
                        d.add_page_break()
                    except Exception:
                        pass
            finally:
                if cover_image_bytes and tmp_path:
                    try:
                        Path(tmp_path).unlink(missing_ok=True)
                    except Exception:
                        pass
        for idx, ch in enumerate(chapters, start=1):
            d.add_heading(str(ch.get("title") or f"Chapter {idx}"), level=1)
            for para in str(ch.get("text", "")).split("\n\n"):
                d.add_paragraph(para.strip())
            d.add_page_break()
        d.save(out_path)
        return
    elif fmt in {"md", "txt"}:
        lines = []
        if title:
            if fmt == "md":
                lines += [f"# {title}", ""]
            else:
                lines += [title.upper(), ""]
        if author:
            if fmt == "md":
                lines += [f"_by {author}_", ""]
            else:
                lines += [f"by {author}", ""]
        for idx, ch in enumerate(chapters, start=1):
            ch_title = str(ch.get("title") or f"Chapter {idx}")
            if fmt == "md":
                lines += [f"\n\n## {ch_title}", ""]
            else:
                lines += ["\n\n" + ch_title, ""]
            lines.append(str(ch.get("text", "")).strip())
        Path(out_path).write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
        return
    elif fmt == "pdf":
        # Very simple PDF: chapters separated with headings
        try:
            from fpdf import FPDF  # type: ignore
        except Exception as e:
            raise RuntimeError(
                "PDF export requires 'fpdf2'. Install with: pip install fpdf2"
            ) from e
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        if title:
            pdf.set_font("Arial", size=16)
            pdf.multi_cell(0, 10, title)
            pdf.ln(4)
        if author:
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 8, f"by {author}")
            pdf.ln(6)
        for idx, ch in enumerate(chapters, start=1):
            pdf.set_font("Arial", size=14)
            pdf.multi_cell(0, 9, str(ch.get("title") or f"Chapter {idx}"))
            pdf.ln(2)
            pdf.set_font("Arial", size=12)
            for para in str(ch.get("text", "")).split("\n\n"):
                for line in para.splitlines():
                    pdf.multi_cell(0, 8, line)
                pdf.ln(3)
            if idx < len(chapters):
                pdf.add_page()
        pdf.output(out_path)
        return
    else:
        raise ValueError(f"Unsupported book format: {fmt}")


def _export_kindle(
    text: str,
    out_path: str,
    target_fmt: str,
    title: Optional[str],
    author: Optional[str],
    cover_image: Optional[str] = None,
    css_file: Optional[str] = None,
    css_text: Optional[str] = None,
):
    conv = shutil.which("ebook-convert")
    if not conv:
        raise RuntimeError(
            "Kindle formats require Calibre's 'ebook-convert' to be installed and on PATH."
        )
    # Make a temp EPUB then convert to target
    fd, tmp_epub = tempfile.mkstemp(prefix="transcript_", suffix=".epub")
    os.close(fd)
    try:
        # Calibre does not reliably support 'azw' as an output extension in all builds.
        # Prefer converting to 'azw3' and then write/copy to the requested path when 'azw' is requested.
        use_ext = target_fmt.lower()
        tmp_out = None
        if use_ext == "azw":
            # Convert to a temp .azw3, then copy to the .azw destination for compatibility
            fd2, tmp_out = tempfile.mkstemp(prefix="kindle_", suffix=".azw3")
            os.close(fd2)
        else:
            tmp_out = out_path
        _export_epub(
            text,
            tmp_epub,
            title=title,
            author=author,
            cover_image=cover_image,
            css_file=css_file,
            css_text=css_text,
        )
        # Choose destination for conversion
        dest = tmp_out or out_path
        subprocess.run([conv, tmp_epub, dest], check=True)
        if use_ext == "azw":
            # Copy converted azw3 bytes into expected .azw file name
            try:
                data = Path(dest).read_bytes()
                Path(out_path).write_bytes(data)
            except Exception:
                # Best-effort: if copy fails, leave the file absent
                raise
    finally:
        try:
            Path(tmp_epub).unlink(missing_ok=True)
        except Exception:
            pass
        if tmp_out and tmp_out != out_path:
            try:
                Path(tmp_out).unlink(missing_ok=True)
            except Exception:
                pass


def _prepare_cover_bytes(path: Path) -> bytes:
    Image = None
    try:
        import importlib

        Image = importlib.import_module("PIL.Image")  # type: ignore
    except Exception:
        try:
            from PIL import Image as PILImage  # type: ignore

            Image = PILImage
        except Exception:
            return path.read_bytes()

    max_size = (1600, 2560)
    try:
        with Image.open(str(path)) as img:
            img = img.convert("RGB")
            img.thumbnail(max_size, Image.LANCZOS)
            import io

            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=85)
            data = buf.getvalue()
            if not data:
                return path.read_bytes()
            return data
    except Exception:
        # If PIL cannot identify or process the image, fall back to raw file bytes
        try:
            return path.read_bytes()
        except Exception:
            return b""


def _format_timestamp(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02}:{m:02}:{s:02},{ms:03}"


def _coalesce_segments(text: str, segments: Optional[list[dict]]) -> list[dict]:
    if segments:
        # ensure required keys
        out = []
        for seg in segments:
            if "start" in seg and "end" in seg and "text" in seg:
                out.append(
                    {
                        "start": float(seg["start"]),
                        "end": float(seg["end"]),
                        "text": str(seg["text"]).strip(),
                        "speaker": seg.get("speaker"),
                    }
                )
        if out:
            return out
    # Fallback: single segment covering unknown duration
    return [{"start": 0.0, "end": 0.0, "text": text.strip()}]


def _export_srt(
    text: str, out_path: str, segments: Optional[list[dict]] = None
) -> None:
    segs = _coalesce_segments(text, segments)
    lines = []
    for i, seg in enumerate(segs, start=1):
        start = _format_timestamp(seg.get("start", 0.0))
        end = _format_timestamp(seg.get("end", 0.0))
        caption = seg.get("text", "").strip()
        spk = seg.get("speaker")
        if spk:
            caption = f"{spk}: {caption}"
        lines.append(str(i))
        lines.append(f"{start} --> {end}")
        lines.append(caption)
        lines.append("")
    Path(out_path).write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _export_vtt(
    text: str, out_path: str, segments: Optional[list[dict]] = None
) -> None:
    segs = _coalesce_segments(text, segments)
    lines = ["WEBVTT", ""]
    for seg in segs:
        start = _format_timestamp(seg.get("start", 0.0)).replace(",", ".")
        end = _format_timestamp(seg.get("end", 0.0)).replace(",", ".")
        caption = seg.get("text", "").strip()
        spk = seg.get("speaker")
        if spk:
            caption = f"{spk}: {caption}"
        lines.append(f"{start} --> {end}")
        lines.append(caption)
        lines.append("")
    Path(out_path).write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def _export_json(
    text: str,
    out_path: str,
    segments: Optional[list[dict]] = None,
    words: Optional[list[dict]] = None,
    title: Optional[str] = None,
    author: Optional[str] = None,
    metadata: Optional[dict] = None,
) -> None:
    import json

    payload = {
        "title": title or "Transcript",
        "author": author,
        "text": text,
        "segments": _coalesce_segments(text, segments),
    }
    if words:
        payload["words"] = words
    if metadata:
        payload["source"] = metadata
    Path(out_path).write_text(
        json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8"
    )


def _export_md(
    text: str, out_path: str, title: Optional[str], author: Optional[str]
) -> None:
    lines = []
    if title:
        lines.append(f"# {title}")
        lines.append("")
    if author:
        lines.append(f"_by {author}_")
        lines.append("")
    for para in text.split("\n\n"):
        lines.append(para.strip())
        lines.append("")
    Path(out_path).write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")
